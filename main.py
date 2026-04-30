import os
import re
import uuid
import asyncio
import datetime as dt
import unicodedata
from typing import Optional, Dict, Any
import glob 
import json
import smtplib
import random
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

# Config Zoho
BREVO_LOGIN = os.getenv("BREVO_LOGIN", "")
BREVO_PASSWORD = os.getenv("BREVO_PASSWORD", "")
SENDER_EMAIL = os.getenv("SENDER_EMAIL", "contact@mycvcopilote.com")

# Stockage temporaire des codes de vérification
# format : { "email@ex.com": {"code": "123456", "expires": datetime} }
email_verification_codes: Dict[str, Dict] = {}
# Rate limiting par IP — max 5 tentatives/heure sur /start
_ip_attempts: Dict[str, list] = {}

def _check_ip_rate_limit(ip: str):
    now = dt.datetime.utcnow()
    history = _ip_attempts.get(ip, [])
    history = [t for t in history if (now - t).seconds < 3600]
    if len(history) >= 5:
        raise HTTPException(
            status_code=429,
            detail="Trop de tentatives. Réessaie dans 1 heure."
        )
    history.append(now)
    _ip_attempts[ip] = history
    
from pydantic import BaseModel

class EmailRequest(BaseModel):
    email: str
    turnstile_token: str = ""  # Cloudflare Turnstile CAPTCHA token

class VerifyCodeRequest(BaseModel):
    email: str
    code: str

from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware

import stripe
import subprocess
import shutil

from openai import OpenAI
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

APP_URL = os.getenv("APP_URL", "")  # ex: https://mycvcopilote.netlify.app
STRIPE_SECRET = os.getenv("STRIPE_SECRET") or os.getenv("STRIPE_SECRET_KEY", "")
STRIPE_WEBHOOK_SECRET = os.getenv("STRIPE_WEBHOOK_SECRET", "")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
PUBLIC_BASE_DOWNLOAD = os.getenv("PUBLIC_BASE_DOWNLOAD", "")  # ex: https://mycvcopilote-api.onrender.com/download
TURNSTILE_SECRET = os.getenv("TURNSTILE_SECRET_KEY", "")  # Cloudflare Turnstile

client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None

from pypdf import PdfReader

def normalize_email(email: str) -> str:
    """
    Normalise un email :
    - lowercase
    - supprime les alias Gmail/Outlook (partie après +)
      ex: user+alias@gmail.com → user@gmail.com
    """
    if not email:
        return email
    email = email.strip().lower()
    if "@" in email:
        local, domain = email.split("@", 1)
        # Supprimer le sous-adressage (+ alias) pour Gmail, Outlook, etc.
        local = local.split("+")[0]
        email = f"{local}@{domain}"
    return email


async def verify_turnstile(token: str, ip: str = "") -> bool:
    """
    Vérifie un token Cloudflare Turnstile côté serveur.
    Retourne True si valide, False sinon.
    Si TURNSTILE_SECRET_KEY n'est pas configuré → bypass (mode dev).
    """
    if not TURNSTILE_SECRET:
        return True  # Bypass si pas configuré
    if not token:
        return False
    try:
        import httpx
        async with httpx.AsyncClient(timeout=5.0) as c:
            resp = await c.post(
                "https://challenges.cloudflare.com/turnstile/v0/siteverify",
                data={"secret": TURNSTILE_SECRET, "response": token, "remoteip": ip},
            )
            data = resp.json()
            return bool(data.get("success"))
    except Exception:
        return True  # En cas d'erreur réseau, on laisse passer


def strip_padding(text: str, is_activity: bool = False) -> str:
    """
    Supprime les subordonnées participiales inventées en fin de bullet ou d'activité.
    Appliqué ligne par ligne après chaque appel LLM.
    """
    if not text:
        return text

    # Patterns de rembourrage en fin de bullet (après virgule)
    BULLET_PADDING = [
        r",\s*(assurant|renforçant|optimisant|consolidant|garantissant|facilitant"
        r"|maximisant|sécurisant|fiabilisant|contribuant ainsi|mettant en évidence"
        r"|présentant des recommandations|proposant des recommandations"
        r"|soutenant|tout en restant|tout en renforçant|tout en assurant"
        r"|tout en optimisant|afin d'assurer|dans l'objectif de"
        r"|augmentant sans précédent|enrichissant|développant|favorisant"
        r"|mobilisant|démontrant|approfondissant|mettant en avant"
        r"|incluant la formation|incluant le coaching|veillant à)[^.]*",
        r",\s*(permettant|contribuant\s+\w*\s*à|participant à l'amélioration)[^.]*",
        r",?\s+afin de (renforcer|optimiser|assurer|garantir|consolider|maximiser)[^.]*",
        r"\s+en veillant à [^.]*",
        r";\s*\w[\w\s]*développée?\.?$",
        r";\s*\w[\w\s]*acquise?\.?$",
        r";\s*\w[\w\s]*renforcée?\.?$",
        r";\s*[^;.]{3,40}développée?\.?$",
    ]

    # Patterns supplémentaires pour les activités - SEULEMENT après virgule
    ACTIVITY_PADDING = [
        r",\s+(développant|renforçant|favorisant|cultivant|enrichissant"
        r"|améliorant|acquérant|permettant de développer|favorisant le développement de"
        r"|approfondissant ainsi|démontrant une|mettant en avant|engagement continu sur"
        r"|développement significatif de|élargissant|stimulant"
        r"|favoriser l'esprit d'équipe|favorisant l'esprit d'équipe"
        r"|développement de compétences en|amélioration des compétences"
        r"|enrichissement des connaissances|développement de la confiance"
        r"|pour le bien-être personnel|pour maintenir le bien-être"
        r"|pour le travail d'équipe|et la stratégie"
        r"|exploration de différents domaines|approfondissement des connaissances"
        r"|pour acquisition de compétences|pour enrichir les compétences"
        r"|pour enrichir ses compétences|pour acquérir des compétences"
        r"|pour garantir le bon déroulement|pour assurer le bon déroulement"
        r"|pour approfondir les connaissances|pour développer les compétences"
        r"|pour renforcer les capacités|pour approfondir sa connaissance"
        r"|participations? régulières?|participations? actives?"
        r"|impliqué dans divers projets|impliqué dans différents projets"
        r"|ce qui renforce|ce qui stimule|ce qui développe|ce qui favorise)[^.]*",
        r"\s+pour (acquérir|enrichir|développer|renforcer|approfondir) [^.]*(?=\.)",
        r",\s+développe (des compétences|le leadership|l'esprit)[^.]*",
        r",\s+suivi (des tendances|régulier|actif)[^.]*",
    ]

    patterns = BULLET_PADDING + (ACTIVITY_PADDING if is_activity else [])

    for pattern in patterns:
        before_strip = text
        text = re.sub(pattern, "", text, flags=re.IGNORECASE)
        # ✅ Si le strip laisse moins de 8 mots, restaurer le texte original
        # (mieux vaut garder du contenu imparfait qu'une activité vide)
        if is_activity and len(text.strip().split()) < 8 and len(before_strip.strip().split()) >= 8:
            text = before_strip

    # Nettoyer la ponctuation résiduelle
    text = re.sub(r"\s*,\s*$", "", text)
    text = re.sub(r"\s*;\s*$", "", text)
    text = text.strip()
    if text and text[-1] not in ".!?":
        text = text + "."

    return text


def apply_strip_padding_to_cv(cv_text: str, payload: dict = None) -> str:
    """
    Applique strip_padding sur chaque bullet et chaque activité du texte CV structuré.
    Pour le droit : efface les DETAILS inventés quand l'input éducation ne contient rien.
    """
    if not cv_text:
        return cv_text

    # Pour le droit : construire la liste des blocs éducation ayant des détails réels
    edu_details_allowed = set()
    if payload and payload.get("sector", "") in ("droit", "legal"):
        raw_edu = payload.get("education", "")
        for block in raw_edu.split("\n\n"):
            lines_b = [l.strip() for l in block.strip().splitlines() if l.strip()]
            if len(lines_b) > 1:  # a des détails réels
                edu_details_allowed.add(lines_b[0][:30].lower())

    lines = cv_text.split("\n")
    result = []
    in_activities = False
    in_experiences = False
    in_education = False
    current_degree_key = None
    skip_details = False

    for line in lines:
        stripped = line.strip()

        if stripped == "ACTIVITIES:":
            in_activities = True
            in_experiences = False
            in_education = False
            result.append(line)
            continue
        elif stripped == "EXPERIENCES:":
            in_experiences = True
            in_activities = False
            in_education = False
            result.append(line)
            continue
        elif stripped == "EDUCATION:":
            in_education = True
            in_activities = False
            in_experiences = False
            result.append(line)
            continue
        elif stripped == "SKILLS:":
            in_activities = False
            in_experiences = False
            in_education = False
            result.append(line)
            continue

        # Droit : gérer les DETAILS inventés dans EDUCATION
        if in_education and stripped.startswith("DEGREE:"):
            degree_val = stripped.replace("DEGREE:", "").strip()[:30].lower()
            current_degree_key = degree_val
            # Vérifier si ce degré avait des détails réels dans l'input
            has_real_details = any(
                current_degree_key in allowed or allowed in current_degree_key
                for allowed in edu_details_allowed
            ) if edu_details_allowed else True
            skip_details = not has_real_details
            result.append(line)
            continue

        if in_education and stripped == "DETAILS:" and skip_details:
            result.append(line)
            continue

        if in_education and stripped.startswith("- ") and skip_details:
            # Remplacer par ligne minimale vide
            result.append("- ")
            skip_details = False  # une seule ligne minimale
            continue

        # Bullets d'expériences — strip_padding seulement
        if stripped.startswith("- ") and in_experiences:
            content = stripped[2:]
            cleaned = strip_padding(content, is_activity=False)
            result.append("- " + cleaned)
            continue

        # Lignes d'activités
        if in_activities and stripped and not stripped.startswith("ACTIVITIES:"):
            cleaned = strip_padding(stripped, is_activity=True)
            result.append(cleaned)
            continue

        result.append(line)

    return "\n".join(result)


def clean_cv_output(cv_text: str) -> str:
    if not cv_text:
        return ""
    lines = cv_text.replace("\r\n", "\n").split("\n")
    out = []
    for ln in lines:
        s = ln.strip()

        if s.startswith("```") or s == "```":
            continue

        if s in {".", "..", "...", "\"", "''", "\"\"", "\"\"\""}:
            continue

        low = s.lower()
        if low.startswith("cette version") or low.startswith("ce cv") or low.startswith("note :"):
            continue

        s = re.sub(r"\[[^\]]+\]", "", s).strip()
        s = re.sub(r"\*\*([^*]*)\*\*", r"\1", s).strip()
        s = re.sub(r"\*([^*]+)\*", r"\1", s).strip()
        s = re.sub(r"^#+\s*", "", s).strip()

        if not s:
            out.append("")
            continue

        out.append(s)
        
    return "\n".join(out).strip()

REQUIRED_SECTIONS = ["EDUCATION:", "EXPERIENCES:", "SKILLS:", "ACTIVITIES:"]

def clean_punctuation_text(text: str) -> str:
    if not text:
        return text

    text = re.sub(r"\s+,", ",", text)
    text = re.sub(r",\.", ".", text)
    text = re.sub(r"\.\.", ".", text)
    text = re.sub(r"\s+\.", ".", text)
    text = re.sub(r"\s+;", ";", text)
    text = re.sub(r";\.", ".", text)
    text = re.sub(r":\.", ".", text)
    text = re.sub(r",\s*$", "", text)   # ✅ enlève une virgule finale
    text = re.sub(r";\s*$", "", text)   # ✅ enlève un point-virgule final

    return text.strip()

def clean_activities_output(activities):
    if not activities:
        return activities

    cleaned = []

    banned_phrases = [
        "rigueur intellectuelle",
        "approfondissement",
        "ouverture d’esprit",
        "analyse critique",
        "passion pour",
        "intérêt pour",
    ]

    for act in activities:
        text = (act.get("text") or "").lower()

        # supprimer activités trop faibles
        if len(text) < 40:
            continue

        if any(b in text for b in banned_phrases):
            continue

        # Majuscule au début du texte
        if act.get("text"):
            t = act["text"].strip()
            if t:
                act = dict(act)
                act["text"] = t[0].upper() + t[1:]

        cleaned.append(act)

    # max 3
    return cleaned[:3]

def normalize_role_text(role: str) -> str:
    if not role:
        return role

    fixes = {
        "tuuteur": "Tuteur",
        "tuetrice": "Tutrice",
        "assitante": "Assistante",
        "stagaire": "Stagiaire",
    }

    # ✅ Rôles trop abrégés que le LLM génère parfois
    role_expansions = {
        r"^m&a$": "Stagiaire M&A",
        r"^finance$": "Stagiaire Finance",
        r"^comptabilité$": "Stagiaire Comptabilité",
        r"^marketing$": "Stagiaire Marketing",
        r"^rh$": "Stagiaire Ressources Humaines",
        r"^juridique$": "Stagiaire Juridique",
        r"^audit$": "Stagiaire Auditeur",
        r"^commercial$": "Chargé de mission commercial",
        r"^communication$": "Chargé de communication",
        r"^contrôle de gestion$": "Stagiaire Contrôle de Gestion",
        r"^controle de gestion$": "Stagiaire Contrôle de Gestion",
        r"^private equity$": "Stagiaire Private Equity",
        r"^développement$": "Chargé de développement",
        r"^developpement$": "Chargé de développement",
        r"^it$": "Stagiaire IT",
        r"^conseil$": "Consultant stagiaire",
        r"^supply chain$": "Stagiaire Supply Chain",
        r"^data$": "Stagiaire Data",
        r"^achats$": "Stagiaire Achats",
    }

    low = role.strip().lower()
    for pattern, replacement in role_expansions.items():
        if re.match(pattern, low):
            return replacement

    if low in fixes:
        return fixes[low]

    for bad, good in fixes.items():
        role = re.sub(rf"(?i)\b{re.escape(bad)}\b", good, role)

    return role.strip()

def has_all_sections(cv_text: str) -> bool:
    t = (cv_text or "")
    return all(sec in t for sec in REQUIRED_SECTIONS)

def safe_apply_llm_edit(old_text: str, new_text: str, payload: dict = None, allow_drop_exp: bool = False) -> str:
    new_clean = clean_cv_output(new_text)
    if not has_all_sections(new_clean):
        return old_text
    if not allow_drop_exp:
        old_role_count = old_text.count("\nROLE:")
        new_role_count = new_clean.count("\nROLE:")
        if new_role_count < old_role_count:
            return old_text
    new_clean = apply_strip_padding_to_cv(new_clean)
    return new_clean

def pdf_page_count(pdf_path: str) -> int:
    reader = PdfReader(pdf_path)
    return len(reader.pages)

def pdf_fill_ratio_first_page(pdf_path: str) -> float:
    """
    Heuristique simple : nombre de lignes non vides extraites de la page 1.
    Sert à détecter "trop vide" (beaucoup d'espace en bas).
    """
    reader = PdfReader(pdf_path)
    if len(reader.pages) == 0:
        return 0.0
    page = reader.pages[0]
    try:
        text = page.extract_text() or ""
    except Exception:
        text = ""
    if not text.strip():
        return 0.0
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    n = len(lines)

    # Calibrage recalibré sur 1cm marges, police 11pt, tableaux 2 colonnes
    # 22 lignes → 0.55 (CV très vide)
    # 45 lignes → 0.80 (CV moyen)
    # 60 lignes → 0.92 (CV bien rempli)
    # 68 lignes → 0.97 (CV pleine page)
    if n <= 22:
        return 0.55
    if n >= 68:
        return 0.97
    return 0.55 + (n - 22) * (0.42 / (68 - 22))

def llm_shrink_cv(cv_text: str) -> str:
    if not client:
        return cv_text

    # Count experiences to include in prompt
    exp_count = cv_text.count("\nROLE:")

    prompt = f"""
Tu dois rendre ce CV LÉGÈREMENT PLUS COURT pour tenir sur 1 page Word, SANS le casser.
Ce CV a {exp_count} expériences. Tu dois TOUTES les conserver.

Stratégie UNIQUE de réduction (dans cet ordre) :
1) Raccourcir les bullets (1-2 lignes au lieu de 2-3), en gardant les chiffres et faits clés
2) Limiter à 2 bullets les expériences secondaires (ancienne, courte, non liée au secteur)
3) Réduire DETAILS dans EDUCATION à 1-2 lignes max par diplôme
NE JAMAIS aller plus loin — ne jamais supprimer une expérience entière.

Règles ABSOLUES :
- Tu gardes exactement les sections : EDUCATION:, EXPERIENCES:, SKILLS:, ACTIVITIES:
- Tu conserves EXACTEMENT le format structuré de chaque expérience : ROLE:, COMPANY:, DATES:, LOCATION:, TYPE:, BULLETS: sur des lignes séparées.
- Tu ne supprimes JAMAIS une expérience entière — chaque ROLE: doit rester présent.
- Tu ne rajoutes AUCUN commentaire ni phrase méta.
- Tu n'inventes rien.
- INTERDIT ABSOLU : chaque bullet doit faire au minimum 8 mots et conserver tous les chiffres et faits précis.
- INTERDIT ABSOLU : chaque activité doit faire au minimum 8 mots.
- INTERDIT ABSOLU : ne jamais écrire un bullet à l'infinitif. Tout bullet commence par un verbe conjugué au passé composé.
- INTERDIT ABSOLU : ne jamais fusionner deux bullets en un seul.

Format ACTIVITIES : 1 activité par ligne, sans puce, forme "Activité : description courte."
Sortie : UNIQUEMENT le CV complet.

CV :
\"\"\"{cv_text}\"\"\"
"""
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
    )
    return resp.choices[0].message.content.strip()

def llm_expand_cv(cv_text: str) -> str:
    if not client:
        return cv_text

    prompt = f"""
Tu dois rendre ce CV plus dense pour remplir une page Word complète.

OBJECTIF PRINCIPAL : ajouter du contenu réel et professionnel pour remplir la page.

Tu peux uniquement :
1) Passer à 3 bullets pour les expériences qui n'en ont que 2 — en développant ce qui est déjà là de manière plus complète.
2) Développer les bullets existants : ajouter contexte, méthode, précision sur ce qui est déjà mentionné.
3) Si une formation a des détails, les développer légèrement.
4) Reformuler les bullets pour qu'ils soient plus complets et professionnels.

INTERDIT ABSOLU :
- Inventer une nouvelle expérience, un nouveau poste, un nouveau projet.
- Ajouter un chiffre non fourni.
- Ajouter un outil non fourni.
- Terminer un bullet par une subordonnée participiale ("assurant", "contribuant à", "favorisant", "renforçant", "permettant", "garantissant", "mobilisant", "démontrant").
- Laisser un fragment de phrase sans verbe principal.

Sortie : UNIQUEMENT le CV complet sans commentaire.

CV :
{cv_text}
"""
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
    )
    return resp.choices[0].message.content.strip()

def llm_expand_cv_droit(cv_text: str) -> str:
    if not client:
        return cv_text

    prompt = f"""
Tu dois rendre ce CV DROIT légèrement plus dense pour mieux remplir 1 page Word,
sans inventer la moindre information.

Règles ABSOLUES :
- Tu gardes exactement les sections : EDUCATION:, EXPERIENCES:, SKILLS:, ACTIVITIES:
- Tu ne rajoutes AUCUN commentaire.
- Tu n'inventes rien.
- Tu ne rajoutes aucune mission, aucun chiffre, aucun outil, aucune matière, aucune activité.
- Tu ne transformes jamais un job étudiant en expérience juridique.
- Tu ne rajoutes jamais de bénéfice implicite, de finalité, d'optimisation ou d'impact.
- Tu conserves absolument tous les éléments académiques explicites déjà présents, notamment :
  mémoire, concours, moot court, mock trial, distinctions, matières, certifications.

Tu peux uniquement :
1) reformuler légèrement une ou deux bullets existantes pour qu'elles soient un peu plus complètes,
2) laisser 3 bullets sur l'expérience la plus pertinente si elles existent déjà,
3) enrichir très légèrement UNE activité existante sans ajouter de fait nouveau,
4) conserver davantage de détails académiques déjà présents dans EDUCATION.

Style :
- sobre
- académique
- crédible
- factuel
- professionnel


RÈGLE LONGUEUR BULLETS : Chaque bullet doit faire entre 20 et 40 mots (environ 1,5 à 2 lignes dans Word). Un bullet trop court (< 15 mots) est insuffisant — enrichis-le avec le contexte, le périmètre ou la méthode utilisée, sans inventer de chiffres.
RÈGLE ACTIVITÉS : Chaque activité doit faire 10-20 mots minimum. Ne jamais laisser une activité en 1-2 mots seuls.

Format ACTIVITIES : 1 activité par ligne, sans puce, forme "Activité : description courte."
Sortie : UNIQUEMENT le CV complet.

CV :
\"\"\"{cv_text}\"\"\"
"""
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
    )
    return resp.choices[0].message.content.strip()

def llm_expand_cv_audit(cv_text: str) -> str:
    if not client:
        return cv_text

    prompt = f"""
Tu dois rendre ce CV AUDIT légèrement plus dense pour mieux remplir une page Word, sans inventer d’information précise.

Règles ABSOLUES :
- Tu gardes exactement les sections : EDUCATION:, EXPERIENCES:, SKILLS:, ACTIVITIES:
- Tu ne rajoutes aucun commentaire.
- Tu n’inventes aucun outil, aucun chiffre, aucune matière, aucune mission nouvelle, aucun pays, aucune compétition, aucun événement.
- Tu peux rendre une mission plus professionnelle et plus valorisante si cela reste très crédible.
- Tu peux faire ressortir une qualité simple comme rigueur, précision, organisation ou fiabilité si elle découle logiquement du texte source.

Tu peux uniquement :
1) ajouter 1 bullet à l’expérience la plus pertinente si elle n’en a que 2,
2) reformuler légèrement les bullets pour les rendre plus professionnelles,
3) conserver davantage de détails académiques déjà présents,
4) enrichir légèrement une activité existante sans ajouter de fait précis.

Interdictions :
- pas de “travaux d’audit”, “contrôle interne”, “conformité”, “états financiers”, “procédures d’audit” si ce n’est pas déjà présent,
- pas de faux bénéfice,
- pas d’optimisation inventée,
- pas de précision artificielle.
- INTERDIT ABSOLU : ne jamais écrire un bullet à l'infinitif (ex : "Analyser", "Rédiger", "Coordonner" seul). Tout bullet commence OBLIGATOIREMENT par un verbe conjugué au passé composé (ex : Réalisé, Coordonné, Piloté, Rédigé, Développé, Géré, Obtenu, Analysé, Structuré, Négocié).
- INTERDIT ABSOLU : ne jamais terminer un bullet par une phrase participiale inventée ("assurant", "contribuant à", "favorisant", "permettant", "garantissant", "renforçant").
- INTERDIT ABSOLU : ne jamais écrire un bullet à l'infinitif. Tout bullet commence par un verbe conjugué au passé composé (Réalisé, Coordonné, Piloté, Géré, Développé, Analysé, Structuré, Négocié...).
- INTERDIT ABSOLU : ne jamais laisser un fragment de phrase sans verbe principal.

Style :
- sobre
- rigoureux
- crédible
- professionnel
- légèrement valorisant


RÈGLE LONGUEUR BULLETS : Chaque bullet doit faire entre 20 et 40 mots (environ 1,5 à 2 lignes dans Word). Un bullet trop court (< 15 mots) est insuffisant — enrichis-le avec le contexte, le périmètre ou la méthode utilisée, sans inventer de chiffres.
RÈGLE ACTIVITÉS : Chaque activité doit faire 10-20 mots minimum. Ne jamais laisser une activité en 1-2 mots seuls.

Format ACTIVITIES : 1 activité par ligne, sans puce, forme "Activité : description courte."
Sortie : UNIQUEMENT le CV complet.

CV :
\"\"\"{cv_text}\"\"\"
"""
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
    )
    return resp.choices[0].message.content.strip()


def llm_expand_cv_management(cv_text: str) -> str:
    if not client:
        return cv_text

    prompt = f"""
Tu dois rendre ce CV MANAGEMENT légèrement plus dense pour mieux remplir une page Word, sans tomber dans un style artificiel ou exagéré.

Règles ABSOLUES :
- Tu gardes exactement les sections : EDUCATION:, EXPERIENCES:, SKILLS:, ACTIVITIES:
- Tu ne rajoutes aucun commentaire.
- Tu n’inventes aucun chiffre, aucun outil, aucune mission nouvelle, aucun pays, aucune fréquence, aucun événement.
- Tu peux rendre une mission plus claire, plus structurée et légèrement plus valorisante.
- Tu peux ajouter une qualité transférable simple et crédible.

Tu peux uniquement :
1) ajouter 1 bullet à l’expérience la plus pertinente si elle n’en a que 2,
2) reformuler légèrement les bullets pour les rendre plus professionnelles,
3) conserver davantage de détails académiques déjà présents,
4) enrichir légèrement une activité existante sans ajouter de fait précis.

Interdictions :
- pas de recommandation stratégique formelle inventée,
- pas de benchmark inventé,
- pas de pilotage inventé,
- pas de jargon type “impact”, “efficacité”, “maximiser”, “haute qualité”, “coordination efficace” si cela sonne artificiel,
- pas de précision fictive.
- INTERDIT ABSOLU : ne jamais écrire un bullet à l'infinitif (ex : "Analyser", "Rédiger", "Coordonner" seul). Tout bullet commence OBLIGATOIREMENT par un verbe conjugué au passé composé (ex : Réalisé, Coordonné, Piloté, Rédigé, Développé, Géré, Obtenu, Analysé, Structuré, Négocié).
- INTERDIT ABSOLU : ne jamais terminer un bullet par une phrase participiale inventée ("assurant", "contribuant à", "favorisant", "permettant", "renforçant", "maximisant").
- INTERDIT ABSOLU : ne jamais écrire un bullet à l'infinitif. Tout bullet commence par un verbe conjugué au passé composé (Réalisé, Coordonné, Piloté, Géré, Développé, Analysé, Structuré, Négocié...).
- INTERDIT ABSOLU : ne jamais laisser un fragment de phrase sans verbe principal.

Style :
- structuré
- crédible
- professionnel
- légèrement valorisant
- simple
- pas de bullshit consulting


RÈGLE LONGUEUR BULLETS : Chaque bullet doit faire entre 20 et 40 mots (environ 1,5 à 2 lignes dans Word). Un bullet trop court (< 15 mots) est insuffisant — enrichis-le avec le contexte, le périmètre ou la méthode utilisée, sans inventer de chiffres.
RÈGLE ACTIVITÉS : Chaque activité doit faire 10-20 mots minimum. Ne jamais laisser une activité en 1-2 mots seuls.

Format ACTIVITIES : 1 activité par ligne, sans puce, forme "Activité : description courte."
Sortie : UNIQUEMENT le CV complet.

CV :
\"\"\"{cv_text}\"\"\"
"""
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
    )
    return resp.choices[0].message.content.strip()

# --- MVP "DB" en mémoire (à remplacer par Postgres plus tard)
# quota[email] = "YYYY-MM" (mois où le gratuit a été consommé)
quota: Dict[str, str] = {}

# ── Rate limiting par IP ──────────────────────────────────────────────────────
# Structure : { ip: [(timestamp, endpoint), ...] }
_ip_hits: Dict[str, list] = {}
_ip_lock = __import__("threading").Lock()

def _check_ip_rate_limit(ip: str, endpoint: str, max_hits: int, window_seconds: int) -> bool:
    """
    Retourne True si la requête est autorisée, False si le rate limit est atteint.
    Fenêtre glissante simple en mémoire.
    """
    now = __import__("time").time()
    with _ip_lock:
        hits = _ip_hits.get(ip, [])
        # Nettoyer les hits hors fenêtre
        hits = [(t, ep) for t, ep in hits if now - t < window_seconds]
        # Compter uniquement les hits pour cet endpoint
        count = sum(1 for t, ep in hits if ep == endpoint)
        if count >= max_hits:
            _ip_hits[ip] = hits
            return False
        hits.append((now, endpoint))
        _ip_hits[ip] = hits
        return True
# jobs[job_id] = {"docx_path":..., "pdf_path":...}
jobs: Dict[str, Dict[str, str]] = {}

# Sessions Stripe en attente : stripe_session_id -> payload utilisateur
pending_stripe_sessions: Dict[str, Dict[str, Any]] = {}

# Price IDs Stripe
STRIPE_PRICE_UNITE = "price_1TMCoaRqCGwhB7YCBKew54tL"
STRIPE_PRICE_MENSUEL = "price_1TMCpGRqCGwhB7YC9NFFc8yx"


app = FastAPI()

# Limite à 3 générations simultanées pour éviter les crashes mémoire
_cv_semaphore = asyncio.Semaphore(3)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[os.getenv("ALLOWED_ORIGIN", "https://mycvcopilote.netlify.app")],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
@app.get("/")
def root():
    return {"ok": True, "service": "mycvcopilot-backend"}

if STRIPE_SECRET:
    stripe.api_key = STRIPE_SECRET

def month_key(now: Optional[dt.datetime] = None) -> str:
    now = now or dt.datetime.utcnow()
    return f"{now.year:04d}-{now.month:02d}"


def sector_to_template(sector: str) -> str:
    s = sector.lower()

    if "finance" in s:
        return "templates/finance.docx"

    if "audit" in s:
        return "templates/audit.docx"

    if "management stratégique" in s or "management strategique" in s or "stratégie" in s or "strategie" in s:
        return "templates/management_strategique.docx"

    if "droit" in s:
        return "templates/droit.docx"

    return "templates/finance.docx"

def is_legal_sector(sector: str) -> bool:
    s = (sector or "").lower()
    return (
        "droit" in s
        or "juridique" in s
        or "juriste" in s
        or "avocat" in s
        or "legal" in s
    )

def is_audit_sector(sector: str) -> bool:
    s = (sector or "").lower()
    return "audit" in s

def is_finance_sector(sector: str) -> bool:
    s = (sector or "").lower()
    return "finance" in s
    
def is_management_sector(sector: str) -> bool:
    s = (sector or "").lower()
    return (
        "management stratégique" in s
        or "management strategique" in s
        or "stratégie" in s
        or "strategie" in s
        or "conseil" in s
    )
    
def sanitize_filename(name: str) -> str:
    name = re.sub(r"[^a-zA-Z0-9_-]+", "_", name.strip())
    return name[:50] or "cv"
    
def build_cv_filename(payload: Dict[str, Any]) -> str:
    full_name = (payload.get("full_name") or "").strip()
    company = (payload.get("company") or "").strip()

    parts = full_name.split()
    if not parts:
        family_name = "CANDIDAT"
    else:
        family_name = "_".join(parts[-2:]) if len(parts) >= 2 else parts[-1]

    family_name = sanitize_filename(family_name).upper()
    company_clean = sanitize_filename(company).upper()

    if company_clean:
        return f"CV-{family_name}-{company_clean}"
    return f"CV-{family_name}"


# ════════════════════════════════════════════════════════════════════
# SYSTÈME D'ADAPTATION À L'OFFRE — EXTRACTION ET VÉRIFICATION
# ════════════════════════════════════════════════════════════════════

# Vocabulaire sectoriel : termes multi-mots prioritaires à extraire
SECTOR_VOCAB = {
    "finance": [
        "modélisation financière", "modélisation lbo", "lbo", "dcf", "leveraged buyout",
        "due diligence", "due-diligence", "data room", "pitch book", "pitchbook",
        "m&a", "fusions acquisitions", "private equity", "venture capital",
        "equity research", "analyse financière", "analyse sectorielle",
        "valorisation", "multiple d'entrée", "taux de rendement",
        "reporting financier", "budget prévisionnel", "prévisionnel financier",
        "trésorerie", "cash flow", "bilan", "compte de résultat",
        "bloomberg terminal", "bloomberg", "factset", "capital iq",
        "excel avancé", "vba", "macros", "tableaux croisés dynamiques", "tcd",
        "python", "sql", "power bi", "crm",
        "anglais courant", "anglais professionnel", "anglais bilingue",
        "suivi de portefeuille", "analyse de portefeuille",
        "sourcing", "deal flow", "term sheet", "cap table",
        "compliance", "kyc", "aml", "conformité",
        "banque d'investissement", "cib", "corporate finance",
        "contrôle de gestion", "fp&a", "consolidation", "ifrs",
        "actifs sous gestion", "aum", "fund accounting",
        "investissement responsable", "esg",
    ],
    "audit": [
        "commissariat aux comptes", "cac", "audit légal", "audit contractuel",
        "contrôle interne", "tests substantifs", "tests analytiques",
        "circularisation", "feuilles de travail", "dossier permanent", "dossier annuel",
        "ifrs", "normes ifrs", "pcg", "normes isa", "normes d'audit",
        "clôture annuelle", "clôture mensuelle", "reforecast",
        "liasse fiscale", "déclaration tva", "tva",
        "auditsoft", "caseware", "drgm", "sage", "coala",
        "big 4", "deloitte", "pwc", "kpmg", "ey", "bdo", "mazars",
        "rigueur", "précision", "fiabilité", "esprit d'analyse",
        "excel avancé", "vba", "tableaux de bord",
        "risk management", "risque opérationnel",
        "rapport d'audit", "recommandations", "points de contrôle",
        "comptabilité générale", "comptabilité analytique",
    ],
    "management": [
        "gestion de projet", "project management", "agile", "scrum",
        "management d'équipe", "leadership", "coordination",
        "stratégie", "analyse stratégique", "business plan", "business case",
        "benchmark", "analyse de marché", "étude de marché",
        "powerpoint", "présentation", "pitching",
        "crm", "salesforce", "hubspot",
        "marketing digital", "seo", "sem", "réseaux sociaux",
        "kpi", "indicateurs de performance", "tableau de bord",
        "supply chain", "logistique", "opérations",
        "relation client", "customer success", "b2b", "b2c",
        "change management", "conduite du changement",
        "excel", "power bi", "notion", "trello", "jira",
        "conseil", "consulting", "mckinsey", "bcg", "bain",
    ],
    "droit": [
        "droit des sociétés", "droit des affaires", "droit social",
        "droit des contrats", "droit fiscal", "droit commercial",
        "rédaction d'actes", "rédaction juridique",
        "recherche jurisprudentielle", "jurisprudence",
        "dalloz", "lexis360", "village justice",
        "contentieux", "arbitrage", "médiation",
        "assemblée générale", "augmentation de capital", "cession",
        "sas", "sarl", "sa", "statuts",
        "due diligence juridique", "audit juridique",
        "propriété intellectuelle", "marques", "brevets",
        "rgpd", "protection des données",
        "fusions acquisitions", "m&a",
        "droit pénal des affaires", "compliance",
        "mémoire", "master", "master 2", "m2",
    ],
}

FRENCH_STOP_WORDS = {
    "le", "la", "les", "un", "une", "des", "du", "de", "d", "l",
    "et", "ou", "mais", "donc", "or", "ni", "car",
    "que", "qui", "quoi", "dont", "où",
    "je", "tu", "il", "elle", "nous", "vous", "ils", "elles",
    "ce", "cet", "cette", "ces", "mon", "ton", "son", "nos", "vos", "leurs",
    "dans", "sur", "sous", "avec", "sans", "pour", "par", "en", "à", "au", "aux",
    "est", "sont", "sera", "seront", "être", "avoir", "faire",
    "plus", "très", "bien", "aussi", "même", "tout", "tous",
    "si", "se", "sa", "lui", "leur",
    "cas", "type", "mise", "lors", "afin",
    "notamment", "notamment", "ainsi", "donc",
    "votre", "notre", "leurs", "vos",
    "the", "and", "or", "of", "in", "to", "a", "an", "for", "on", "with",
}

def _normalize_for_matching(text: str) -> str:
    """Normalise un texte pour la comparaison (lowercase, accents, ponctuations)."""
    import unicodedata
    text = text.lower().strip()
    # Supprimer les accents
    text = "".join(
        c for c in unicodedata.normalize("NFD", text)
        if unicodedata.category(c) != "Mn"
    )
    # Normaliser la ponctuation
    text = re.sub(r"['\-–]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _term_in_text(term: str, text_normalized: str) -> bool:
    """Vérifie si un terme (ou une variante proche) est dans le texte normalisé."""
    term_n = _normalize_for_matching(term)
    if not term_n:
        return False
    # Match exact
    if term_n in text_normalized:
        return True
    # Match avec dérivés courants (pluriel, conjugaison)
    # Ex: "modélisation" → "modelis" (stem de 7 chars)
    if len(term_n) >= 6:
        stem = term_n[:max(5, len(term_n) - 3)]
        if stem in text_normalized:
            return True
    return False


def build_keyword_mapping(
    job_posting: str,
    raw_experiences: str,
    raw_education: str,
    raw_skills: str,
    sector: str,
) -> dict:
    """
    Extrait les mots-clés de l'offre et les mappe sur le profil réel de l'utilisateur.
    
    Retourne:
    {
        "applicable": [(term, reason), ...],   # termes validés dans le profil
        "absent": [term, ...],                  # termes absents du profil (à ne pas inventer)
        "sector_terms": [...],                  # termes sectoriels extraits de l'offre
    }
    """
    if not job_posting or not job_posting.strip():
        return {"applicable": [], "absent": [], "sector_terms": []}

    # Contenu réel de l'utilisateur (tout normalisé)
    user_raw = f"{raw_experiences} {raw_education} {raw_skills}"
    user_norm = _normalize_for_matching(user_raw)

    # Déterminer le vocabulaire sectoriel à utiliser
    sector_low = sector.lower()
    if any(s in sector_low for s in ["finance", "banque", "private equity", "audit"]):
        vocab = SECTOR_VOCAB.get("finance", []) + SECTOR_VOCAB.get("audit", [])
    elif "audit" in sector_low:
        vocab = SECTOR_VOCAB.get("audit", [])
    elif any(s in sector_low for s in ["management", "stratégique", "conseil", "marketing"]):
        vocab = SECTOR_VOCAB.get("management", [])
    elif "droit" in sector_low or "juridique" in sector_low:
        vocab = SECTOR_VOCAB.get("droit", [])
    else:
        vocab = SECTOR_VOCAB.get("finance", [])

    job_norm = _normalize_for_matching(job_posting)

    # Étape 1 : termes du vocabulaire sectoriel présents dans l'offre
    sector_terms_in_offer = []
    for term in vocab:
        if _term_in_text(term, job_norm):
            sector_terms_in_offer.append(term)

    # Étape 2 : extraction de termes libres depuis l'offre (mots significatifs)
    # Bigrams et trigrams de l'offre
    offer_words = [w for w in re.split(r"[\s,;.()\[\]]+", job_posting.lower()) if w]
    offer_significant = []
    for w in offer_words:
        w_clean = _normalize_for_matching(w)
        if len(w_clean) >= 5 and w_clean not in FRENCH_STOP_WORDS:
            offer_significant.append(w_clean)

    # Bigrams significatifs de l'offre
    bigrams = []
    for i in range(len(offer_words) - 1):
        bg = _normalize_for_matching(f"{offer_words[i]} {offer_words[i+1]}")
        if len(bg) >= 8 and not any(sw in bg.split() for sw in list(FRENCH_STOP_WORDS)[:20]):
            bigrams.append(bg)

    all_candidates = list(set(sector_terms_in_offer + offer_significant[:30] + bigrams[:20]))

    # Étape 3 : mapper sur le profil utilisateur
    applicable = []
    absent = []

    for term in all_candidates:
        if len(term) < 3:
            continue
        in_profile = _term_in_text(term, user_norm)

        # Déterminer la raison de l'applicabilité
        if in_profile:
            # Trouver où dans le profil
            user_norm_snippet = ""
            for chunk in user_norm.split("."):
                if _term_in_text(term, chunk):
                    user_norm_snippet = chunk.strip()[:80]
                    break
            applicable.append((term, user_norm_snippet))
        else:
            # Terme dans l'offre mais PAS dans le profil
            # Ne marquer absent que si c'est un terme technique précis (pas un mot générique)
            if len(term) >= 6 and term not in FRENCH_STOP_WORDS:
                absent.append(term)

    # Trier par longueur décroissante (les termes les plus spécifiques en premier)
    applicable.sort(key=lambda x: len(x[0]), reverse=True)
    absent = sorted(set(absent), key=len, reverse=True)

    # Limiter pour ne pas surcharger le prompt
    applicable = applicable[:20]
    absent = absent[:15]

    return {
        "applicable": applicable,
        "absent": absent,
        "sector_terms": sector_terms_in_offer[:10],
    }


def build_keyword_injection(mapping: dict) -> str:
    """
    Génère le bloc d'instruction à injecter dans le prompt de génération.
    Version renforcée : exemples concrets de transformation + instruction impérative.
    """
    if not mapping["applicable"] and not mapping["absent"]:
        return ""

    applicable = mapping["applicable"]
    absent = mapping["absent"]

    lines = []
    lines.append("")
    lines.append("╔══════════════════════════════════════════════════════════════╗")
    lines.append("║         ADAPTATION OBLIGATOIRE À L'OFFRE D'EMPLOI          ║")
    lines.append("╚══════════════════════════════════════════════════════════════╝")
    lines.append("")
    lines.append("Le système a vérifié les mots-clés de l'offre contre le profil réel.")
    lines.append("Tu DOIS appliquer les règles suivantes AVANT de rédiger les bullets.")
    lines.append("")

    if applicable:
        lines.append("✅ TERMES VÉRIFIÉS — UTILISE-LES OBLIGATOIREMENT dans les bullets :")
        lines.append("   (Ces termes sont dans l'offre ET dans le profil réel du candidat)")
        lines.append("")
        for term, ctx in applicable[:12]:
            lines.append(f"   → \"{term}\"")
        lines.append("")
        lines.append("   MÉTHODE : Pour chaque expérience, si un terme ci-dessus décrit")
        lines.append("   ce que le candidat a fait → remplace le vocabulaire générique par")
        lines.append("   ce terme exact. Exemples :")
        lines.append("   • \"modélisation financière\" → \"modélisation LBO et DCF\"")
        lines.append("   • \"note pour le comité\" → \"mémo d'investissement\"")
        lines.append("   • \"tests\" → \"tests de contrôle interne et tests substantifs\"")
        lines.append("   • \"vérification documents\" → \"circularisation créances clients\"")
        lines.append("")

    if absent:
        lines.append("❌ TERMES INTERDITS — ABSENTS du profil, ne jamais les inventer :")
        for term in absent[:8]:
            lines.append(f"   ✗ \"{term}\"")
        lines.append("")

    lines.append("LONGUEUR BULLETS : Chaque bullet = 20-35 mots minimum (1,5 lignes)")
    lines.append("VERBES : Passé composé OBLIGATOIRE (Réalisé, Développé, Coordonné...)")
    lines.append("╚═══════════════════════════════════════════════════════════════╝")
    lines.append("")

    return "\n".join(lines)


def build_prompt(payload: Dict[str, Any]) -> str:
    return f"""
Tu es un expert en recrutement.
Tu dois générer un CV FRANÇAIS d'1 page maximum, ultra sobre, ATS-friendly, clair et crédible.

Le CV doit être adapté :
- au secteur : {payload["sector"]}
- au poste : {payload["role"]}
- à l’entreprise : {payload["company"]}
- à l’offre d’emploi ci-dessous

OFFRE D'EMPLOI :
\"\"\"{payload["job_posting"]}\"\"\"

PROFIL UTILISATEUR :
Nom : {payload["full_name"]}
Ville : {payload["city"]}
Email : {payload["email"]}
Téléphone : {payload["phone"]}
LinkedIn : {payload.get("linkedin","")}

FORMATION :
{payload["education"]}

EXPÉRIENCES :
{payload["experiences"]}

COMPÉTENCES :
{payload["skills"]}

LANGUES :
{payload["languages"]}

CENTRES D’INTÉRÊT :
{payload.get("interests","")}

RÈGLES ABSOLUES :
- Tu n’inventes rien.
- Tu n’ajoutes ni chiffres, ni missions, ni outils, ni distinctions non fournis.
- Tu restes crédible, professionnel et sobre.
- Tu reformules intelligemment pour valoriser le profil sans mentir.
- Chaque expérience contient 3 bullet points si l'utilisateur en a fourni 3, et 2 uniquement si le profil est très chargé. Ne supprime jamais un bullet fourni par l'utilisateur.
- Chaque bullet doit être concret, court et orienté action.
- Si le CV semble trop vide, tu densifies d’abord les expériences, puis les activités, sans inventer.
- Si une expérience est peu détaillée, tu la rends professionnelle sans extrapoler.
- Tu n’ajoutes jamais de finalité business, de bénéfice, de recommandation ou d’amélioration non explicitement fournis.
- Les langues ne doivent JAMAIS être une section séparée.
- Les langues doivent être intégrées dans SKILLS, sur une ligne commençant par "Langues :".
- La section SKILLS doit contenir 2 à 4 lignes maximum parmi :
  "Certifications : ..."
  "Maîtrise des logiciels : ..."
  "Capacités professionnelles : ..."
  "Langues : ..."
- La section ACTIVITIES doit contenir uniquement des centres d’intérêt personnels.
- Chaque activité doit tenir sur une ligne sous la forme :
  "Activité : pratique factuelle ; qualité développée"
- Tu n’écris aucun commentaire, aucune introduction, aucune phrase méta.

FORMAT DE SORTIE OBLIGATOIRE :
EDUCATION:
<contenu>

EXPERIENCES:
<contenu>

SKILLS:
<contenu incluant les langues>

ACTIVITIES:
<contenu>

IMPORTANT :
- Tu ne dois rien écrire avant EDUCATION:
- Tu ne dois rien écrire après ACTIVITIES:
- Tu ne génères surtout PAS de section LANGUAGES:
"""
    
def build_prompt_finance(payload: Dict[str, Any]) -> str:
    # Construire le mapping mots-clés offre ↔ profil réel
    _exp_anchor = build_mandatory_experience_anchor(payload)
    _exp_anchor = build_mandatory_experience_anchor(payload)
    _exp_anchor = build_mandatory_experience_anchor(payload)
    _exp_anchor = build_mandatory_experience_anchor(payload)
    _kw_map = build_keyword_mapping(
        job_posting=payload.get("job_posting", ""),
        raw_experiences=payload.get("experiences", ""),
        raw_education=payload.get("education", ""),
        raw_skills=payload.get("skills", ""),
        sector=payload.get("sector", "finance"),
    )
    _kw_injection = build_keyword_injection(_kw_map)

    return f"""
Tu es un ancien recruteur en banque d’investissement et en Big 4.
Tu sélectionnes uniquement les 10% meilleurs profils étudiants.
Tu élimines immédiatement les CV vagues, imprécis ou sans résultats chiffrés.

OBJECTIF :
Générer un CV FINANCE français d’1 page maximum, ultra structuré, minimal et stratégique.

Le CV doit être adapté :
- au type de finance visé : {payload.get("finance_type", "Non précisé")}
- au poste : {payload["role"]}
- à l’entreprise : {payload["company"]}
- à l’offre d’emploi

OFFRE D’EMPLOI :
\"\"\"{payload["job_posting"]}\"\"\"
{_exp_anchor}
{_kw_injection}

RÈGLES :
- 1 page maximum (ABSOLUMENT aucune 2e page).
- Format de dates homogène, toujours sous la forme "MMM YYYY – MMM YYYY"
  (exemple : "Sept 2023 – Juin 2025") et jamais "09/2023", "2023-2025" ou "au".
- Chaque bullet = Verbe fort + action concrète issue de l'expérience source.
- Tu peux professionnaliser la formulation, mais tu n'ajoutes jamais d'impact business ou de bénéfice implicite non fourni.
- 2 à 3 bullets maximum par expérience (3 par défaut, 2 uniquement pour les expériences les moins pertinentes).
- Interdiction des mots : assisted, helped, worked on.
- Ton professionnel, précis, sobre.
- Classe les expériences de la plus pertinente à la moins pertinente par rapport au poste visé.
- Les expériences de tutorat / soutien scolaire sont plus pertinentes qu’un job de caisse générique et doivent être placées AU-DESSUS des jobs étudiants alimentaires.
- Les expériences en finance / audit / assurance / banque / analyse financière doivent être tout en haut, même si elles sont plus anciennes.
- Les jobs étudiants génériques (supermarché, baby-sitting, barista, etc.) doivent toujours être en bas de la section EXPÉRIENCES, même s’ils sont plus récents.
- Si le contenu commence à être trop long pour tenir sur une page, tu SUPPRIMES d’abord les expériences les moins pertinentes (jobs étudiants génériques) et tu raccourcis les bullets les moins importantes.
- Le CV doit être rédigé intégralement en français (même si l’offre ou les intitulés sont en anglais).
- Tous les bullet points doivent être écrits en français.
- LONGUEUR BULLETS : chaque bullet doit faire 20 à 35 mots (environ 1,5 ligne). Un bullet court (< 15 mots) = insuffisant. Enrichis avec le contexte, le périmètre, la méthode.
- prioriser ces verbes : analyser, évaluer, structurer, modéliser, préparer, synthétiser, présenter, suivre
- éviter ces verbes: aider, assister, participer, contribuer

RÈGLES STRICTES :
Ces règles priment sur toutes les autres instructions.
- Tu n’inventes AUCUN chiffre.
- Tu n’inventes AUCUNE mission.
- Tu n’inventes AUCUN outil.
- Si une information est absente, tu restes général sans ajouter de détails fictifs.
- Si aucun résultat chiffré n’est fourni, tu reformules sans métriques.
- Tu utilises uniquement les informations présentes dans le profil utilisateur.
- Interdiction totale d’inventer pour “améliorer” le CV.
- Si une expérience contient trop peu d'informations, tu la rends professionnelle mais concise, sans extrapolation.
- Évite les verbes faibles (participé, aidé, effectué, travaillé sur).
- Privilégie des verbes orientés impact et responsabilité.
- Chaque bullet doit refléter une contribution concrète.
- Tu peux reformuler une expérience existante pour la rendre plus claire et plus professionnelle.
- Tu ne dois jamais inventer un impact chiffré, un résultat business précis, une recommandation formelle ou une finalité stratégique lourde si ce n’est pas explicitement fourni.
- En revanche, tu peux reformuler une mission existante de manière légèrement plus valorisante et plus professionnelle si cela reste directement crédible au regard du texte source.
- Tu ne dois jamais inventer une activité, un projet, un événement, un impact, une recommandation ou un bénéfice business.
- Tu peux améliorer la formulation pour la rendre plus professionnelle, plus concise et plus crédible.
- Tu peux faire ressortir une qualité transférable ou une compétence utile au poste uniquement si elle découle directement d’un fait fourni.
- Exemple autorisé :
  "Équitation pratiquée à niveau national pendant 15 ans" peut devenir
  "Équitation : pratique de haut niveau développant discipline, patience et résilience."
- Exemple interdit :
  ajouter une fréquence, un club, un palmarès, un encadrement, un rôle ou une performance non fournis.
- Tu enrichis le style, jamais les faits.

HALLUCINATIONS (INTERDICTION ABSOLUE) :
- Dans EDUCATION : interdiction d’ajouter des séminaires, conférences, ateliers, études de cas, projets, classements, GPA/moyenne, prix, bourses, matières, cours, spécialisations, options ou modules
  SI ce n’est pas explicitement écrit dans le champ FORMATION utilisateur.
- Interdiction absolue d’ajouter une matière "logique" ou "proche du secteur" si elle n’est pas fournie mot pour mot ou clairement présente dans le champ FORMATION.
- Dans EXPERIENCES : interdiction d’ajouter des impacts inventés ("augmentant", "optimisant", "améliorant", "permettant", "renforçant", "contribuant à", "garantissant", "assurant", "identifiant", "mettant en évidence", "présentant des recommandations", "proposant des recommandations")
  SI l’impact, la finalité ou la recommandation n’est pas explicitement présente dans l’expérience brute.
- Dans ACTIVITIES : interdiction d’ajouter un niveau ("compétition", "national", "régional", "club", "championnat", "hebdomadaire", "quotidien")
  SI ce n’est pas explicitement écrit dans CENTRES D’INTÉRÊT utilisateur.
  
INTERDICTION ABSOLUE d’ajouter :
- Classement
- GPA
- Moyenne
- Distinction académique
- Prix
- Bourse
SI ces informations ne sont pas explicitement présentes dans le profil utilisateur.

BDE / ASSOCIATIONS / PROJETS ÉTUDIANTS :
- Tu DOIS les mettre dans "EXPERIENCES" (même si ce n’est pas une entreprise).
- Tu les écris comme une expérience (titre + dates si disponibles + 2-3 bullets).
- INTERDICTION ABSOLUE d’inventer des chiffres : aucun %, aucun volume, aucun "5 sponsors", aucun "100 participants" si ce n’est pas fourni.

SECTION SKILLS (COMPÉTENCES & OUTILS) :
- Tu produis EXACTEMENT 2 à 4 lignes sous "SKILLS:" :
  1) "Certifications : ..."
  2) "Maîtrise des logiciels : ..."
  3) "Capacités professionnelles : ..." (facultatif si peu d'infos)
  4) "Langues : ..."
- Si aucune certification n’est fournie, tu n’écris JAMAIS "Certifications : ...".
- Dans chaque ligne, les éléments sont séparés par des virgules (PAS de "|").
- "Certifications" : tests ou validations concrètes (Excel, PIX, etc.).
- "Maîtrise des logiciels" : Excel, PowerPoint, VBA, outils spécifiques.
- "Capacités professionnelles" : 3–4 compétences maximum, simples, sobres et directement liées à l’offre (ex : analyse financière, reporting, gestion des priorités, communication professionnelle).
- Interdiction d’utiliser des formulations trop valorisantes comme "avancé", "approfondi", "complexe", "percutant", "stratégique", "excellente maîtrise", sauf si explicitement fourni.
- Les langues doivent être intégrées ici sur une ligne "Langues : ...".
- Les tests de langues officiels peuvent apparaître dans cette ligne s’ils sont explicitement fournis.


SECTION ACTIVITIES (CENTRES D’INTÉRÊT) :
- Tu n’y mets QUE des centres d’intérêt / activités personnelles (sport, voyages, engagements associatifs non listés en expérience, hobbies).
- INTERDICTION d’y mettre BDE / associations / projets déjà listés dans EXPÉRIENCES.
- Pas de doublons : si c’est dans EXPÉRIENCES, tu ne le répètes pas ailleurs.
- Tu n’utilises JAMAIS de Markdown (**texte**, *texte*). Tu écris simplement le texte brut.
- Format de chaque activité sur UNE LIGNE :
  Nom de l’activité en gras (nous ferons le gras côté Word), suivi de ":" puis une phrase courte et factuelle.

- La phrase doit décrire concrètement la pratique :
  - niveau (loisir, régulier, intensif, compétition, etc.) si disponible,
  - fréquence ou cadence si disponible (ex : 2 à 3 fois par semaine),
  - contexte si pertinent (club, voyages, événements, etc.).

- Si ces informations ne sont pas fournies, tu restes factuel sans inventer.

- Tu peux mentionner au maximum UNE qualité simple et crédible (ex : rigueur, discipline, persévérance), mais uniquement si elle est directement cohérente avec l’activité.

- Interdiction d’utiliser un ton RH générique ou trop valorisant.

IMPORTANT :
- Toute la sortie (EDUCATION, EXPERIENCES, SKILLS, ACTIVITIES)
  doit être rédigée EN FRANÇAIS.
- Si tu écris une phrase en anglais, tu la traduis immédiatement en français.
- Seuls les noms propres (noms d’écoles, diplômes officiels, logiciels, intitulés exacts de postes)
  peuvent rester en anglais.

RÈGLES DE SORTIE (TRÈS IMPORTANT) :
- Tu génères UNIQUEMENT les sections suivantes, dans cet ordre :
  EDUCATION:
  EXPERIENCES:
  SKILLS:
  ACTIVITIES:
- Tu ne génères PAS de titre de section supplémentaire.
- Tu ne génères PAS le nom.
- Tu ne génères PAS les coordonnées.
- Tu ne génères PAS d'accroche.
- Tu ne génères JAMAIS de section "LANGUAGES:" ou "LANGUES:" séparée.
- Les langues doivent toujours être incluses dans SKILLS sur une ligne "Langues : ...".

FORMAT EXACT À RESPECTER :

1️⃣ TU DOIS ABSOLUMENT PRODUIRE CES 4 BLOCS DANS CET ORDRE EXACT,
   CHAQUE EN-TÊTE SUR SA PROPRE LIGNE :

EDUCATION:
<contenu éducation>

EXPERIENCES:
<contenu expériences>

SKILLS:
<contenu compétences incluant les langues>

ACTIVITIES:
<contenu activités>

2️⃣ TU NE DOIS RIEN ÉCRIRE AVANT "EDUCATION:" NI APRÈS LA DERNIÈRE LIGNE D’ACTIVITIES.
   PAS DE COMMENTAIRES, PAS DE TEXTE EXPLICATIF, PAS D’INTRODUCTION, RIEN.

3️⃣ FORMAT PRÉCIS DE CHAQUE BLOC :

EDUCATION:
DEGREE: <intitulé du diplôme ou programme>
SCHOOL: <école ou université>
LOCATION: <Ville, Pays>
DATES: <MMM YYYY – MMM YYYY ou MMM YYYY – Present>
DETAILS:
- <ligne de détail 1 (ex : Matières fondamentales : ... )>
- <ligne de détail 2>
- <ligne de détail 3>

DEGREE: <autre diplôme ou échange académique>
SCHOOL: <école ou université>
LOCATION: <Ville, Pays>
DATES: <MMM YYYY – MMM YYYY ou MMM YYYY – Present>
DETAILS:
- <détail 1>
- <détail 2>

EXPERIENCES:
ROLE: <intitulé exact du poste>
COMPANY: <nom exact de l’entreprise ou de l’association>
DATES: <MMM YYYY – MMM YYYY ou MMM YYYY – Present>
LOCATION: <Ville, Pays>
TYPE: <Internship / Alternance / CDI / etc. si fourni sinon vide>
BULLETS:
- <bullet 1>
- <bullet 2>
- <bullet 3>

ROLE: <autre poste>
COMPANY: ...
DATES: ...
LOCATION: ...
TYPE: ...
BULLETS:
- ...
- ...

SKILLS:
<2 à 4 lignes, chacune commençant par "Certifications :", "Maîtrise des logiciels :", "Capacités professionnelles :" ou "Langues :">

ACTIVITIES:
<une activité par ligne, sans puces, sous la forme "Nom de l’activité : description">

CONTRAINTE LONGUEUR (INTELLIGENTE) :

Le CV doit tenir STRICTEMENT sur UNE SEULE page Word.
Tu dois viser une densité pro optimale :
- ni trop vide
- ni surchargé
- une page pleine mais aérée.

RÈGLE STRUCTURELLE DE BASE :
- Maximum 9 bullet points au total (jamais plus de 9).
- 2 bullet points par défaut par expérience.
- 3 bullet points uniquement pour les 1 ou 2 expériences **les plus pertinentes** pour l’offre.
- Tu ne crées pas plus de 4 expériences au total (hors éventuellement 1 job étudiant très court).
- Tu ne supprimes **jamais** une expérience en finance / audit / banque / BDE / projet important, sauf si le profil en contient vraiment trop.

RÈGLE D’AJUSTEMENT AUTOMATIQUE :

1️⃣ Si le contenu devient trop long :
- Tu réduis d’abord les expériences les **moins pertinentes** (jobs étudiants génériques, etc.).
- Tu limites à 2 bullet points maximum pour les expériences secondaires.
- Tu raccourcis les formulations (phrases plus directes, une seule idée par bullet).
- Tu supprimes **uniquement en dernier recours** un job étudiant générique (caisse, vente, etc.), jamais une expérience en finance / audit / BDE / projet sérieux.
- Tu gardes toujours au moins 3 expériences au total si possible.

2️⃣ Si le CV semble trop court (moins d’une page) :
- Tu passes à 3 bullet points pour les expériences les plus pertinentes.
- Tu reformules les éléments existants de manière plus précise et plus professionnelle.
- Tu peux expliciter une compétence déjà implicite dans une expérience ou une activité.
- Tu ne dois JAMAIS ajouter de nouvelle matière, de nouveau logiciel, de nouvelle langue, de nouvelle activité, de nouveau projet ou de nouvel événement.
- Si une section manque d’informations, tu la laisses sobre au lieu d’inventer.

RÈGLES D’ÉCRITURE :
- Phrases courtes, une seule idée par bullet.
- Tu évites les répétitions entre les bullets et entre les expériences.
- Dans EDUCATION, chaque bloc DOIT contenir DETAILS: avec au moins 1 ligne "- ...".
- Tu dois reprendre les lignes "Cours : ..." fournies dans le profil et les convertir en DETAILS.

PROFIL :
Nom : {payload["full_name"]}
Ville : {payload["city"]}

FORMATION :
{payload["education"]}

EXPÉRIENCES :
{payload["experiences"]}

COMPÉTENCES :
{payload["skills"]}

LANGUES :
{payload["languages"]}

CENTRES D’INTÉRÊT :
{payload.get("interests","")}

Génère uniquement le CV structuré.
"""

def build_prompt_audit(payload: Dict[str, Any]) -> str:
    _kw_map = build_keyword_mapping(
        job_posting=payload.get("job_posting", ""),
        raw_experiences=payload.get("experiences", ""),
        raw_education=payload.get("education", ""),
        raw_skills=payload.get("skills", ""),
        sector=payload.get("sector", "audit"),
    )
    _kw_injection = build_keyword_injection(_kw_map)

    return f"""
Tu es un ancien recruteur en audit financier et en Big 4.
Tu sélectionnes uniquement les profils étudiants crédibles, rigoureux et structurés.

OBJECTIF :
Générer un CV AUDIT français d’1 page maximum, ultra structuré, sobre et professionnel.

Le CV doit être adapté :
- au poste : {payload["role"]}
- à l’entreprise : {payload["company"]}
- à l’offre d’emploi

OFFRE D’EMPLOI :
\"\"\"{payload["job_posting"]}\"\"\"
{_exp_anchor}
{_kw_injection}

RÈGLES :
- 1 page maximum.
- Format de dates homogène, toujours sous la forme "MMM YYYY – MMM YYYY".
- Chaque bullet = Verbe fort + action concrète issue de l'expérience source.
- Tu peux professionnaliser la formulation, mais tu n'ajoutes jamais de finalité professionnelle ou de bénéfice implicite non fourni.
- 2 à 3 bullets maximum par expérience.
- Ton professionnel, précis, rigoureux, sobre.
- Classe les expériences de la plus pertinente à la moins pertinente par rapport au poste visé.
- Les expériences en audit, comptabilité, contrôle de gestion, finance, conformité ou trésorerie doivent être tout en haut.
- Les expériences associatives avec gestion de budget ou organisation peuvent être valorisées.
- Les jobs étudiants génériques restent en bas.

PRIORITÉS MÉTIER AUDIT :
- prioriser les verbes : analyser, contrôler, réviser, vérifier, préparer, documenter, suivre, fiabiliser
- éviter les verbes : aider, assister, participer, contribuer
- valoriser :
  - revue de cycles
  - contrôle interne
  - vérification documentaire
  - analyse comptable et financière
  - préparation de feuilles de travail
  - suivi de procédures
  - rigueur, fiabilité, précision

RÈGLES STRICTES :
- Tu n’inventes AUCUN chiffre.
- Tu n’inventes AUCUNE mission.
- Tu n’inventes AUCUN outil.
- Tu n’utilises que les informations fournies.
- Si une expérience contient peu d’informations, tu la reformules proprement sans extrapoler.
- Tu peux professionnaliser une expérience existante et légèrement enrichir sa formulation si cela reste directement cohérent avec le texte source.
- Tu n’inventes jamais de nouveau projet, de nouvel événement, de mission entièrement nouvelle ni de résultat chiffré.

HALLUCINATIONS (INTERDICTION ABSOLUE) :
- Dans EDUCATION : interdiction d’ajouter séminaires, classements, GPA, prix, bourses, projets, matières, cours, spécialisations, options ou modules non fournis.
- Interdiction absolue d’ajouter une matière ou un cours simplement parce qu’il paraît cohérent avec l’audit.
- Dans EXPERIENCES : interdiction d’ajouter des impacts, finalités ou bénéfices inventés ("améliorant", "optimisant", "renforçant", "garantissant", "assurant", "fiabilisant", "facilitant", "permettant", "sécurisant", "mettant en évidence", etc.) si ce n’est pas explicitement fourni.
- Tu n’ajoutes jamais "conformité", "normes", "contrôle interne", "procédures d'audit", "travaux d'audit" ou "états financiers" si ces notions ne figurent pas déjà dans l’expérience source.
- Dans ACTIVITIES : interdiction d’ajouter compétition, club, fréquence ou niveau non fourni.

SECTION SKILLS (COMPÉTENCES & OUTILS) :
- Tu produis EXACTEMENT 2 à 4 lignes sous "SKILLS:" :
  1) "Certifications : ..."
  2) "Maîtrise des logiciels : ..."
  3) "Capacités professionnelles : ..."
  4) "Langues : ..."
- Si aucune certification n’est fournie, tu n’écris jamais "Certifications : ...".
- Les éléments sont séparés par des virgules.
- Les langues sont toujours intégrées dans "Langues : ...".

SECTION ACTIVITIES :
- Tu n’y mets que des centres d’intérêt personnels.
- Format : "Activité : description courte et factuelle".
- Tu peux mentionner une seule qualité simple et crédible, jamais plusieurs.
- Interdiction d’utiliser un ton RH générique ou trop valorisant.

RÈGLES DE SORTIE :
- Tu génères UNIQUEMENT :
  EDUCATION:
  EXPERIENCES:
  SKILLS:
  ACTIVITIES:
- Pas de nom, pas de coordonnées, pas d’accroche.
- Pas de section LANGUAGES séparée.

FORMAT EXACT :

EDUCATION:
DEGREE: <intitulé>
SCHOOL: <école>
LOCATION: <Ville, Pays>
DATES: <MMM YYYY – MMM YYYY>
DETAILS:
- <ligne 1>
- <ligne 2>

EXPERIENCES:
ROLE: <poste>
COMPANY: <entreprise>
DATES: <MMM YYYY – MMM YYYY>
LOCATION: <Ville, Pays>
TYPE: <type>
BULLETS:
- <bullet 1>
- <bullet 2>
- <bullet 3>

SKILLS:
<2 à 4 lignes>

ACTIVITIES:
<une activité par ligne>

PROFIL :
Nom : {payload["full_name"]}
Ville : {payload["city"]}

FORMATION :
{payload["education"]}

EXPÉRIENCES :
{payload["experiences"]}

COMPÉTENCES :
{payload["skills"]}

LANGUES :
{payload["languages"]}

CENTRES D’INTÉRÊT :
{payload.get("interests","")}

Génère uniquement le CV structuré.
"""

def build_prompt_management(payload: Dict[str, Any]) -> str:
    _kw_map = build_keyword_mapping(
        job_posting=payload.get("job_posting", ""),
        raw_experiences=payload.get("experiences", ""),
        raw_education=payload.get("education", ""),
        raw_skills=payload.get("skills", ""),
        sector=payload.get("sector", "management"),
    )
    _kw_injection = build_keyword_injection(_kw_map)

    return f"""
Tu es un recruteur en conseil, stratégie et management.
Tu sélectionnes les profils étudiants les plus structurés, analytiques et crédibles.

OBJECTIF :
Générer un CV MANAGEMENT STRATÉGIQUE français d’1 page maximum, clair, structuré et professionnel.

Le CV doit être adapté :
- au poste : {payload["role"]}
- à l’entreprise : {payload["company"]}
- à l’offre d’emploi

OFFRE D’EMPLOI :
\"\"\"{payload["job_posting"]}\"\"\"
{_exp_anchor}
{_kw_injection}

RÈGLES :
- 1 page maximum.
- Format de dates homogène, toujours sous la forme "MMM YYYY – MMM YYYY".
- Chaque bullet = Verbe fort + action concrète issue de l'expérience source.
- Tu peux professionnaliser la formulation, mais tu n'ajoutes jamais de recommandation, de finalité ou de bénéfice implicite non fourni.
- 2 à 3 bullets maximum par expérience.
- Ton professionnel, structuré, analytique, orienté résolution de problèmes.
- Classe les expériences de la plus pertinente à la moins pertinente par rapport au poste visé.
- Valorise particulièrement :
  - analyse
  - benchmark
  - diagnostic
  - recommandations
  - coordination
  - gestion de projet
  - communication
  - vision d’ensemble

PRIORITÉS MÉTIER MANAGEMENT STRATÉGIQUE :
- prioriser les verbes : analyser, structurer, coordonner, préparer, recommander, piloter, présenter, suivre
- éviter les verbes : aider, assister, participer, contribuer
- valoriser :
  - analyse de marché
  - synthèse d’informations
  - coordination de projet
  - élaboration de recommandations
  - organisation
  - résolution de problèmes
  - communication professionnelle

RÈGLES STRICTES :
- Tu n’inventes AUCUN chiffre.
- Tu n’inventes AUCUNE mission.
- Tu n’inventes AUCUN outil.
- Tu n’utilises que les informations fournies.
- Si une expérience est peu détaillée, tu la professionnalises sans extrapoler.
- Tu peux reformuler une expérience existante de manière plus structurée, plus professionnelle et légèrement plus valorisante si cela reste crédible.
- Tu dois privilégier des formulations simples, directes et naturelles.
- Tu n’inventes jamais de projet, d’événement, de recommandation stratégique formelle ni d’impact chiffré.
- Tu évites les expressions artificielles comme : besoins d'un client spécifique, décisions stratégiques, planification efficace, environnement collaboratif, portefeuille clients, processus de vente, service orienté satisfaction, fidélisation.

HALLUCINATIONS (INTERDICTION ABSOLUE) :
- Dans EDUCATION : interdiction d’ajouter classements, GPA, distinctions, projets, matières, cours, spécialisations, options ou modules non fournis.
- Interdiction absolue d’ajouter une matière ou un cours simplement parce qu’il paraît cohérent avec la stratégie ou le management.
- Dans EXPERIENCES : interdiction d’ajouter des impacts, recommandations, diagnostics, optimisations, opportunités identifiées ou bénéfices inventés.
- Tu n’ajoutes jamais "recommandations stratégiques", "diagnostic", "benchmark", "pilotage", "coordination de projet", "parties prenantes", "roadmap", "CRM", "visibilité", "efficacité", "traçabilité" ou "performance" si ces notions ne figurent pas déjà dans le texte source.
- Dans ACTIVITIES : interdiction d’ajouter un niveau, une fréquence ou un engagement non fourni.

SECTION SKILLS (COMPÉTENCES & OUTILS) :
- Tu produis EXACTEMENT 2 à 4 lignes sous "SKILLS:" :
  1) "Certifications : ..."
  2) "Maîtrise des logiciels : ..."
  3) "Capacités professionnelles : ..."
  4) "Langues : ..."
- Si aucune certification n’est fournie, tu n’écris jamais "Certifications : ...".
- Les langues sont intégrées dans "Langues : ...".

SECTION ACTIVITIES :
- Tu n’y mets que des centres d’intérêt personnels.
- Format : "Activité : description courte et factuelle".
- Tu peux mentionner une seule qualité simple et crédible, jamais plusieurs.
- Interdiction d’utiliser un ton RH générique ou trop valorisant.

RÈGLES DE SORTIE :
- Tu génères UNIQUEMENT :
  EDUCATION:
  EXPERIENCES:
  SKILLS:
  ACTIVITIES:
- Pas de nom, pas de coordonnées, pas d’accroche.
- Pas de section LANGUAGES séparée.

FORMAT EXACT :

EDUCATION:
DEGREE: <intitulé>
SCHOOL: <école>
LOCATION: <Ville, Pays>
DATES: <MMM YYYY – MMM YYYY>
DETAILS:
- <ligne 1>
- <ligne 2>

EXPERIENCES:
ROLE: <poste>
COMPANY: <entreprise>
DATES: <MMM YYYY – MMM YYYY>
LOCATION: <Ville, Pays>
TYPE: <type>
BULLETS:
- <bullet 1>
- <bullet 2>
- <bullet 3>

SKILLS:
<2 à 4 lignes>

ACTIVITIES:
<une activité par ligne>

PROFIL :
Nom : {payload["full_name"]}
Ville : {payload["city"]}

FORMATION :
{payload["education"]}

EXPÉRIENCES :
{payload["experiences"]}

COMPÉTENCES :
{payload["skills"]}

LANGUES :
{payload["languages"]}

CENTRES D’INTÉRÊT :
{payload.get("interests","")}

Génère uniquement le CV structuré.
"""

def build_prompt_droit(payload: Dict[str, Any]) -> str:
    _kw_map = build_keyword_mapping(
        job_posting=payload.get("job_posting", ""),
        raw_experiences=payload.get("experiences", ""),
        raw_education=payload.get("education", ""),
        raw_skills=payload.get("skills", ""),
        sector=payload.get("sector", "droit"),
    )
    _kw_injection = build_keyword_injection(_kw_map)

    return f"""
Tu es un recruteur juridique exigeant en cabinet d’avocats, direction juridique et stages juridiques.
Tu sélectionnes des profils étudiants sobres, rigoureux, crédibles et précis.

OBJECTIF :
Générer un CV DROIT français d’1 page maximum, structuré, lisible, académique et crédible.

Le CV doit être adapté :
- au poste : {payload["role"]}
- à l’entreprise : {payload["company"]}
- à l’offre d’emploi

OFFRE D’EMPLOI :
\"\"\"{payload["job_posting"]}\"\"\"
{_exp_anchor}
{_kw_injection}

RÈGLES GÉNÉRALES :
- 1 page maximum.
- Le CV doit être rédigé intégralement en français.
- Format de dates homogène : "MMM YYYY – MMM YYYY" ou "MMM YYYY – Present".
- Ton sobre, académique, précis, sans marketing personnel.
- Aucune phrase méta, aucune introduction, aucun commentaire.
- Tu n’écris rien avant EDUCATION: et rien après ACTIVITIES:.

RÈGLES ABSOLUES :
- Tu n’inventes rien.
- Tu n’ajoutes aucune mission, aucun chiffre, aucune matière, aucun outil, aucune distinction, aucun mémoire, aucune audience, aucun contrat, aucune veille, aucun acte, aucune note si ce n’est pas explicitement fourni.
- Tu n’ajoutes jamais de bénéfice, d’impact, de recommandation, d’optimisation ou d’amélioration si cela n’est pas explicitement indiqué.
- Tu utilises uniquement les informations présentes dans le profil utilisateur.

SECTION EDUCATION :
- En droit, la formation est centrale.
- Tu retranscris UNIQUEMENT ce qui est explicitement écrit dans le champ FORMATION :
  - l'intitulé exact du diplôme
  - l'université / école
  - le lieu et les dates
  - les matières SI et SEULEMENT SI elles sont écrites mot pour mot dans l'input
  - la mention SI et SEULEMENT SI elle est écrite mot pour mot dans l'input
- INTERDIT ABSOLU — même si cela paraît logique ou probable :
  - inventer un mémoire, une thèse, un sujet de recherche
  - inventer un concours, une compétition, une participation
  - inventer un classement, un prix, une distinction
  - inventer des matières non fournies
  - ajouter "avec la rédaction d'un mémoire", "centré sur un sujet juridique", "encadré par un professeur"
- Si le bloc de formation ne contient QUE diplôme + université (sans matières ni mention) : DETAILS: avec "- " (vide). NE PAS inventer de contenu, NE PAS écrire "Formation juridique." ni aucune phrase.
- INTERDIT ABSOLU dans DETAILS : "avec un mémoire", "centré sur", "portant sur des thèmes", "axée sur", "orientée vers" ou toute phrase inventée.
- Chaque bloc EDUCATION doit contenir DETAILS:.

SECTION EXPERIENCES :
- 2 bullet points par défaut par expérience.
- Classe les expériences de la plus pertinente à la moins pertinente pour le poste visé.
- Pour un poste en droit social, les expériences liées aux RH, au droit du travail, à la gestion de dossiers, à la rédaction formelle, à l’administratif structuré ou aux responsabilités associatives passent avant les jobs de vente ou d’accueil.
- 3 bullet points maximum uniquement pour les expériences les plus pertinentes.
- Chaque bullet doit être court, factuel, professionnel.
- Chaque bullet doit reprendre STRICTEMENT l’idée présente dans l’expérience brute, sans ajouter de finalité, de bénéfice, de conformité, d’efficacité, d’optimisation ou d’impact implicite.
- Verbes à privilégier seulement s’ils correspondent réellement au contenu :
  rédiger, analyser, rechercher, synthétiser, préparer, constituer, qualifier, assister, interpréter, mettre en conformité, assurer la veille, préparer des dossiers, participer à la rédaction
- Si le texte source contient un volume, une fréquence, un nombre de dossiers, de notes, de contrats, d’audiences, de pièces ou un délai, tu le conserves car ce sont de très bons signaux en droit.
- Si aucun chiffre n’est fourni, tu n’en inventes pas.
- Interdiction d’inventer :
  - audiences
  - contrats
  - actes
  - consultations
  - notes de synthèse
  - recherches jurisprudentielles
  - clients
  - délais
  - nombre de dossiers
  - domaines juridiques non fournis
- Si l’expérience est non juridique, tu la reformules de manière sobre et transférable, sans la transformer artificiellement en expérience juridique.
- Pour une expérience non juridique, tu n’ajoutes jamais de vocabulaire pseudo-juridique comme conformité, réglementation, sécurité juridique, analyse contractuelle ou contentieux sauf si ces mots figurent explicitement dans le texte source.
- Tu ne transformes jamais un job étudiant en faux stage juridique.
- Si peu d’informations sont fournies, tu restes simple et crédible.

SECTION SKILLS :
- Tu produis entre 2 et 4 lignes maximum sous "SKILLS:" parmi :
  1) "Certifications : ..."
  2) "Maîtrise des logiciels : ..."
  3) "Capacités professionnelles : ..."
  4) "Langues : ..."
- La ligne "Certifications :" peut inclure, si explicitement fournis :
  PIX, certifications numériques, concours de plaidoirie, moot courts, mock trials, certifications ou examens utiles au poste.
- Tu n’ajoutes jamais Dalloz, LexisNexis, Doctrine, Légifrance, Word avancé, Excel basique, ni aucun autre outil juridique ou bureautique s’ils ne sont pas explicitement fournis.
- Tu n’ajoutes jamais de domaine du droit maîtrisé s’il n’est pas explicitement présent dans la formation, les expériences ou les compétences fournies.
- Tu peux reprendre un moot court, mock trial ou concours de plaidoirie dans "Certifications :" seulement s’il est explicitement fourni comme élément distinct de la formation.
- Les tests et scores de langue (TOEIC, TOEFL, IELTS, Cambridge, etc.) ne doivent JAMAIS apparaître dans "Certifications :".
- Les tests et scores de langue doivent toujours être intégrés dans la ligne "Langues :".
- Tu n’inventes jamais une certification, un concours ou un examen.
- Si rien n’est fourni, tu n’écris pas la ligne "Certifications :".
- Les éléments sont séparés par des virgules.
- Tu n’ajoutes aucun outil juridique non fourni.
- Les langues doivent être intégrées dans "Langues : ...".
- Si aucune certification n’est fournie, tu n’écris pas "Certifications : ...".
- Si aucune capacité professionnelle claire n’est fournie, tu n’écris pas "Capacités professionnelles : ...".
- Tu dois toujours écrire au minimum :
  "Maîtrise des logiciels : ..."
  "Langues : ..."

SECTION ACTIVITIES :
- Tu n’y mets QUE des centres d’intérêt personnels réels.
- Chaque activité doit obligatoirement contenir :
    1. une pratique concrète (ex : compétition, engagement, fréquence, projet, expérience)
    2. un contexte ou niveau (ex : club, association, voyage, événement, durée)
    3. un lien implicite avec des qualités utiles (sans exagération)

- Format obligatoire :
  "Activité : description concrète + impact ou apprentissage"

- Interdiction ABSOLUE :
  - "développement de la rigueur"
  - "approfondissement des connaissances"
  - "ouverture d’esprit"
  - "analyse critique"
  - "passion pour"
  - "intérêt pour"
  - toute formulation vague ou académique

- Interdiction de faire des activités vides :
  ❌ "course à pied : préparation d’un 5 km"
  ❌ "lecture : loisir"
  ❌ "cinéma : passion"

- Exemples attendus :
  ✔️ "Course à pied : entraînement régulier et participation à des courses locales, développant endurance et discipline"
  ✔️ "Piano : pratique depuis 5 ans, apprentissage progressif et rigoureux"
  ✔️ "Voyages : découverte de plusieurs pays, développant adaptabilité et ouverture culturelle"
  ✔️ "Bénévolat : engagement associatif ponctuel, gestion d’événements étudiants"

- Maximum 2 à 3 activités
- Chaque activité doit apporter une information utile ou valorisante

RÈGLES DE STYLE :
- Phrases courtes.
- Une idée par bullet.
- Aucun adjectif vide : motivé, dynamique, passionné, polyvalent, excellent.
- Aucun ton promotionnel.
- Aucun markdown.

FORMAT DE SORTIE OBLIGATOIRE :

EDUCATION:
DEGREE: <intitulé du diplôme>
SCHOOL: <école ou université>
LOCATION: <Ville, Pays>
DATES: <MMM YYYY – MMM YYYY ou MMM YYYY – Present>
DETAILS:
- <détail 1>
- <détail 2>

EXPERIENCES:
ROLE: <intitulé du poste>
COMPANY: <nom de la structure>
DATES: <MMM YYYY – MMM YYYY ou MMM YYYY – Present>
LOCATION: <Ville, Pays>
TYPE: <Stage / Alternance / Job étudiant / Projet associatif / etc. si fourni sinon vide>
BULLETS:
- <bullet 1>
- <bullet 2>
- <bullet 3>

SKILLS:
<2 à 4 lignes>

ACTIVITIES:
<une activité par ligne>

CONTRAINTES DE SORTIE :
- Tu génères UNIQUEMENT les sections suivantes, dans cet ordre exact :
  EDUCATION:
  EXPERIENCES:
  SKILLS:
  ACTIVITIES:
- Tu ne génères PAS de section LANGUAGES: ou LANGUES: séparée.
- Tu ne génères PAS le nom.
- Tu ne génères PAS les coordonnées.
- Tu ne génères PAS d’accroche.
- Tu ne génères PAS de texte explicatif.

PROFIL :
Nom : {payload["full_name"]}
Ville : {payload["city"]}

FORMATION :
{payload["education"]}

EXPÉRIENCES :
{payload["experiences"]}

COMPÉTENCES :
{payload["skills"]}

LANGUES :
{payload["languages"]}

CENTRES D’INTÉRÊT :
{payload.get("interests","")}

Génère uniquement le CV structuré.
"""
    

def build_mandatory_experience_anchor(payload: dict) -> str:
    """
    Construit la liste FIGÉE des expériences que le LLM DOIT inclure.
    Le LLM peut reformuler les bullets mais ne peut JAMAIS supprimer une expérience.
    """
    raw_exps = parse_raw_experiences_input(payload.get("experiences", ""))
    if not raw_exps:
        return ""
    
    lines = []
    lines.append("")
    lines.append("⚠️ LISTE OBLIGATOIRE DES EXPÉRIENCES — NE PAS SUPPRIMER :")
    lines.append("Tu DOIS inclure toutes les expériences suivantes dans le CV généré.")
    lines.append("Tu peux reformuler les bullets mais tu ne peux JAMAIS omettre une expérience.")
    lines.append("")
    for i, exp in enumerate(raw_exps, 1):
        role = (exp.get("role") or "").strip()
        company = (exp.get("company") or "").strip()
        dates = (exp.get("dates") or "").strip()
        lines.append(f"  {i}. {role} — {company} ({dates})")
    lines.append("")
    lines.append(f"TOTAL : {len(raw_exps)} expériences à inclure TOUTES.")
    lines.append("")
    return "\n".join(lines)

def generate_cv_text(payload: Dict[str, Any]) -> str:
    if not client:
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY manquante sur le serveur.")

    sector = (payload.get("sector") or "").lower()

    if "finance" in sector:
        prompt = build_prompt_finance(payload)
    elif "audit" in sector:
        prompt = build_prompt_audit(payload)
    elif is_management_sector(sector):
        prompt = build_prompt_management(payload)
    elif "droit" in sector or "juridique" in sector or "juriste" in sector or "avocat" in sector:
        prompt = build_prompt_droit(payload)
    else:
        prompt = build_prompt(payload)

    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
    )

    cv_text = resp.choices[0].message.content.strip()
    cv_text = clean_cv_output(cv_text)
    cv_text = apply_strip_padding_to_cv(cv_text, payload=payload)

    if is_legal_sector(payload.get("sector", "")):
        if "DEGREE:" not in cv_text or "ROLE:" not in cv_text:
            print("=== WARNING DROIT: FORMAT STRUCTURÉ INCOMPLET ===")

    print("=== RAW CV TEXT ===")
    print(cv_text)
    print("=== END RAW CV TEXT ===")

    expected_edu_blocks = count_education_blocks(payload.get("education", ""))
    actual_edu_blocks = cv_text.count("DEGREE:")

    if actual_edu_blocks < expected_edu_blocks:
        print("=== WARNING EDUCATION: BLOCS MANQUANTS ===")

    cv_text = ensure_required_sections(cv_text, payload)
    return cv_text

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
ITEM_SPACING = Pt(0.2)   # espace entre 2 formations / 2 expériences
SECTION_SPACING = Pt(0)  # le titre de section a déjà space_before via normalize

from docx.oxml.ns import qn

def count_education_blocks(raw_education: str) -> int:
    blocks = []
    current = []
    for line in (raw_education or "").splitlines():
        if line.strip():
            current.append(line.strip())
        else:
            if current:
                blocks.append(current)
                current = []
    if current:
        blocks.append(current)
    return len(blocks)

def count_experience_blocks(raw_experiences: str) -> int:
    blocks = []
    current = []

    for line in (raw_experiences or "").splitlines():
        if line.strip():
            current.append(line.strip())
        else:
            if current:
                blocks.append(current)
                current = []

    if current:
        blocks.append(current)

    return len(blocks)

def rebuild_education_from_input(raw_education: str) -> list[str]:
    """
    Reconvertit l'input brut formation en pseudo-format structuré minimal.
    Gère 3 formats :
    A) Multi-lignes : chaque champ sur sa propre ligne (format front natif)
       "Master in Management\nESSCA\nSept 2022 – Mai 2027\nLyon, France\nSpécialisation Finance"
    B) Tout sur une ligne : "ESSCA grande ecole sept 2022 juin 2027 lyon france specialisation finance"
    C) Mix : infos partiellement sur plusieurs lignes
    """
    DATE_PAT = re.compile(
        r"(?:jan|fév|fev|mar|avr|apr|mai|may|juin|jun|juil|jul|août|aout|aug|sept|sep|oct|nov|déc|dec)"
        r"\.?\s+\d{4}\s*[–\-]\s*"
        r"(?:(?:jan|fév|fev|mar|avr|apr|mai|may|juin|jun|juil|jul|août|aout|aug|sept|sep|oct|nov|déc|dec)"
        r"\.?\s+\d{4}|aujourd'hui|present|en cours|\d{4})",
        re.IGNORECASE,
    )
    YEAR_RANGE_PAT = re.compile(r"\b(\d{4})\s*[–\-]\s*(\d{4}|\bau(jourd'hui|jourdhui)?\b|present|today)", re.IGNORECASE)
    DEGREE_KEYWORDS = re.compile(
        r"^(master|bachelor|licence|bba|mba|ms\b|msc\b|llm\b|m1\b|m2\b|m1/m2|"
        r"programme grande[- ]é?cole|grande[- ]é?cole|programme grande ecole|grande ecole|"
        r"baccalauréat|baccalaureat|bac\b|cpge|prépa|prepa|"
        r"exchange program|exchange semester|échange académique|semester abroad|study abroad|"
        r"visiting student|diplôme|diplome|certificate|dut\b|but\b|bts\b|deug\b|dea\b|des\b)",
        re.IGNORECASE,
    )
    SCHOOL_KEYWORDS = re.compile(
        r"(university|université|universite|school|école|ecole|institute|institut|"
        r"college|collège|iep\b|sciences[- ]?po|hec\b|edhec|essec|escp|emlyon|em[- ]|"
        r"ieseg|ieseg|inseec|audencia|kedge|skema|neoma|rennes\s*sb|grenoble\s*em|"
        r"esg\b|iseg\b|iseg|sup[- ]?de[- ]?co|ict|ict\b|ict\s|"
        r"ensae|ensai|polytechnique|normale[- ]?sup|"
        r"panthéon|pantheon|sorbonne|paris[- ]?\d|dauphine|assas|nanterre|"
        r"lyon\s*\d|bordeaux|strasbourg|montpellier|toulouse|aix-marseille|"
        r"lycée|lycee)",
        re.IGNORECASE,
    )

    # ── Étape 1 : splitter par blocs (lignes vides ou nouveaux diplômes) ──────
    raw_blocks: list[list[str]] = []
    current: list[str] = []
    for raw_line in (raw_education or "").splitlines():
        line = raw_line.strip()
        if not line:
            if current:
                raw_blocks.append(current)
                current = []
        else:
            current.append(line)
    if current:
        raw_blocks.append(current)

    # ── Étape 2 : si 1 seul gros bloc, tenter de sub-splitter ───────────────
    if len(raw_blocks) == 1 and len(raw_blocks[0]) > 2:
        split_blocks: list[list[str]] = []
        cur_block: list[str] = []
        cur_has_anchor = False
        for line in raw_blocks[0]:
            is_new_degree = cur_block and DEGREE_KEYWORDS.match(line.strip())
            is_new_school = cur_block and cur_has_anchor and SCHOOL_KEYWORDS.search(line.strip()) and not DATE_PAT.search(line)
            if is_new_degree or is_new_school:
                split_blocks.append(cur_block)
                cur_block = [line]
                cur_has_anchor = bool(DEGREE_KEYWORDS.match(line.strip()) or SCHOOL_KEYWORDS.search(line.strip()))
            else:
                cur_block.append(line)
                if not cur_has_anchor and (DEGREE_KEYWORDS.match(line.strip()) or SCHOOL_KEYWORDS.search(line.strip())):
                    cur_has_anchor = True
        if cur_block:
            split_blocks.append(cur_block)
        if len(split_blocks) > 1:
            raw_blocks = split_blocks

    def _extract_dates_and_strip(text: str) -> tuple[str, str]:
        """Extrait dates du texte et renvoie (dates_str, text_sans_dates)."""
        m = DATE_PAT.search(text)
        if m:
            return m.group(0).strip(), (text[:m.start()].rstrip(" ,–-") + text[m.end():]).strip()
        # Fallback: année – année
        m2 = YEAR_RANGE_PAT.search(text)
        if m2:
            return m2.group(0).strip(), (text[:m2.start()].rstrip(" ,–-") + text[m2.end():]).strip()
        return "", text

    COUNTRY_WORDS = {"france", "allemagne", "espagne", "italie", "portugal", "belgique",
                     "suisse", "royaume-uni", "états-unis", "etats-unis", "canada", "pays-bas",
                     "luxembourg", "autriche", "finlande", "suède", "danemark", "irlande",
                     "uk", "usa", "gb", "netherlands", "germany", "spain", "italy"}
    LOCATION_WORDS = COUNTRY_WORDS | {"paris", "lyon", "bordeaux", "marseille", "lille", "nantes",
                                       "toulouse", "strasbourg", "nice", "rennes", "montpellier",
                                       "london", "new york", "amsterdam", "berlin", "madrid",
                                       "milan", "rome", "lisbon", "lisbonne", "milan", "munich",
                                       "hong kong", "singapour", "singapore", "dubai", "dubai",
                                       "boston", "chicago", "toronto", "montreal", "cergy", "gex",
                                       "chavannes", "bruxelles", "brussels", "zurich", "genève",
                                       "warwick", "coventry", "edinburgh", "Glasgow", "oxford",
                                       "cambridge", "rotterdam", "stockholm", "oslo", "copenhague"}

    def _looks_like_location(text: str) -> bool:
        words = {w.lower().strip(".,;") for w in text.split()}
        return bool(words & LOCATION_WORDS)

    def _parse_block(block: list[str]) -> tuple[str, str, str, str, list[str]]:
        """
        Renvoie (degree, school, dates, location, details).
        Gère aussi bien le format multi-lignes que tout-sur-une-ligne.
        """
        degree = school = dates = location = ""
        details: list[str] = []

        # Mode A : plusieurs lignes bien séparées
        if len(block) >= 2:
            first = block[0].strip()
            second = block[1].strip() if len(block) > 1 else ""
            dates, first_stripped = _extract_dates_and_strip(first)

            # ✅ Détecter si la 1ère ligne est l'ÉCOLE et la 2ème le DIPLÔME
            # (format front natif : "EDHEC Business School\nBachelor in Business Administration...")
            first_is_school = bool(SCHOOL_KEYWORDS.search(first_stripped)) and not bool(DEGREE_KEYWORDS.match(first_stripped))
            second_is_degree = bool(DEGREE_KEYWORDS.match(second)) if second else False

            if first_is_school and second_is_degree:
                # Cas "École\nDiplôme\nDates\nLieu\nDétails"
                school = first_stripped
                degree = second
                # Traiter les lignes à partir de la 3ème
                for line in block[2:]:
                    line_stripped = line.strip()
                    if not line_stripped:
                        continue
                    d, stripped = _extract_dates_and_strip(line_stripped)
                    if d and not dates:
                        dates = d
                        if stripped and _looks_like_location(stripped) and not location:
                            location = stripped
                        continue
                    if not location and _looks_like_location(line_stripped) and len(line_stripped.split()) <= 5:
                        location = line_stripped
                        continue
                    detail_clean = line_stripped.lstrip("-•").strip()
                    if detail_clean:
                        details.append(detail_clean)
            else:
                # Cas standard "Diplôme – École\nDates\nLieu\nDétails" ou "Diplôme\nÉcole\nDates..."
                # Chercher "Degree – School" sur la 1ère ligne
                for sep in [" – ", " - "]:
                    if sep in first_stripped:
                        parts = first_stripped.split(sep, 1)
                        if SCHOOL_KEYWORDS.search(parts[1]) or not SCHOOL_KEYWORDS.search(second):
                            degree = parts[0].strip()
                            school = parts[1].strip()
                        else:
                            degree = first_stripped
                        break
                else:
                    degree = first_stripped

                # Lignes suivantes : détecter school, dates, location, details
                for line in block[1:]:
                    line_stripped = line.strip()
                    if not line_stripped:
                        continue

                    d, stripped = _extract_dates_and_strip(line_stripped)
                    if d and not dates:
                        dates = d
                        if stripped and not school and SCHOOL_KEYWORDS.search(stripped):
                            school = stripped
                        elif stripped and _looks_like_location(stripped) and not location:
                            location = stripped
                        continue

                    if not school and SCHOOL_KEYWORDS.search(line_stripped) and not DATE_PAT.search(line_stripped):
                        school = line_stripped
                        continue

                    if not location and _looks_like_location(line_stripped) and len(line_stripped.split()) <= 5:
                        location = line_stripped
                        continue

                    detail_clean = line_stripped.lstrip("-•").strip()
                    if detail_clean:
                        details.append(detail_clean)

        else:
            # Mode B : tout sur 1 ligne — parser token par token
            line = block[0] if block else ""
            dates, line = _extract_dates_and_strip(line)

            # Extraire location (mots LOCATION_WORDS à la fin ou après virgule)
            # Simple heuristique : derniers tokens après la date
            tokens = line.split(",")
            loc_candidates = []
            remaining_tokens = []
            for tok in reversed(tokens):
                if _looks_like_location(tok) and len(tok.strip().split()) <= 4:
                    loc_candidates.insert(0, tok.strip())
                else:
                    remaining_tokens.insert(0, tok.strip())
                    break
            # Si loc trouvée
            if loc_candidates:
                location = ", ".join(loc_candidates)
                line = ", ".join(remaining_tokens) if remaining_tokens else line

            # Chercher school vs degree dans ce qui reste
            # Si le SCHOOL_KEYWORD est dans les premiers mots → school en premier
            words = line.strip().split()
            if words and SCHOOL_KEYWORDS.search(" ".join(words[:4])):
                # Format "ESSCA grande ecole spécialisation finance"
                school_words = []
                degree_words = []
                found_degree = False
                for w in words:
                    if not found_degree and (DEGREE_KEYWORDS.match(w) or w.lower() in {"grande", "ecole", "école", "d'"}):
                        degree_words.append(w)
                        if DEGREE_KEYWORDS.match(w):
                            found_degree = True
                    elif not found_degree:
                        school_words.append(w)
                    else:
                        degree_words.append(w)

                school = " ".join(school_words).strip()
                degree_raw = " ".join(degree_words).strip()
                # Le reste après les mots-clés de diplôme = détails
                degree_match = DEGREE_KEYWORDS.search(degree_raw)
                if degree_match:
                    degree = degree_raw[:degree_match.end()].strip()
                    rest = degree_raw[degree_match.end():].strip().lstrip("–-,").strip()
                    if rest:
                        details.append(rest)
                else:
                    degree = degree_raw
            else:
                # Chercher séparateur "–" ou "-"
                for sep in [" – ", " - "]:
                    if sep in line:
                        idx = line.index(sep)
                        degree = line[:idx].strip()
                        rest = line[idx + len(sep):].strip()
                        if SCHOOL_KEYWORDS.search(rest):
                            school = rest
                        else:
                            details.append(rest)
                        break
                else:
                    degree = line.strip()

        return degree, school, dates, location, details

    out: list[str] = []
    for block in raw_blocks:
        degree, school, dates, location, details = _parse_block(block)

        # Normaliser les dates
        dates = translate_months_fr(dates) if dates else ""

        # ✅ Si le "degree" contient en fait des matières/sujets (bac), le déplacer en détail
        BAC_SUBJECT_WORDS = {
            "physique", "chimie", "mathématiques", "mathematiques", "maths",
            "svt", "histoire", "géographie", "philosophie", "philo",
            "économie", "economie", "ses", "informatique", "arts", "français",
            "moyenne", "mention",
        }
        degree_words_low = {w.lower().strip(".,;") for w in degree.split()}
        if len(degree_words_low & BAC_SUBJECT_WORDS) >= 2:
            # Le "degree" est en fait des matières/spécialités → déplacer en details
            details = [degree] + details
            # Déduire le vrai degré depuis l'école si possible
            if "baccalauréat" in school.lower() or "baccalaureat" in school.lower() or "bac" in school.lower():
                degree = "Baccalauréat général"
            else:
                degree = "Baccalauréat"

        # ✅ Nettoyer le nom d'école des mots de diplôme redondants
        for bac_word in ["baccalauréat général", "baccalaureat general", "bac général", "bac general",
                          "baccalauréat", "baccalaureat"]:
            school = re.sub(rf"(?i)\s*{re.escape(bac_word)}\s*", " ", school).strip()
        # Capitaliser
        if school:
            school = school[0].upper() + school[1:]

        out.append(f"DEGREE: {degree}")
        out.append(f"SCHOOL: {school}")
        out.append(f"LOCATION: {location}")
        out.append(f"DATES: {dates}")
        out.append("DETAILS:")
        if details:
            # Joindre les lignes de détails : si plusieurs courtes → une seule séparée par virgules
            if len(details) == 1 or all(len(d) < 50 for d in details):
                out.append("- " + ", ".join(d.rstrip(".,") for d in details))
            else:
                for d in details:
                    out.append(f"- {d}")
        else:
            out.append("- ")
        out.append("")

    return out
    """
    Reconvertit l'input brut formation en pseudo-format structuré minimal.
    Supporte le format : "Degree – School, Dates, Location" sur une ligne, puis détails.
    Chaque bloc éducation est séparé par une ligne vide OU commence par un pattern de diplôme.
    """
    # D'abord on split par ligne vide
    raw_blocks = []
    current = []
    for raw in (raw_education or "").splitlines():
        line = (raw or "").strip()
        if not line:
            if current:
                raw_blocks.append(current)
                current = []
        else:
            current.append(line)
    if current:
        raw_blocks.append(current)

    # Si un seul gros bloc, on tente de le découper intelligemment
    if len(raw_blocks) == 1 and len(raw_blocks[0]) > 2:
        date_pat = re.compile(
            r"(?:Jan|Fév|Feb|Mar|Avr|Apr|Mai|May|Juin|Jun|Juil|Jul|Août|Aug|Sept|Sep|Oct|Nov|Déc|Dec)"
            r"\.?\s+\d{4}\s*[–\-]",
            re.IGNORECASE
        )
        # Mots clés qui indiquent TOUJOURS un nouveau bloc de formation
        new_block_keywords = re.compile(
            r"^(exchange program|exchange semester|échange académique|semester abroad|study abroad"
            r"|master\s|master\d|bachelor\s|licence\s|bba\s|mba\s|programme grande école|cpge|prépa|baccalauréat)",
            re.IGNORECASE
        )
        split_blocks = []
        current_block = []
        current_has_degree = False  # pour savoir si le bloc courant a déjà un diplôme
        for line in raw_blocks[0]:
            is_keyword_new = current_block and new_block_keywords.match(line.strip())
            # On split sur dates SEULEMENT si le bloc courant a déjà un diplôme
            is_date_new = current_block and current_has_degree and date_pat.search(line)
            if is_keyword_new or is_date_new:
                split_blocks.append(current_block)
                current_block = [line]
                current_has_degree = bool(new_block_keywords.match(line.strip()))
            else:
                current_block.append(line)
                if not current_has_degree and (
                    new_block_keywords.match(line.strip()) or
                    re.search(r"(master|bachelor|programme|licence|bba|mba)", line, re.IGNORECASE)
                ):
                    current_has_degree = True
        if current_block:
            split_blocks.append(current_block)
        if len(split_blocks) > 1:
            raw_blocks = split_blocks

    def parse_edu_first_line(line: str):
        """Parse 'Degree – School, Dates, Location' depuis une ligne."""
        degree, school, dates, location = "", "", "", ""

        date_pat = re.search(
            r"((?:Jan|Fév|Feb|Mar|Avr|Apr|Mai|May|Juin|Jun|Juil|Jul|Août|Aug|Sept|Sep|Oct|Nov|Déc|Dec)"
            r"\.?\s+\d{4}\s*[–\-]\s*"
            r"(?:(?:Jan|Fév|Feb|Mar|Avr|Apr|Mai|May|Juin|Jun|Juil|Jul|Août|Aug|Sept|Sep|Oct|Nov|Déc|Dec)"
            r"\.?\s+\d{4}|Aujourd'hui|Present|En cours|\d{4}))",
            line, re.IGNORECASE
        )
        if date_pat:
            dates = date_pat.group(1).strip()
            line = (line[:date_pat.start()].rstrip(" ,–-") + line[date_pat.end():]).strip()

        loc_pat = re.search(
            r",\s*([A-ZÀ-Ÿa-zà-ÿ\u00C0-\u017E][A-ZÀ-Ÿa-zà-ÿ\u00C0-\u017E\s\-']+(?:,\s*[A-ZÀ-Ÿa-zà-ÿ\u00C0-\u017E][A-ZÀ-Ÿa-zà-ÿ\u00C0-\u017E\s\-']+)?)\s*$",
            line
        )
        if loc_pat:
            candidate = loc_pat.group(1).strip()
            parts = [p.strip() for p in candidate.split(",")]
            if all(1 <= len(p.split()) <= 4 for p in parts):
                location = candidate
                line = line[:loc_pat.start()].rstrip(" ,").strip()

        for sep in [" – ", " - "]:
            if sep in line:
                idx = line.index(sep)
                degree = line[:idx].strip()
                school = line[idx + len(sep):].strip()
                break
        else:
            degree = line.strip()

        return degree, school, dates, location

    out = []
    for block in raw_blocks:
        first_line = block[0] if block else ""
        details = block[1:] if len(block) > 1 else []

        degree, school, dates, location = parse_edu_first_line(first_line)

        out.append(f"DEGREE: {degree}")
        out.append(f"SCHOOL: {school}")
        out.append(f"LOCATION: {location}")
        out.append(f"DATES: {dates}")
        out.append("DETAILS:")

        detail_lines = [d.strip().lstrip("-").strip() for d in details if d.strip()]
        if detail_lines:
            # Joindre les lignes de détails en une seule ligne séparée par des virgules
            # sauf si c'est déjà une phrase longue (> 60 chars)
            joined = ", ".join(detail_lines)
            out.append(f"- {joined}")
        else:
            out.append("- ")
        out.append("")

    return out
def rebuild_experiences_from_input(raw_experiences: str) -> list[str]:
    exps = parse_raw_experiences_input(raw_experiences)
    out = []

    for exp in exps:
        out.append(f"ROLE: {(exp.get('role') or '').strip()}")
        out.append(f"COMPANY: {(exp.get('company') or '').strip()}")
        out.append(f"DATES: {(exp.get('dates') or '').strip()}")
        out.append(f"LOCATION: {(exp.get('location') or '').strip()}")
        out.append(f"TYPE: {(exp.get('type') or '').strip()}")
        out.append("BULLETS:")

        bullets = [b.strip() for b in (exp.get("bullets") or []) if b and b.strip()]
        if bullets:
            for b in bullets:
                out.append(f"- {b}")
        else:
            out.append("- Expérience professionnelle.")
        out.append("")

    return out


def rebuild_activities_from_input(raw_interests: str) -> list[str]:
    lines = []
    for raw in (raw_interests or "").splitlines():
        text = clean_punctuation_text((raw or "").strip())
        if not text:
            continue
        if ":" in text:
            head, tail = text.split(":", 1)
            head = head.strip()
            tail = tail.strip().rstrip(".")
            lines.append(f"{head} : {tail}.")
        else:
            lines.append(text)
    return lines


def build_skills_from_payload(payload: Dict[str, Any]) -> list[str]:
    lines = []

    raw_certifications = clean_punctuation_text((payload.get("certifications") or "").strip())
    raw_skills = clean_punctuation_text((payload.get("skills") or "").strip())
    raw_languages = clean_punctuation_text((payload.get("languages") or "").strip())

    if raw_certifications:
        lines.append(f"Certifications : {raw_certifications}")

    if raw_skills:
        lines.append(f"Maîtrise des logiciels : {raw_skills}")
    else:
        lines.append("Maîtrise des logiciels : Pack Office")

    if raw_languages:
        lines.append(f"Langues : {raw_languages}")
    else:
        lines.append("Langues : Français")

    return lines


def ensure_required_sections(cv_text: str, payload: Dict[str, Any]) -> str:
    sections = _split_sections(cv_text)

    education_lines = sections.get("EDUCATION") or []
    experience_lines = sections.get("EXPERIENCES") or []
    skills_lines = sections.get("SKILLS") or []
    activity_lines = sections.get("ACTIVITIES") or []

    expected_edu_blocks = count_education_blocks(payload.get("education", ""))
    actual_edu_blocks = sum(
        1 for line in education_lines
        if (line or "").strip().startswith("DEGREE:")
    )

    # ✅ si EDUCATION manque ou si le LLM a oublié un ou plusieurs diplômes,
    # on reconstruit depuis l'input utilisateur
    if (
        not education_lines
        or actual_edu_blocks < expected_edu_blocks
    ):
        education_lines = rebuild_education_from_input(payload.get("education", ""))

    expected_exp_blocks = count_experience_blocks(payload.get("experiences", ""))
    actual_exp_blocks = sum(
        1 for line in experience_lines
        if (line or "").strip().startswith("ROLE:")
    )
    
    if not experience_lines or actual_exp_blocks == 0:
        experience_lines = rebuild_experiences_from_input(payload.get("experiences", ""))
    else:
        # Vérifie que les bullets utilisateur ne sont pas supprimés
        user_exps = parse_raw_experiences_input(payload.get("experiences", ""))
        for i, user_exp in enumerate(user_exps):
            user_bullets = [b for b in (user_exp.get("bullets") or []) if b.strip()]
            if len(user_bullets) >= 3:
                # Cherche l'expérience correspondante dans le LLM output
                role_key = (user_exp.get("role") or "").lower()[:20]
                for j, line in enumerate(experience_lines):
                    if role_key and role_key in (line or "").lower():
                        # Compte les bullets après ce ROLE:
                        bullet_count = 0
                        k = j + 1
                        while k < len(experience_lines) and not (experience_lines[k] or "").startswith("ROLE:"):
                            if (experience_lines[k] or "").strip().startswith("-"):
                                bullet_count += 1
                            k += 1
                        if bullet_count < len(user_bullets):
                            print(f"=== WARNING: bullets supprimés pour {role_key}, reconstruction ===")
                        break

    if not skills_lines:
        skills_lines = build_skills_from_payload(payload)

    if not activity_lines:
        activity_lines = rebuild_activities_from_input(payload.get("interests", ""))

    rebuilt = [
        "EDUCATION:",
        *education_lines,
        "",
        "EXPERIENCES:",
        *experience_lines,
        "",
        "SKILLS:",
        *skills_lines,
        "",
        "ACTIVITIES:",
        *activity_lines,
    ]

    return clean_cv_output("\n".join(rebuilt))

def _keep_lines(paragraph: Paragraph, keep_lines=True, keep_next=False):
    """
    Empêche Word/LibreOffice de couper ce paragraphe sur 2 pages,
    et optionnellement le colle au paragraphe suivant.
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()

    if keep_lines:
        el = OxmlElement("w:keepLines")
        pPr.append(el)

    if keep_next:
        el = OxmlElement("w:keepNext")
        pPr.append(el)

def _row_cant_split(row):
    """
    Empêche une ligne de tableau d’être coupée entre 2 pages.
    C’est LE truc qui évite le rendu “moche/coupé”.
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)

def translate_months_fr(text: str) -> str:
    """
    Normalise les mois :
    - Anglais complet ou abrégé -> abréviation FR (Janv, Fév, Mars, Avr, Mai, Juin, Juil, Août, Sept, Oct, Nov, Déc)
    - Français complet -> abréviation FR
    On évite l'effet 'Septt' en ne remplaçant que des mots entiers.
    """
    # ✅ Normaliser "Present", "Today", "Actuellement" → "Aujourd'hui"
    text = re.sub(r"(?i)\b(present|today|actuellement|en cours)\b", "Aujourd'hui", text)

    # ✅ Normaliser toutes les dates tout-en-majuscules (ex: "AVR 2024", "AOUT 2024", "ETE 2023")
    text = re.sub(r"\b([A-ZÉÀÂÄÈÊËÎÏÔÙÛÜ]{3,})\b",
                  lambda m: m.group(1).capitalize(),
                  text)

    # ✅ "Été 2022" / "Eté 2023" / "ETE 2022" → convertir en mois réels
    text = re.sub(r"(?i)\b[EÉ]t[EÉ]\s+(\d{4})\s*[–\-]\s*[EÉ]t[EÉ]\s+(\d{4})", r"Juin \1 – Août \2", text)
    text = re.sub(r"(?i)\b[EÉ]t[EÉ]\s+(\d{4})", r"Juin \1", text)
    text = re.sub(r"(?i)\bsummer\s+(\d{4})", r"Juin \1", text)

    # Normaliser la casse (Sept au lieu de SEPT)
    text = re.sub(r"\b(SEPT|OCT|NOV|DÉC|DEC|JANV|FÉV|FEV|AVR|JUIN|JUIL|AOÛT|AOUT)\b",
                  lambda m: m.group(0).capitalize(),
                  text)
    if not text:
        return text

    patterns = {
        # EN full → FR abrégé avec point (mois vraiment abrégés) ou sans (mois courts)
        r"(?i)\bJanuary\b": "Janv.",
        r"(?i)\bFebruary\b": "Fév.",
        r"(?i)\bMarch\b": "Mars",
        r"(?i)\bApril\b": "Avr.",
        r"(?i)\bMay\b": "Mai",
        r"(?i)\bJune\b": "Juin",
        r"(?i)\bJuly\b": "Juil.",
        r"(?i)\bAugust\b": "Août",
        r"(?i)\bSeptember\b": "Sept.",
        r"(?i)\bOctober\b": "Oct.",
        r"(?i)\bNovember\b": "Nov.",
        r"(?i)\bDecember\b": "Déc.",

        # EN short
        r"(?i)\bJan\b": "Janv.",
        r"(?i)\bFeb\b": "Fév.",
        r"(?i)\bMar\b": "Mars",
        r"(?i)\bApr\b": "Avr.",
        r"(?i)\bJun\b": "Juin",
        r"(?i)\bJul\b": "Juil.",
        r"(?i)\bAug\b": "Août",
        r"(?i)\bSep\b": "Sept.",
        r"(?i)\bOct\b": "Oct.",
        r"(?i)\bNov\b": "Nov.",
        r"(?i)\bDec\b": "Déc.",

        # FR full → abrégé
        r"(?i)\bJanvier\b": "Janv.",
        r"(?i)\bFévrier\b": "Fév.",
        r"(?i)\bFevrier\b": "Fév.",
        r"(?i)\bMars\b": "Mars",
        r"(?i)\bAvril\b": "Avr.",
        r"(?i)\bMai\b": "Mai",
        r"(?i)\bJuin\b": "Juin",
        r"(?i)\bJuillet\b": "Juil.",
        r"(?i)\bAoût\b": "Août",
        r"(?i)\bAout\b": "Août",
        r"(?i)\bSeptembre\b": "Sept.",
        r"(?i)\bOctobre\b": "Oct.",
        r"(?i)\bNovembre\b": "Nov.",
        r"(?i)\bDécembre\b": "Déc.",
        r"(?i)\bDecembre\b": "Déc.",

        # FR déjà abrégé sans point → ajouter le point
        r"(?i)\bJanv\b(?!\.)": "Janv.",
        r"(?i)\bFév\b(?!\.)": "Fév.",
        r"(?i)\bFev\b(?!\.)": "Fév.",
        r"(?i)\bAvr\b(?!\.)": "Avr.",
        r"(?i)\bJuil\b(?!\.)": "Juil.",
        r"(?i)\bSept\b(?!\.)": "Sept.",
        r"(?i)\bOct\b(?!\.)": "Oct.",
        r"(?i)\bNov\b(?!\.)": "Nov.",
        r"(?i)\bDéc\b(?!\.)": "Déc.",
        r"(?i)\bDec\b(?!\.)": "Déc.",
    }

    for pattern, repl in patterns.items():
        text = re.sub(pattern, repl, text)

    return text

def soften_overclaiming(text: str) -> str:
    if not text:
        return text

    t = text.strip()

    replacements = [
        (r"(?i)\bfavorisant la compréhension des données\b", "pour analyse"),
        (r"(?i)\bfacilitant la prise de décisions stratégiques\b", "pour l'équipe"),
        (r"(?i)\bcontribuant aux travaux d'analyse\b", "pour le suivi des analyses"),
        (r"(?i)\brenforçant le réseau associatif\b", "pour développer les partenariats"),
        (r"(?i)\bétablissant des relations durables\b", "assurant le suivi des échanges"),
        (r"(?i)\bassurant une planification efficace\b", "pour le bon déroulement des événements"),
        (r"(?i)\bassurant un service de qualité\b", "dans le respect du flux en caisse"),
        (r"(?i)\bdéveloppant des compétences en stress\b", "mobilisant réactivité et organisation"),
        (r"(?i)\bfavorisant leur progression académique\b", "dans leur apprentissage"),
        (r"(?i)\bstimulant l'intérêt des élèves\b", "dans la compréhension des méthodes"),
        (r"(?i)\bscrupuleusement\b", ""),
        (r"(?i)\bdocumentation exhaustive\b", "documentation claire"),
        (r"(?i)\bsaine et responsable\b", "structurée"),
        (r"(?i)\bconcerts scolaires\b", "pratique collective"),
        (r"(?i)\bpromotion d'un environnement accueillant\b", "accueil des visiteurs"),
        (r"(?i)\bengagement constant pour l'actualité\b", "intérêt pour l'actualité"),
        (r"(?i)\bcultivant discipline, créativité et confiance en soi\b", "développant discipline et constance"),
        (r"(?i)\boptimisant\b", "soutenant"),
        (r"(?i)\bmaximisant\b", "renforçant"),
        (r"(?i)\bgarantissant\b", "assurant"),
        (r"(?i)\baméliorant\b", "soutenant"),
        (r"(?i)\bcréant des présentations percutantes\b", "pour présentation"),
        (r"(?i)\bfacilitant l[''](analyse|information|accès)\b", "pour l'équipe"),
        (r"(?i)\bassurant une communication claire\b", ""),
        (r"(?i)\bcontribuant à l[''](amélioration|optimisation) des processus\b", "dans le cadre du suivi"),
        (r"(?i)\bun approvisionnement optimal\b", "le réapprovisionnement"),
        (r"(?i)\bune expérience client satisfaisante\b", "l'accueil des clients"),
        (r"(?i)\bcompétences interpersonnelles\b", ""),
        (r"(?i)\bsurveillance efficace et proactive\b", "suivi"),
        (r"(?i)\brenforçant la fiabilité des résultats\b", ""),
        (r"(?i)\btransparence financière\b", ""),
        (r"(?i)\bpour une efficacité accrue\b", ""),
        (r"(?i)\boptimisation des ressources\b", ""),
        (r"(?i)\bsoutenant la clarté analytique\b", ""),
        (r"(?i)\bfacilitant ainsi une prise de décision éclairée\b", ""),
        (r"(?i)\brenforçant l[''](expertise|engagement|impact|visibilité)\b", ""),
        (r"(?i)\bfavorisant (leur |la |ainsi )(réussite|progression|dynamique)\b", "dans leur apprentissage"),
        (r"(?i)\bstimulant ainsi leur progression\b", ""),
        (r"(?i)\bcontribuant ainsi à (garantir|assurer|maintenir)\b", ""),
        (r"(?i)\bœuvrant à garantir\b", "assurant"),
        (r"(?i)\bassurant la qualité et la (précision|fiabilité)\b", "avec rigueur"),
        (r"(?i)\bdans le cadre du suivi\b", ""),
        (r"(?i)\bune expérience (positive|mémorable|enrichissante)\b", ""),
        (r"(?i)\bexpérience client (positive et mémorable|positive)\b", "accueil des clients"),
        (r"(?i)\brenforçant ainsi la visibilité\b", ""),
        (r"(?i)\bsoutenant l[''](expérience client)\b", ""),
        (r"(?i)\bune visibilité accrue\b", ""),
        (r"(?i)\bainsi\b", ""),
        (r"(?i)\bdémontrant un engagement\b", ""),
        (r"(?i)\bvisant à (favoriser|améliorer|renforcer)\b", ""),
        (r"(?i)\bpour la réussite scolaire\b", ""),
        (r"(?i)\bau travers de\b", "à travers"),
        (r"(?i)\bencourageant le développement personnel\b", ""),
        (r"(?i)\bfavorisant l[''](épanouissement|excellence)\b", ""),
        (r"(?i)\bpour leur bien-être\b", ""),
        (r"(?i)\bdans un cadre familial\b", ""),
    ]

    for pattern, repl in replacements:
        t = re.sub(pattern, repl, t)

    t = re.sub(r"\s+", " ", t).strip()
    return clean_punctuation_text(t)

def soften_legal_overclaiming(text: str) -> str:
    if not text:
        return text

    t = text.strip()

    replacements = [
        (r"(?i)\bfacilitant (leur|la) compréhension\b", "à destination de l'équipe"),
        (r"(?i)\bgarantissant la conformité\b", "dans le respect des documents fournis"),
        (r"(?i)\bgarantissant (leur|une) précision\b", "avec rigueur"),
        (r"(?i)\bfavorisant\b", "dans le cadre de"),
        (r"(?i)\bcontribuant à\b", ""),
        (r"(?i)\bvisant à\b", ""),
        (r"(?i)\bassurant une gestion optimale\b", "participant au suivi"),
        (r"(?i)\bdéveloppant des stratégies de défense\b", "analysant des cas pratiques"),
        (r"(?i)\bpermettant\b", ""),
        (r"(?i)\boptimale\b", "structurée"),
        (r"(?i)\befficace\b", "rigoureux"),
    ]

    for pattern, repl in replacements:
        t = re.sub(pattern, repl, t)

    t = re.sub(r"\s+", " ", t).strip()
    return clean_punctuation_text(t)

def filter_education_details(details: list[str], raw_education_input: str, is_legal: bool = False) -> list[str]:
    out = []

    # Récupère les lignes "Cours :" exactes depuis l'input utilisateur
    source_courses = []
    for line in (raw_education_input or "").splitlines():
        if line.lower().startswith("cours"):
            _, _, after = line.partition(":")
            source_courses.extend([x.strip() for x in after.split(",") if x.strip()])

    for d in (details or []):
        t = (d or "").strip()
        low = t.lower()

        # En DROIT : si le détail parle de matières/cours, on remplace par les cours exacts utilisateur
        if is_legal and (
            low.startswith("matières fondamentales")
            or low.startswith("cours")
            or "droit du travail" in low
            or "relations collectives" in low
            or "procédure civile" in low
            or "droit des obligations" in low
        ):
            if source_courses:
                t = "Matières fondamentales : " + ", ".join(source_courses) + "."
                if t not in out:
                    out.append(t)
            continue

        banned_keywords = [
            "séminaire", "seminar", "conférence", "conference", "atelier", "workshop",
            "étude de cas", "case study", "participation à",
            "développement des compétences", "capacité à réaliser",
            "préparation approfondie", "acquisition de compétences",
            "compétences fondamentales", "large éventail", "travaux de recherche",
            "exercices de débat", "mémoire à rédiger", "concours professionnels",
            "solide capacité", "méthodologique",
            "travaux pratiques", "analyses approfondies",
            "formation approfondie", "formation complète",
            "formation théorique", "approfondissement",
            "plusieurs matières",
            "formation spécialisée", "formation specialisee",
            "formation en gestion générale", "formation en gestion generale",
            "formation académique", "formation générale", "formation en",
        ]

        if any(k in low for k in banned_keywords):
            continue

        out.append(t)

    # En DROIT : si rien n’a survécu mais qu’il y a des cours source, on force une ligne propre
    if is_legal and not out and source_courses:
        out.append("Matières fondamentales : " + ", ".join(source_courses) + ".")

    return out
def is_course_detail_line(text: str) -> bool:
    if not text:
        return False

    low = text.strip().lower()

    course_markers = [
        "matières fondamentales",
        "matieres fondamentales",
        "cours",
        "key coursework",
    ]

    if any(low.startswith(marker) for marker in course_markers):
        return True

    if "droit du travail" in low or "protection sociale" in low or "relations collectives" in low:
        return True
    if "comptabilité" in low or "comptabilite" in low or "analyse financière" in low or "analyse financiere" in low:
        return True
    if "contrôle de gestion" in low or "controle de gestion" in low:
        return True
    if "stratégie" in low or "strategie" in low or "analyse de marché" in low or "analyse de marche" in low:
        return True

    return False


def normalize_detail_for_dedupe(text: str) -> str:
    if not text:
        return ""
    t = text.strip().lower()
    t = t.replace("’", "'")
    t = re.sub(r"\s+", " ", t)
    t = re.sub(r"[.;:,]+$", "", t)
    return t


def dedupe_preserve_order(lines: list[str]) -> list[str]:
    out = []
    seen = set()

    for line in lines or []:
        txt = clean_punctuation_text((line or "").strip())
        if not txt:
            continue

        key = normalize_detail_for_dedupe(txt)
        if key in seen:
            continue

        seen.add(key)
        out.append(txt)

    return out

def _remove_paragraph(p: Paragraph):
    if p is None:
        return
    el = getattr(p, "_element", None)
    if el is None:
        return
    parent = el.getparent()
    if parent is None:
        return
    parent.remove(el)
    p._p = p._element = None

def _add_table_after(paragraph: Paragraph, rows: int, cols: int):
    """
    Ajoute un tableau juste après le paragraphe.

    Objectifs :
    - 2 colonnes : texte formation à gauche, dates à droite
    - Largeur TOTALE légèrement réduite pour éviter l'effet "dates collées à la marge"
    - Largeurs forcées sur les colonnes (Word + LibreOffice)
    """
    doc = paragraph.part.document
    table = doc.add_table(rows=rows, cols=cols)

    # On ne laisse pas Word/LibreOffice recalculer les largeurs
    table.autofit = False

    if cols == 2:
        try:
        
            # 15,1 cm de texte + 3,9 cm pour les dates
            # → texte bien large + plus de place pour la date (évite qu'elle casse)
            # Largeur totale ≈ 19 cm : très proche du bord mais sans dépasser
            widths = [Cm(15.1), Cm(3.9)]

            # Largeur sur les colonnes
            for col, w in zip(table.columns, widths):
                col.width = w

            # Sécurité : on force aussi la largeur sur chaque cellule
            for row in table.rows:
                for j, w in enumerate(widths):
                    row.cells[j].width = w
        except Exception:
            pass

    # On aligne le tableau à gauche pour qu'il commence au même endroit que le texte normal
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Insérer le tableau juste après le paragraphe ancre
    paragraph._p.addnext(table._tbl)

    # ✅ Empêche les lignes de tableau de se couper entre 2 pages
    try:
        for row in table.rows:
            _row_cant_split(row)
    except Exception:
        pass
    
    return table

def _insert_spacer_after_table(table, parent, space_after):
    spacer_elt = OxmlElement("w:p")
    table._tbl.addnext(spacer_elt)
    spacer = Paragraph(spacer_elt, parent)
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = space_after
    return spacer

def parse_finance_experiences(lines: list[str]) -> list[dict]:
    exps = []
    cur = None
    mode = None

    def push():
        nonlocal cur
        if cur and (cur.get("role") or cur.get("bullets")):
            exps.append(cur)
        cur = None

    for raw in lines:
        line = (raw or "").strip()
        if not line:
            continue

        if line.startswith("ROLE:"):
            push()
            cur = {
                "role": line.replace("ROLE:", "").strip(),
                "company": "",
                "dates": "",
                "location": "",
                "type": "",
                "bullets": [],
            }
            mode = None
            continue

        if not cur:
            continue

        if line.startswith("COMPANY:"):
            cur["company"] = line.replace("COMPANY:", "").strip()
        elif line.startswith("DATES:"):
            cur["dates"] = line.replace("DATES:", "").strip()
        elif line.startswith("LOCATION:"):
            cur["location"] = line.replace("LOCATION:", "").strip()
        elif line.startswith("TYPE:"):
            cur["type"] = line.replace("TYPE:", "").strip()
        elif line.startswith("BULLETS:"):
            mode = "bullets"
        elif mode == "bullets" and line.lstrip().startswith("-"):
            # On tolère les espaces avant le tiret ("  - bullet")
            stripped = line.lstrip()
            bullet_text = stripped[1:].strip()
            if bullet_text:
                bullet_text = re.sub(r"(?i)^participé\s+à\s+", "Contribué à ", bullet_text)
                bullet_text = re.sub(r"(?i)^aidé\s+à\s+", "Soutenu ", bullet_text)
                
                cur["bullets"].append(bullet_text)

    push()
    return exps


PLACEHOLDERS = [
    "%%FULL_NAME%%",
    "%%CV_TITLE%%",
    "%%CONTACT_LINE%%",
    "%%EDUCATION%%",
    "%%EXPERIENCE%%",
    "%%SKILLS%%",
    "%%LANGUAGES%%",
    "%%INTERESTS%%",
]

def parse_raw_experiences_input(raw_text: str) -> list[dict]:
    """
    Parse les expériences directement depuis le texte brut utilisateur.
    Format attendu par bloc :
    ligne 1 = rôle
    ligne 2 = société
    ligne 3 = lieu
    ligne 4 = dates
    ligne 5 = type
    puis bullets commençant par "-"
    """
    blocks = []
    current = []

    for raw in (raw_text or "").splitlines():
        line = (raw or "").strip()
        if not line:
            if current:
                blocks.append(current)
                current = []
            continue
        current.append(line)

    if current:
        blocks.append(current)

    exps = []

    for block in blocks:
        meta = []
        bullets = []

        for line in block:
            if line.startswith("-"):
                bullets.append(line[1:].strip())
            else:
                meta.append(line)

        role = meta[0] if len(meta) > 0 else ""
        company = meta[1] if len(meta) > 1 else ""
        location = meta[2] if len(meta) > 2 else ""
        dates = meta[3] if len(meta) > 3 else ""
        type_ = meta[4] if len(meta) > 4 else ""

        if role or bullets:
            exps.append({
                "role": role,
                "company": company,
                "dates": dates,
                "location": location,
                "type": type_,
                "bullets": bullets,
            })

    return exps
def parse_education_structured(lines: list[str]) -> list[dict]:
    """
    Parse une section EDUCATION structurée avec les tags :
    DEGREE:, SCHOOL:, LOCATION:, DATES:, DETAILS:
    """
    programs = []
    cur = None
    mode = None

    def push():
        nonlocal cur
        if cur and (cur.get("degree") or cur.get("school")):
            # On s'assure d'avoir toujours une liste de détails
            cur.setdefault("details", [])
            programs.append(cur)
        cur = None

    for raw in (lines or []):
        line = (raw or "").strip()
        if not line:
            continue

        if line.startswith("DEGREE:"):
            push()
            cur = {
                "degree": line.replace("DEGREE:", "").strip(),
                "school": "",
                "location": "",
                "dates": "",
                "details": [],
            }
            mode = None
            continue

        if not cur:
            continue

        if line.startswith("SCHOOL:"):
            cur["school"] = line.replace("SCHOOL:", "").strip()
        elif line.startswith("LOCATION:"):
            cur["location"] = line.replace("LOCATION:", "").strip()
        elif line.startswith("DATES:"):
            cur["dates"] = line.replace("DATES:", "").strip()
        elif line.startswith("DETAILS:"):
            mode = "details"
        elif mode == "details" and line.lstrip().startswith("-"):
            txt = line.lstrip()[1:].strip()
            if txt:
                cur["details"].append(txt)

    push()
    return programs

def extract_source_courses_by_education_block(raw_education_input: str) -> list[list[str]]:
    blocks = []
    current = []

    for raw in (raw_education_input or "").splitlines():
        line = (raw or "").strip()
        if not line:
            if current:
                blocks.append(current)
                current = []
            continue
        current.append(line)

    if current:
        blocks.append(current)

    out = []
    for block in blocks:
        courses = []
        for line in block:
            if line.lower().startswith("cours"):
                _, _, after = line.partition(":")
                courses = [x.strip() for x in after.split(",") if x.strip()]
                break
        out.append(courses)

    return out

def extract_non_course_details_by_education_block(raw_education_input: str) -> list[list[str]]:
    blocks = []
    current = []

    for raw in (raw_education_input or "").splitlines():
        line = (raw or "").strip()
        if not line:
            if current:
                blocks.append(current)
                current = []
            continue
        current.append(line)

    if current:
        blocks.append(current)

    out = []
    for block in blocks:
        extra_details = []
        for idx, line in enumerate(block):
            low = line.lower()
            if idx < 4:
                continue
            if low.startswith("cours"):
                continue
            extra_details.append(line.strip())
        out.append(extra_details)

    return out
    
def _no_space_len(s: str) -> int:
    """Longueur d'un texte sans compter les espaces."""
    return len(re.sub(r"\s+", "", s or ""))

def shorten_experience_bullets_with_llm(
    exps: list[dict],
    max_no_space_per_bullet: int = 90,
) -> list[dict]:
    """
    RÉÉCRIT les bullets via l'API pour qu'elles soient plus courtes,
    SANS changer le sens, SANS inventer, SANS '...'
    et SANS JAMAIS changer le nombre de bullets.

    Si l'IA ne respecte pas ça -> on garde la version ORIGINALE.
    """
    if not client:
        return exps  # pas d'API dispo -> on ne touche rien

    # On prépare une version simplifiée pour l'IA
    simple_exps = []
    for e in exps:
        simple_exps.append({
            "role": e.get("role", ""),
            "company": e.get("company", ""),
            "bullets": e.get("bullets", []),
        })

    payload = {
        "max_no_space": max_no_space_per_bullet,
        "experiences": simple_exps,
    }

    prompt = f"""
Tu es recruteur en finance.

On te donne des expériences avec leurs bullet points au format JSON.

POUR CHAQUE BULLET :
- tu réécris la phrase en français,
- tu gardes exactement le même sens (aucune nouvelle mission, aucun nouveau chiffre, aucun nouvel outil),
- tu gardes la structure : verbe d'action + moyen + résultat,
- INTERDIT de commencer par : "Participé", "Aidé", "Effectué", "Travaillé",
- la phrase est complète et se termine par un point,
- maximum {max_no_space_per_bullet} caractères SANS espaces,
- JAMAIS de points de suspension ("...").

INTERDIT :
- changer le nombre d'expériences,
- changer le nombre de bullets,
- réordonner les bullets,
- inventer des éléments.
- toute affirmation d'impact mesurable si elle n'est pas explicitement fournie
- toute amélioration inventée (ex : améliorant la performance, augmentant l'efficacité)

Tu dois renvoyer UNIQUEMENT un JSON de la forme :

{{"experiences": [
  {{"bullets": ["...", "..."]}},
  {{"bullets": ["...", "..."]}}
]}}

Voici le JSON d'entrée :

{json.dumps(payload, ensure_ascii=False)}
"""

    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
    )
    content = resp.choices[0].message.content.strip()

    try:
        data = json.loads(content)
        new_exps = data.get("experiences", [])
    except Exception:
        # Si l'IA ne répond pas en JSON -> on garde TOUT tel quel
        return exps

    # Sécurité maximale : si la longueur ne colle pas, on ne change RIEN
    if not isinstance(new_exps, list) or len(new_exps) != len(exps):
        return exps

    result: list[dict] = []

    for old, new in zip(exps, new_exps):
        if not isinstance(new, dict):
            result.append(old)
            continue

        old_bullets = old.get("bullets") or []
        new_bullets = new.get("bullets")

        if not isinstance(new_bullets, list):
            result.append(old)
            continue

        # ⚠️ SI le nombre de bullets ne correspond pas -> on garde l'ancien
        if len(new_bullets) != len(old_bullets):
            result.append(old)
            continue

        cleaned_bullets = []
        invalid = False
        for b in new_bullets:
            txt = (b or "").strip()
            if not txt:
                invalid = True
                break
            cleaned_bullets.append(txt)
        if invalid:
            result.append(old)
            continue

        updated = dict(old)
        updated["bullets"] = cleaned_bullets
        result.append(updated)

    # Par sécurité : si on a un souci, on renvoie l'original
    if len(result) != len(exps):
        return exps

    return result

def enrich_activities_with_llm(lines: list[str], sector: str = "") -> list[str]:
    try:
        cleaned = [(l or "").strip() for l in (lines or []) if (l or "").strip()]
        if not cleaned:
            return []

        sector_low = (sector or "").lower()

        if "finance" in sector_low:
            sector_hint = "Valorise des qualités comme discipline, rigueur, résilience, persévérance, esprit de compétition, gestion de la pression."
        elif "audit" in sector_low:
            sector_hint = "Valorise des qualités comme rigueur, discipline, constance, précision, persévérance."
        elif "management" in sector_low or "stratégie" in sector_low or "strategie" in sector_low or "conseil" in sector_low:
            sector_hint = "Valorise des qualités comme esprit d'équipe, esprit critique, aisance orale, persévérance, adaptabilité."
        elif "droit" in sector_low or "juridique" in sector_low or "juriste" in sector_low or "avocat" in sector_low:
            sector_hint = "Valorise des qualités comme rigueur, discipline, persévérance, esprit critique, capacité d'analyse."
        else:
            sector_hint = "Valorise des qualités simples, crédibles et cohérentes avec l'activité."

        prompt = f"""
Tu es un expert en rédaction de CV premium.

Ta mission :
Transformer des centres d’intérêt bruts en lignes de CV plus valorisantes et plus élégantes,
sans ajouter de faits faux ou absurdes.

OBJECTIF PRODUIT :
- On doit apporter de la valeur.
- On doit rendre l’activité plus intéressante pour un recruteur.
- On peut faire ressortir des qualités transférables, même si elles ne sont pas écrites mot pour mot,
  à condition qu’elles soient LOGIQUES et cohérentes avec l’activité.
- On n’invente jamais un niveau, une fréquence, un club, un championnat, un événement ou un contexte précis
  qui n’existent pas dans l’entrée.

RÈGLES :
- Tu gardes EXACTEMENT le même nombre de lignes.
- Une ligne en entrée = une ligne en sortie.
- Tu écris en français.
- Aucun markdown.
- Aucun commentaire.
- Pas de puces.
- Pas de guillemets.
- Pas de ton RH cliché.

AUTORISÉ :
- reformuler de manière plus premium
- rendre l’activité plus professionnelle
- ajouter 1 à 3 qualités transférables logiques
- transformer un hobby simple en ligne plus valorisante

INTERDIT :
- inventer une compétition, un club, un niveau, une fréquence, une durée, un événement, un voyage précis
- écrire des clichés comme :
  "culture générale", "perspective internationale", "enrichit la vision du monde",
  "analyse des récits", "forme physique et mentale", "ouverture sur le monde"
- faire des phrases lourdes ou scolaires
- faire trop long

IMPORTANT :
- Si l’entrée contient déjà un niveau précis, tu peux le reprendre.
- Si l’entrée est simple ("Lecture", "Voyages", "Running"), tu peux enrichir intelligemment
  avec des qualités cohérentes, sans inventer de faits précis.
- Exemple attendu :
  "Équitation : pratique à haut niveau développant discipline, résilience et concentration."
- Exemple attendu :
  "Football : pratique développant esprit d’équipe et esprit de compétition."
- Exemple attendu :
  "Lecture : intérêt personnel développant curiosité et esprit d’analyse."
- Exemple attendu :
  "Voyages : découverte de nouvelles cultures développant adaptabilité et ouverture."
- Exemple attendu :
  "Running : pratique régulière développant discipline et persévérance."
- Exemple interdit :
  "Lecture : participation à des clubs de lecture..."
  si ce n'est pas dans l'entrée.
- Exemple interdit :
  "Football : compétitions régionales"
  si ce n'est pas dans l'entrée.
- IMPORTANT : si un chiffre est présent dans l'entrée (ex : "13 pays", "15 ans", "500 personnes"), tu DOIS le conserver dans la sortie.

INDICATION SECTEUR :
{sector_hint}

FORMAT DE SORTIE :
Nom activité : phrase valorisante concise.

ACTIVITÉS :
{chr(10).join(cleaned)}
"""

        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0.3,
            messages=[{"role": "user", "content": prompt}],
        )

        text = resp.choices[0].message.content.strip()
        result = [l.strip() for l in text.split("\n") if l.strip()]

        if len(result) != len(cleaned):
            return cleaned

        return result

    except Exception:
        return lines

def enrich_experience_bullets_with_llm(exps: list[dict], sector: str = "") -> list[dict]:
    if not client:
        return exps

    try:
        flat_bullets = []
        bullet_counts = []

        for exp in exps:
            bullets = [b.strip() for b in (exp.get("bullets") or []) if b and b.strip()]
            bullet_counts.append(len(bullets))
            flat_bullets.extend(bullets)

        if not flat_bullets:
            return exps

        prompt = f"""
Tu es un expert en rédaction de CV juridiques.

Ta mission :
Réécrire légèrement des bullet points d'expérience pour les rendre plus professionnels, sobres et juridiquement crédibles.

RÈGLES STRICTES :
- Tu gardes exactement le même sens.
- Tu n’inventes aucune nouvelle mission.
- Tu n’ajoutes aucun chiffre.
- Tu n’ajoutes aucun outil non mentionné.
- Tu n’ajoutes aucun impact, aucun bénéfice, aucune amélioration implicite.
- Tu n’ajoutes jamais :
  conformité, sécurité juridique, optimisation, efficacité, amélioration continue,
  gain de temps, réduction des risques, fiabilisation, cadre légal, réglementation
  sauf si ces notions sont déjà présentes dans le bullet source.
- Tu ne transformes jamais une expérience non juridique en expérience juridique.
- Tu peux seulement :
  - reformuler,
  - rendre la phrase plus fluide,
  - préciser légèrement le geste déjà écrit, sans dépasser son sens.

IMPORTANT :
- Tu dois garder EXACTEMENT le même nombre de bullet points.
- Tu ne fusionnes jamais deux bullets.
- Une ligne en sortie = un bullet point.
- Tu ne dois rien écrire d’autre.

BULLETS :
{chr(10).join(flat_bullets)}
        """

        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0.2,
            messages=[{"role": "user", "content": prompt}],
        )

        lines = [
            l.strip().lstrip("-").strip()
            for l in resp.choices[0].message.content.split("\n")
            if l.strip()
        ]

        if len(lines) != len(flat_bullets):
            return exps

        idx = 0
        new_exps = []

        for exp, count in zip(exps, bullet_counts):
            copied = dict(exp)
            copied["bullets"] = lines[idx:idx + count]
            idx += count
            new_exps.append(copied)

        return new_exps

    except Exception:
        return exps

def apply_density_to_experiences(
    exps: list[dict],
    is_cv_long: bool = False,
    is_cv_short: bool = False,
    keep_three_for_short: int = 2,
    keep_three_for_normal: int = 2,
    keep_three_for_long: int = 1,
) -> list[dict]:
    for i, exp in enumerate(exps):
        bullets = [b.strip() for b in (exp.get("bullets") or []) if (b or "").strip()]

        if is_cv_short:
            limit = 3
        elif is_cv_long:
            limit = 3 if i < keep_three_for_long else 2
        else:
            limit = 3 if i < keep_three_for_normal else 2

        exp["bullets"] = bullets[:limit]

    return exps

def trim_finance_experiences(
    exps: list[dict],
    is_cv_long: bool = True,
    max_experiences: int = 4,
    max_total_bullets: int = 8,
    min_experiences: int = 2,
    max_no_space_per_bullet: int = 90,
) -> list[dict]:
    """
    NOUVELLE VERSION :
    - NE SUPPRIME PLUS JAMAIS d'expérience.
    - NE SUPPRIME PLUS JAMAIS de bullet.
    - Se contente de nettoyer les vides et, si le CV est long,
      de faire RÉÉCRIRE les bullets via l'API pour les raccourcir.
    """

    # 1) Nettoyage des expériences VRAIMENT vides
    cleaned: list[dict] = []
    for e in exps:
        role = (e.get("role") or "").strip()
        bullets = [b for b in (e.get("bullets") or []) if (b or "").strip()]
        if not role and not bullets:
            continue  # là c'est du vide total, ça ne sert à rien
        e["role"] = role
        e["bullets"] = bullets
        cleaned.append(e)

    if not cleaned:
        return []

    # 2) Si le CV n'est pas long -> on densifie légèrement les premières expériences
    if not is_cv_long:
        for i, e in enumerate(cleaned):
            bullets = [b for b in (e.get("bullets") or []) if (b or "").strip()]
            if len(bullets) >= 3:
                e["bullets"] = bullets[:3]
            else:
                e["bullets"] = bullets
        return cleaned

    # 3) Si le CV est long -> on raccourcit PAR RÉÉCRITURE (pas par suppression)
    cleaned = shorten_experience_bullets_with_llm(
        cleaned,
        max_no_space_per_bullet=max_no_space_per_bullet,
    )

    return cleaned

def trim_experiences_droit(
    exps: list[dict],
    is_cv_long: bool = True,
    is_cv_short: bool = False,
) -> list[dict]:
    cleaned: list[dict] = []

    for e in exps:
        role = normalize_role_text((e.get("role") or "").strip())
        company = (e.get("company") or "").strip()
        dates = (e.get("dates") or "").strip()
        location = (e.get("location") or "").strip()
        type_ = (e.get("type") or "").strip()
        bullets = [b.strip() for b in (e.get("bullets") or []) if (b or "").strip()]

        if not role and not bullets:
            continue

        cleaned.append({
            "role": role,
            "company": company,
            "dates": dates,
            "location": location,
            "type": type_,
            "bullets": bullets,
        })

    def legal_score(exp: dict) -> int:
        text = " ".join([
            exp.get("role", ""),
            exp.get("company", ""),
            exp.get("type", ""),
            " ".join(exp.get("bullets", []))
        ]).lower()

        score = 0

        strong = [
            "jurid", "droit", "social", "veille", "note", "recherche",
            "dossier", "rh", "relations sociales", "travail",
            "administratif", "documents", "rédaction"
        ]
        medium = [
            "suivi", "coordination", "association", "partenariat",
            "tuteur", "tutrice", "tutorat"
        ]
        weak = [
            "vente", "magasin", "encaissement", "stock", "client"
        ]

        for k in strong:
            if k in text:
                score += 5
        for k in medium:
            if k in text:
                score += 2
        for k in weak:
            if k in text:
                score -= 2

        return score

    cleaned.sort(key=legal_score, reverse=True)
    cleaned = cleaned[:4]
    cleaned = apply_density_to_experiences(
        cleaned,
        is_cv_long=is_cv_long,
        is_cv_short=is_cv_short,
        keep_three_for_short=2,
        keep_three_for_normal=2,
        keep_three_for_long=1,
    )
    return cleaned

def trim_experiences_audit(
    exps: list[dict],
    is_cv_long: bool = True,
    is_cv_short: bool = False,
) -> list[dict]:
    cleaned = []

    for e in exps:
        role = normalize_role_text((e.get("role") or "").strip())
        company = (e.get("company") or "").strip()
        dates = (e.get("dates") or "").strip()
        location = (e.get("location") or "").strip()
        type_ = (e.get("type") or "").strip()
        bullets = [b.strip() for b in (e.get("bullets") or []) if (b or "").strip()]

        if not role and not bullets:
            continue

        cleaned.append({
            "role": role,
            "company": company,
            "dates": dates,
            "location": location,
            "type": type_,
            "bullets": bullets,
        })

    def audit_score(exp: dict) -> int:
        text = " ".join([
            exp.get("role", ""),
            exp.get("company", ""),
            exp.get("type", ""),
            " ".join(exp.get("bullets", []))
        ]).lower()

        score = 0

        strong = [
            "audit", "comptabilité", "comptable", "contrôle",
            "contrôle de gestion", "reporting", "analyse financière",
            "trésorerie", "procédure", "vérification", "documentation"
        ]
        medium = [
            "excel", "suivi", "budget", "association", "administratif"
        ]
        weak = [
            "vente", "magasin", "encaissement", "stock", "client"
        ]

        for k in strong:
            if k in text:
                score += 5
        for k in medium:
            if k in text:
                score += 2
        for k in weak:
            if k in text:
                score -= 2

        return score

    cleaned.sort(key=audit_score, reverse=True)
    cleaned = cleaned[:4]
    cleaned = apply_density_to_experiences(
        cleaned,
        is_cv_long=is_cv_long,
        is_cv_short=is_cv_short,
        keep_three_for_short=4,
        keep_three_for_normal=3,
        keep_three_for_long=2,
    )
    return cleaned

def trim_experiences_management(
    exps: list[dict],
    is_cv_long: bool = True,
    is_cv_short: bool = False,
) -> list[dict]:
    cleaned = []

    for e in exps:
        role = normalize_role_text((e.get("role") or "").strip())
        company = (e.get("company") or "").strip()
        dates = (e.get("dates") or "").strip()
        location = (e.get("location") or "").strip()
        type_ = (e.get("type") or "").strip()
        bullets = [b.strip() for b in (e.get("bullets") or []) if (b or "").strip()]

        if not role and not bullets:
            continue

        cleaned.append({
            "role": role,
            "company": company,
            "dates": dates,
            "location": location,
            "type": type_,
            "bullets": bullets,
        })

    def management_score(exp: dict) -> int:
        text = " ".join([
            exp.get("role", ""),
            exp.get("company", ""),
            exp.get("type", ""),
            " ".join(exp.get("bullets", []))
        ]).lower()

        score = 0

        strong = [
            "analyse", "benchmark", "coordination", "gestion de projet",
            "recommandation", "stratégie", "communication", "prospection",
            "partenariat", "présentation", "synthèse", "étude de marché"
        ]
        medium = [
            "association", "événement", "organisation", "suivi", "rédaction"
        ]
        weak = [
            "vente", "magasin", "encaissement", "stock"
        ]

        # pénalise les associations purement étudiantes face à de vrais stages pro
        assoc_student = ["bde", "association étudiante", "enactus", "junior entreprise"]
        for k in assoc_student:
            if k in text:
                score -= 4
        for k in strong:
            if k in text:
                score += 5
        for k in medium:
            if k in text:
                score += 2
        for k in weak:
            if k in text:
                score -= 2

        return score

    cleaned.sort(key=management_score, reverse=True)
    cleaned = cleaned[:4]
    cleaned = apply_density_to_experiences(
        cleaned,
        is_cv_long=is_cv_long,
        is_cv_short=is_cv_short,
        keep_three_for_short=4,
        keep_three_for_normal=3,
        keep_three_for_long=2,
    )
    return cleaned

def shorten_activities_with_llm(
    lines: list[str],
    max_no_space_per_activity: int = 90,
) -> list[str]:
    """
    Réécrit chaque activité pour qu'elle tienne sur une ligne,
    phrase complète, sans '...', SANS jamais changer le NOMBRE d'activités.

    Si l'IA ne respecte pas ça -> on garde les lignes d'origine.
    """
    if not client:
        return lines

    activities = [(l or "").strip() for l in lines if (l or "").strip()]
    if not activities:
        return []

    payload = {
        "max_no_space": max_no_space_per_activity,
        "activities": activities,
    }

    prompt = f"""
Tu es recruteur exigeant.
Tu réécris des bullet points de CV de manière sobre, factuelle et crédible.

On te donne une liste d'activités / centres d'intérêt.

POUR CHAQUE ACTIVITÉ :
- tu gardes UNE activité par ligne (pas de fusion),
- tu réécris en français en gardant le sens,
- style CV (PAS de "je", PAS de "nous", PAS de phrase à la première personne),
- formulation orientée finance : pratique + discipline / exigence / rigueur,
- tu fais une phrase complète qui se termine par un point,
- tu ne mets JAMAIS de points de suspension ("..."),
- la phrase doit faire au maximum {max_no_space_per_activity} caractères SANS espaces.
- INTERDIT d’ajouter un niveau ou une fréquence si ce n’est pas dans l’activité d’origine (ex : "compétition", "national", "régional", "club", "championnat", "hebdomadaire", "quotidien").
- INTERDIT d’ajouter des événements caritatifs, clubs, tournois, compétitions si non mentionnés.
- Structure obligatoire : "<Activité> : <pratique factuelle (sans inventer)> ; <qualités utiles en finance (rigueur, discipline, stress, priorités)>."

INTERDIT :
- changer le nombre d'activités,
- fusionner plusieurs activités en une seule.
- INTERDIT d'ajouter "membre", "équipe", "amateur", "hebdomadaire", "régulière", "occasionnelle"
  si ce n'est pas explicitement dans l'entrée.
  
Réponds UNIQUEMENT avec un JSON de la forme :
{{"activities": ["Activité 1 : ...", "Activité 2 : ...", ...]}}

Voici les activités :

{json.dumps(payload, ensure_ascii=False)}
"""

    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
    )
    content = resp.choices[0].message.content.strip()

    try:
        data = json.loads(content)
        new_acts = data.get("activities", [])
    except Exception:
        # L'IA n'a pas répondu comme il faut -> on garde tout
        return activities

    # Sécurité : si le nombre ne correspond pas -> on garde l'original
    if not isinstance(new_acts, list) or len(new_acts) != len(activities):
        return activities

    out: list[str] = []
    for old, new in zip(activities, new_acts):
        txt = (new or "").strip()
        if not txt:
            # si une activité disparaît -> on annule tout
            return activities
        out.append(txt)

    return out


def trim_activities(
    lines: list[str],
    cv_is_long: bool,
    ideal_max: int = 4,
    cv_is_short: bool = False,
) -> list[str]:
    cleaned = [(l or "").strip() for l in (lines or []) if (l or "").strip()]
    if not cleaned:
        return []

    out = []
    banned_fragments = [
        "clubs de lecture",
        "perspective internationale",
        "enrichit la perspective",
        "analyse des récits",
        "forme physique et mentale",
        "culture générale",
        "vision du monde",
    ]

    for line in cleaned:
        low = line.lower().strip()

        if any(b in low for b in banned_fragments):
            continue

        line = clean_punctuation_text(line)
        line = re.sub(r"(?i), impliquant .*?$", ".", line)
        line = re.sub(r"(?i), avec une préférence marquée .*?$", ".", line)

        if line and ":" in line:
            head, tail = line.split(":", 1)
            head = head.strip()
            tail = tail.strip().rstrip(".")
            line = f"{head} : {tail}."
        elif line and " : " not in line:
            line = line.rstrip(".") + "."

        out.append(line)

    out = dedupe_preserve_order(out)

    if cv_is_short:
        return out[:4]
    
    if cv_is_long:
        return out[:4] if len(out) <= 4 else out[:3]
    
    if len(out) <= 4:
        return out
    
    return out[:4]

def trim_activities_droit(
    lines: list[str],
    ideal_max: int = 4,
    cv_is_short: bool = False,
    cv_is_long: bool = False,
) -> list[str]:
    cleaned = [(l or "").strip() for l in (lines or []) if (l or "").strip()]
    if not cleaned:
        return []

    out = []
    banned_fragments = [
        "clubs de lecture",
        "perspective internationale",
        "enrichit la perspective",
        "analyse des récits",
        "forme physique et mentale",
        "culture générale",
        "vision du monde",
    ]

    for line in cleaned:
        low = line.lower().strip()

        if any(b in low for b in banned_fragments):
            continue

        line = clean_punctuation_text(line)
        line = re.sub(r"(?i), impliquant .*?$", ".", line)
        line = re.sub(r"(?i), avec une préférence marquée .*?$", ".", line)
        low_after = line.lower()

        weak_legal_hobbies = ["musique", "cinéma", "cinema", "shopping"]
        if any(h in low_after for h in weak_legal_hobbies):
            if cv_is_long:
                continue

        if line and ":" in line:
            head, tail = line.split(":", 1)
            head = head.strip()
            tail = tail.strip().rstrip(".")
            line = f"{head} : {tail}."
        elif line and " : " not in line:
            line = line.rstrip(".") + "."

        out.append(line)

    out = dedupe_preserve_order(out)

    if cv_is_short:
        return out[:4]
    
    if cv_is_long:
        return out[:4] if len(out) <= 4 else out[:3]
    
    if len(out) <= 4:
        return out
    
    return out[:4]
        
def clean_skills_lines(lines: list[str]) -> list[str]:
    if not lines:
        return []

    banned_fragments = [
        "présentations percutantes",
        "compréhension avancée",
        "outils analytiques avancés",
        "résolution de problèmes complexes",
        "expertise avancée",
        "connaissance approfondie",
        "maîtrise approfondie",
        "excellente maîtrise",
        "approche orientée résultats",
        "communication interculturelle",
        "pensée critique",
        "leadership",
        "esprit stratégique",
        "sens stratégique",
        "avec une utilisation avancée",
        "utilisation avancée des fonctionnalités",
        "communication efficace",
        "coordination de projets",
        "outils de gestion de projet",
        "analyse financière",
        "immersion locale",
    ]

    cleaned = []
    seen = set()

    for raw in lines:
        txt = clean_punctuation_text((raw or "").strip())
        low = txt.lower()

        if not txt:
            continue

        if any(b in low for b in banned_fragments):
            continue

        key = low.strip()
        if key in seen:
            continue
        seen.add(key)

        cleaned.append(txt)

    return cleaned

def _find_paragraph_containing(doc: Document, needle: str):
    for p in doc.paragraphs:
        if needle in (p.text or ""):
            return p
    return None


def _clear_paragraph(p):
    p.text = ""


def _insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)

    if text:
        new_para.add_run(text)

    if style:
        try:
            new_para.style = style
        except Exception:
            pass

    return new_para


def _insert_lines_after(paragraph, lines, make_bullets=False):
    last = paragraph
    for line in lines:
        line = (line or "").rstrip()

        if not line:
            last = _insert_paragraph_after(last, "")
            continue

        if make_bullets and line.lstrip().startswith("-"):
            text = line.lstrip()[1:].strip()
            last = _insert_paragraph_after(last, text, style="List Bullet")
        else:
            last = _insert_paragraph_after(last, line)

    return last


def _split_sections(cv_text: str) -> dict:
    t = (cv_text or "").replace("\r\n", "\n").strip()

    tags = [
        "EDUCATION:", "FORMATION:",
        "EXPERIENCES:", "EXPÉRIENCES:", "EXPERIENCE:",
        "SKILLS:", "COMPETENCES:", "COMPÉTENCES:",
        "LANGUAGES:", "LANGUES:",
        "INTERESTS:", "ACTIVITIES:", "ACTIVITÉS:"
    ]
    pos = {tag: t.find(tag) for tag in tags}

    if all(pos[tag] == -1 for tag in tags):
        return {
            "EDUCATION": t.splitlines(),
            "EXPERIENCES": [],
            "SKILLS": [],
            "LANGUAGES": [],
            "INTERESTS": [],
            "ACTIVITIES": [],
        }

    present = [(tag, pos[tag]) for tag in tags if pos[tag] != -1]
    present.sort(key=lambda x: x[1])

    sections = {}
    for i, (tag, start) in enumerate(present):
        end = present[i + 1][1] if i + 1 < len(present) else len(t)
        block = t[start:end].strip().splitlines()

        if block and block[0].strip() == tag:
            block = block[1:]

        while block and not block[0].strip():
            block = block[1:]
        while block and not block[-1].strip():
            block = block[:-1]

        sections[tag.replace(":", "")] = block

    if not sections.get("SKILLS"):
        sections["SKILLS"] = sections.get("COMPETENCES") or sections.get("COMPÉTENCES") or []

    if not sections.get("LANGUAGES"):
        sections["LANGUAGES"] = sections.get("LANGUES") or []

    if not sections.get("EXPERIENCES"):
        sections["EXPERIENCES"] = sections.get("EXPÉRIENCES") or sections.get("EXPERIENCE") or []

    # 🔴 IMPORTANT : si le modèle écrit "FORMATION:" au lieu de "EDUCATION:"
    if not sections.get("EDUCATION"):
        sections["EDUCATION"] = sections.get("FORMATION") or sections.get("EDUCATION") or []

    return sections
def _render_education(anchor: Paragraph, lines: list[str]):
    """
    Rend la section FORMATION de façon un peu plus premium :
    - Première ligne de chaque bloc en gras
    - 'Cours pertinents' -> 'Matières fondamentales'
    - 'Matières fondamentales :' souligné
    - Dans la section EDUCATION, chaque diplôme ou programme est sur son propre paragraphe, séparé par UNE LIGNE VIDE du suivant (ex : Programme Grande École, ligne vide, puis Baccalauréat, etc.
    """
    last = anchor
    first_in_block = True

    for raw in (lines or []):
        line = (raw or "").strip()

        # ligne vide = séparation entre deux formations
        if not line:
            last = _insert_paragraph_after(last, "")
            first_in_block = True
            continue

        # Remplace le texte
        if "Cours pertinents" in line or "Key coursework" in line:
            line = line.replace("Cours pertinents", "Matières fondamentales")
            line = re.sub(r"(?i)key coursework", "Matières fondamentales", line)

        # Première ligne du bloc = nom d'école / programme -> gras
        if first_in_block:
            para = _insert_paragraph_after(last, "")
            run = para.add_run(line)
            run.bold = True
            para.paragraph_format.space_after = Pt(0)
            last = para
            first_in_block = False
            continue

        # Ligne "Matières fondamentales : ..." avec le label souligné
        if "Matières fondamentales" in line:
            para = _insert_paragraph_after(last, "")
            before, sep, after = line.partition(":")
            label = before + sep  # "Matières fondamentales:"
            r1 = para.add_run(label + " ")
            r1.underline = True
            if after:
                para.add_run(after.strip())
            last = para
            continue

        # Autres lignes normales
        last = _insert_paragraph_after(last, line)

    return last

def _render_interests(anchor: Paragraph, lines: list[str]):
    """
    Rend la section ACTIVITIES / CENTRES D'INTÉRÊT :
    - Chaque ligne -> puce
    - Nom de l'activité en gras avant ':' ou ' - '
    """
    last = anchor

    for raw in (lines or []):
        text = clean_punctuation_text((raw or "").strip())
        text = re.sub(r"(?i)^je\s+", "", text).strip()
        text = re.sub(r"(?i)^j['']\s*", "", text).strip()
        # supprime la première personne embarquée dans la phrase
        text = re.sub(r"(?i),?\s+ce qui m[''](a|ont)\s+permis\s+de\s+", ", permettant de ", text)
        text = re.sub(r"(?i),?\s+pour\s+(approfondir mes|rester informée|développer ma|développer mon)\s+", ", développant ", text)
        text = re.sub(r"(?i)\bmes\s+(connaissances|compétences)\b", "les compétences", text)
        text = re.sub(r"(?i)\bmon\s+esprit\b", "l'esprit", text)
        text = re.sub(r"(?i)\bma\s+capacité\b", "la capacité", text)
        if not text:
            last = _insert_paragraph_after(last, "")
            continue

        # Nouveau paragraphe en mode liste à puces
        new_p = _insert_paragraph_after(last, "")

        head = text
        tail = ""

        if ":" in text:
            head, tail = text.split(":", 1)
        elif " - " in text:
            left, right = text.split(" - ", 1)
            # On considère que la partie gauche est le "nom" si elle est courte
            if len(left.split()) <= 4:
                head, tail = left, right
            else:
                head, tail = text, ""

        head = head.strip()
        tail = tail.strip()

        # Nettoyage des éventuels **...** ou *...* venant du modèle
        while head.startswith("*") and head.endswith("*") and len(head) > 2:
            head = head[1:-1].strip()

        r_head = new_p.add_run(head)
        r_head.bold = True

        if tail:
            new_p.add_run(" : " + tail)

        last = new_p

    return last

def dedupe_language_items(items: list[str]) -> list[str]:
    if not items:
        return []

    normalized = []
    seen_exact = set()
    seen_language_bases = set()
    seen_test_keywords = set()

    language_roots = [
        "anglais", "français", "francais", "espagnol", "allemand",
        "italien", "chinois", "mandarin", "cantonais", "japonais",
        "coréen", "coreen", "arabe", "portugais", "russe"
    ]

    test_keywords = ["toeic", "toefl", "ielts", "cambridge"]

    def extract_language_root(txt: str) -> str | None:
        """Extrait la racine de langue dans un texte, même verbeux comme 'compétences en Espagnol'."""
        low = txt.lower()
        # Chercher la racine n'importe où dans le texte
        for root in language_roots:
            if re.search(rf"\b{root}\b", low):
                return root
        return None

    for raw in items:
        txt = clean_punctuation_text((raw or "").strip())
        if not txt:
            continue

        low = txt.lower()
        low = low.replace("niveau ", "")
        low = re.sub(r"\s+", " ", low).strip()

        # Supprimer les entrées verbeuses qui répètent une langue déjà vue
        # ex: "compétences en Espagnol de niveau A2" quand "Espagnol A2" est déjà là
        root_found = extract_language_root(txt)
        if root_found and root_found in seen_language_bases:
            continue  # langue déjà présente → doublon verbeux

        # si c'est un test officiel, on le garde tel quel une seule fois
        if any(k in low for k in test_keywords):
            matched_test = next((k for k in test_keywords if k in low), None)
            if matched_test and matched_test in seen_test_keywords:
                continue
            if low not in seen_exact:
                seen_exact.add(low)
                if matched_test:
                    seen_test_keywords.add(matched_test)
                normalized.append(txt)
                if root_found:
                    seen_language_bases.add(root_found)
            continue

        matched_root = None
        for root in language_roots:
            if low.startswith(root):
                matched_root = root
                break

        if matched_root:
            if matched_root in seen_language_bases:
                continue
            seen_language_bases.add(matched_root)
            normalized.append(txt)
            continue

        if low not in seen_exact:
            seen_exact.add(low)
            normalized.append(txt)

    return normalized

def validate_skills_completeness(skills_line: str, payload: dict) -> str:
    """
    Vérifie que les skills critiques du payload ne sont pas absents du CV généré.
    Si Excel/PowerPoint/outils clés sont dans le payload mais absents du CV → les réinjecter.
    """
    if not skills_line:
        return skills_line

    raw_skills = (payload.get("skills") or "").lower()
    skills_lower = skills_line.lower()

    # Outils critiques qu'on ne peut jamais perdre
    CRITICAL_TOOLS = {
        "excel": "Excel",
        "powerpoint": "PowerPoint",
        "word": "Word",
    }

    missing = []
    for key, label in CRITICAL_TOOLS.items():
        if key in raw_skills and key not in skills_lower:
            missing.append(label)

    if missing:
        # Insérer les outils manquants après "Maîtrise des logiciels :"
        if "Maîtrise des logiciels :" in skills_line:
            skills_line = skills_line.replace(
                "Maîtrise des logiciels :",
                "Maîtrise des logiciels : " + ", ".join(missing) + ","
            )
        else:
            skills_line = "Maîtrise des logiciels : " + ", ".join(missing) + ", " + skills_line

    return skills_line


def build_software_line_from_payload(payload: dict) -> str:
    raw_skills = payload.get("skills") or ""
    items = [clean_punctuation_text(x.strip()) for x in re.split(r",|;", raw_skills) if x.strip()]
    items = dedupe_preserve_order(items)

    if not items:
        items = ["Pack Office"]

    return "Maîtrise des logiciels : " + ", ".join(items)

def _clean_user_annotation(text: str) -> str:
    """
    Supprime les annotations personnelles que les utilisateurs ajoutent dans les champs
    (ex: "je prépare aussi", "pas encore obtenu", "je me débrouille", "pas super fort", etc.)
    """
    if not text:
        return text
    # Patterns courants d'annotations personnelles
    patterns = [
        r"\bje (prépare|prepare|fais|suis|connais|parle|maîtrise|maitrise|me débrouille|me debrouille)[^,;.]*",
        r"\bpas (encore|super|très|tres|trop|forcément|forcement)[^,;.]*",
        r"\ben (cours|préparation|preparation)[^,;.]*",
        r"\baussi\b[^,;.]*",
        r"\bunpeu\b[^,;.]*",
        r"\bun peu\b[^,;.]*",
        r"\bnotions?\b(?!\s+de\s+\w)",  # "notions" seul = gardé, "notions de X" = gardé
        r"\bà améliorer\b[^,;.]*",
        r"\bà perfectionner\b[^,;.]*",
        r"\bmais (je|j')[^,;.]*",
    ]
    for pat in patterns:
        text = re.sub(pat, "", text, flags=re.IGNORECASE)
    # Nettoyer ponctuation résiduelle
    text = re.sub(r"\s*,\s*,", ",", text)
    text = re.sub(r",\s*$", "", text.strip())
    return text.strip()


def normalize_skills_block(lines: list, payload: dict) -> list:
    raw = " ".join((x or "").strip() for x in (lines or []) if (x or "").strip())
    raw = re.sub(r"\s+", " ", raw).strip()

    raw = re.sub(r"(?i)\bcertifications\s*:", "Certifications :", raw)
    raw = re.sub(r"(?i)\bma[iî]trise des logiciels\s*:", "Maîtrise des logiciels :", raw)
    raw = re.sub(r"(?i)\bcapacités professionnelles\s*:", "Capacités professionnelles :", raw)
    raw = re.sub(r"(?i)\bcapacites professionnelles\s*:", "Capacités professionnelles :", raw)
    raw = re.sub(r"(?i)\blangues\s*:", "Langues :", raw)

    labels = [
        "Certifications :",
        "Maîtrise des logiciels :",
        "Capacités professionnelles :",
        "Langues :",
    ]

    chunks = []
    current = raw

    while current:
        next_positions = []
        for label in labels:
            pos = current.find(label)
            if pos != -1:
                next_positions.append((pos, label))

        if not next_positions:
            if current.strip():
                chunks.append(current.strip())
            break

        next_positions.sort(key=lambda x: x[0])
        first_pos, _ = next_positions[0]

        if first_pos > 0:
            orphan = current[:first_pos].strip(" ,")
            if orphan:
                chunks.append(orphan)

        current = current[first_pos:]

        next_positions = []
        for label in labels:
            pos = current.find(label)
            if pos != -1:
                next_positions.append((pos, label))
        next_positions.sort(key=lambda x: x[0])

        if len(next_positions) >= 2:
            chunk = current[next_positions[0][0]:next_positions[1][0]].strip()
            current = current[next_positions[1][0]:]
        else:
            chunk = current.strip()
            current = ""

        if chunk:
            chunks.append(chunk)

    payload_certifications = [_clean_user_annotation(x.strip()) for x in re.split(r",|;", payload.get("certifications", "") or "") if x.strip() and len(_clean_user_annotation(x.strip())) > 2]
    payload_languages = clean_punctuation_text(_clean_user_annotation((payload.get("languages") or "").strip()))
    payload_skills = clean_punctuation_text(_clean_user_annotation((payload.get("skills") or "").strip()))

    cleaned = []
    seen = set()

    language_tests = []
    certifications_items = []

    for chunk in chunks:
        chunk = clean_punctuation_text(chunk)
        if not chunk:
            continue

        low = chunk.lower()

        if low.startswith("certifications :"):
            content = chunk.split(":", 1)[1].strip() if ":" in chunk else ""
            for item in [x.strip() for x in content.split(",") if x.strip()]:
                item_low = item.lower()
                if any(k in item_low for k in ["toeic", "toefl", "ielts", "cambridge"]):
                    language_tests.append(item)
                else:
                    certifications_items.append(item)
            continue

        if low.startswith("langues :"):
            content = chunk.split(":", 1)[1].strip() if ":" in chunk else ""
            if content:
                # ✅ Splitter par virgule HORS des parenthèses pour gérer "Anglais (courant, TOEIC 910)"
                parts = []
                depth = 0
                current_part = ""
                for ch in content:
                    if ch == "(":
                        depth += 1
                        current_part += ch
                    elif ch == ")":
                        depth -= 1
                        current_part += ch
                    elif ch == "," and depth == 0:
                        if current_part.strip():
                            parts.append(current_part.strip())
                        current_part = ""
                    else:
                        current_part += ch
                if current_part.strip():
                    parts.append(current_part.strip())

                for p in parts:
                    if p not in language_tests:
                        language_tests.append(p)
            continue

        key = low
        if key not in seen:
            seen.add(key)
            cleaned.append(chunk)

    for item in payload_certifications:
        item_low = item.lower()
        if any(k in item_low for k in ["toeic", "toefl", "ielts", "cambridge"]):
            if item not in language_tests:
                language_tests.append(item)
        else:
            if item not in certifications_items:
                certifications_items.append(item)

    # ✅ Utiliser la ligne logiciels du LLM si elle existe (bien formatée), sinon fallback payload
    has_llm_software_line = any(x.lower().startswith("maîtrise des logiciels") for x in cleaned)
    if not has_llm_software_line:
        cleaned.insert(0, build_software_line_from_payload(payload))

    certifications_items = dedupe_preserve_order(certifications_items)

    if certifications_items:
        cert_line = "Certifications : " + ", ".join(certifications_items)
        if not any(x.lower().startswith("certifications :") for x in cleaned):
            cleaned.insert(0, cert_line)

    def _clean_lang_item(item: str) -> str:
        item = re.sub(r",?\s+avec\s+capacité.*$", "", item, flags=re.IGNORECASE)
        item = re.sub(r",?\s+permettant\s+de.*$", "", item, flags=re.IGNORECASE)
        item = re.sub(r",?\s+et\s+(une\s+)?compréhension.*$", "", item, flags=re.IGNORECASE)
        return item.strip()

    # ✅ Re-parser toutes les langues via parse_languages_smart pour gérer le texte brut du LLM
    # Ex: "francais natif anglais courant TOEFL 105 italien B2" → liste propre
    reparsed = []
    for item in language_tests:
        if "," in item or len(item.split()) <= 4:
            reparsed.append(item)
        else:
            # Texte brut long → re-parser
            parsed = parse_languages_smart(item)
            reparsed.extend(parsed if parsed else [item])
    language_tests = reparsed

    # ✅ Payload languages uniquement en fallback — si le LLM n'a fourni aucune langue
    if payload_languages and not language_tests:
        language_tests = parse_languages_smart(payload_languages)

    language_tests = dedupe_language_items(language_tests)

    # Français toujours en premier
    french_items = [x for x in language_tests if re.search(r"(?i)\bfran[cç]ais\b", x)]
    other_items = [x for x in language_tests if not re.search(r"(?i)\bfran[cç]ais\b", x)]

    # Trier les autres langues par niveau décroissant
    def lang_level_score(lang: str) -> int:
        l = lang.lower()
        if "natif" in l or "native" in l: return 10
        if "c2" in l: return 9
        if "c1" in l: return 8
        if "courant" in l or "fluent" in l: return 7
        if "b2" in l: return 6
        if "b1" in l or "intermédiaire" in l: return 5
        if "a2" in l: return 3
        if "a1" in l or "notions" in l or "débutant" in l: return 2
        return 4  # score par défaut si pas de niveau

    other_items = sorted(other_items, key=lang_level_score, reverse=True)
    language_tests = french_items + other_items

    if language_tests:
        lang_line = "Langues : " + ", ".join(language_tests)
    else:
        lang_line = "Langues : Français"

    cleaned = [x for x in cleaned if not x.lower().startswith("langues :")]
    cleaned.append(lang_line)

    final = []
    seen_final = set()
    for line in cleaned:
        line = clean_punctuation_text(line)
        if not line:
            continue
        key = line.lower()
        if key in seen_final:
            continue
        seen_final.add(key)
        final.append(line)

    return final

def _render_skills(anchor: Paragraph, lines: list[str], payload: dict = None):
    """
    Rend la section COMPÉTENCES & OUTILS :
    - Pas de puces
    - Sous-titres en gras (Certifications, Maîtrise des logiciels, Capacités professionnelles)
    - Les éléments sont séparés par des virgules
    """
    last = anchor
    is_first = True  # ✅ pour ajouter un petit espace avant la 1ère ligne
    cleaned = []

    # Si l'utilisateur n'a fourni aucune certification, on supprime toute ligne Certifications
    user_certifications = (payload or {}).get("certifications", "").strip() if payload else ""

    for line in lines:
        txt = line.strip()
        
        if txt.lower().startswith("certifications"):
            # Si l'utilisateur n'a rien fourni → on supprime
            if not user_certifications:
                continue
            # Si l'utilisateur a fourni des certifications → on affiche toujours
            # (sans filtre keywords car l'utilisateur sait ce qu'il a fourni)
    
        cleaned.append(txt)

    normalized = []
    labels = [
        "Certifications :",
        "Maîtrise des logiciels :",
        "Capacités professionnelles :",
        "Langues :",
    ]

    for txt in cleaned:
        current = txt.strip()
        split_done = True

        while split_done:
            split_done = False
            for label in labels:
                pos = current.find(label)
                if pos > 0:
                    left = current[:pos].strip().rstrip(",")
                    right = current[pos:].strip()
                    if left:
                        normalized.append(left)
                    current = right
                    split_done = True
                    break

        if current:
            normalized.append(current)
    
    cleaned = normalized

    for raw in (cleaned or []):
        text = clean_punctuation_text((raw or "").strip())
        if not text:
            last = _insert_paragraph_after(last, "")
            continue

        # On remplace les ' | ' par des virgules si jamais le modèle en met encore
        text = text.replace(" | ", ", ")

        new_p = _insert_paragraph_after(last, "")

        # ✅ petit espace juste au début de la section
        if is_first:
            is_first = False
        
        head = text
        tail = ""

        # ✅ normalisation des libellés (le LLM varie souvent)
        hlow = head.lower()
        if hlow in {"capacités", "capacites"}:
            head = "Capacités professionnelles"
        if hlow in {"logiciels"}:
            head = "Maîtrise des logiciels"

        if ":" in text:
            head, tail = text.split(":", 1)
        elif " - " in text:
            left, right = text.split(" - ", 1)
            if len(left.split()) <= 4:
                head, tail = left, right
            else:
                head, tail = text, ""

        head = head.strip()
        tail = tail.strip()

        r_head = new_p.add_run(head)
        r_head.bold = True

        if tail:
            new_p.add_run(" : " + tail)

        last = new_p

    return last
    
def _education_end_year(block: list[str]) -> int:
    """
    Récupère l'année de fin à partir de la première ligne du bloc.
    On prend simplement la DERNIÈRE année à 4 chiffres trouvée dans la ligne.
    Ex :
      'Programme Grande École – ESCP — Sept 2022 – Juin 2026' -> 2026
      'Classe préparatoire ECG – Lycée du Parc (2020-2022)'   -> 2022
    """

    if not block:
        return 0

    first_line = (block[0] or "").strip()

    # On cherche toutes les années à 4 chiffres dans la ligne complète
    years = re.findall(r"(?:19|20)\d{2}", first_line)

    if not years:
        return 0

    try:
        # Dernière année = année de fin
        return int(years[-1])
    except ValueError:
        return 0

def _is_bac_block(block: list[str]) -> bool:
    """Retourne True si le bloc correspond à un baccalauréat classique."""
    if not block:
        return False
    first = (block[0] or "").lower()
    return "baccalauréat" in first or "baccalaureat" in first


def _keep_bac_block(block: list[str]) -> bool:
    """
    On garde le bac si :
    1) lycée d'exception (Henri IV, Louis-le-Grand, lycée international, etc.)
    2) bac / diplôme international (IB, Abibac, maturité suisse, etc.)
    3) mention d'honneur / félicitations du jury
    4) moyenne ou note explicitement fournie (ex : 15.5, mention bien)
    """
    text = " ".join(block).lower()
    # Cas spécifiques : honneurs / honeurs du jury
    if "honneurs du jury" in text or "honeurs du jury" in text:
        return True

    # ✅ Garde si une moyenne ou note est mentionnée (l'utilisateur a mis sa note → valeur pour lui)
    if re.search(r"\bmoyenne\b", text):
        return True
    if re.search(r"\b(mention\s+(bien|très bien|assez bien|très\s+bien))\b", text):
        return True
    # Note brute type 15, 16, 17...
    if re.search(r"\b(1[4-9]|20)[\.,]\d", text):
        return True

    elite_keywords = [
        "henri iv", "henri-iv", "henry iv",
        "louis-le-grand", "louis le grand",
        "lycée international", "lycee international",
        "lycée du parc", "lycee du parc",
        "stanislas", "lycée stanislas",
        "janson de sailly",
        "franklin", "lycée franklin",
        "fénelon", "fenelon",
        "charlemagne",
        "buffon",
        "condorcet",
        "sainte-geneviève", "sainte genevieve", "ginette",
        "le parc",
        "masséna", "massena",
        "thiers",
        "hoche",
        "kléber", "kleber",
        "clemenceau",
        "du parc",
        "chateaubriand",
        "berthelot",
        "pierre de fermat",
        "montaigne",
        "descartes",
        "champollion",
    ]

    intl_keywords = [
        "baccalauréat international", "baccalaureat international",
        "international baccalaureate", "ib diploma", "ib programme",
        "abibac", "esabac",
        "maturité suisse", "maturite suisse", "maturité gymnasiale",
        "matura",
        " ib ",
        "cess",  # Belgique
        "certificat d'enseignement secondaire supérieur",
        "certificat d'enseignement secondaire superieur",
    ]

    honours_keywords = [
        "félicitations du jury", "felicitations du jury", "honneurs du jury"
    ]

    if any(k in text for k in elite_keywords):
        return True
    if any(k in text for k in intl_keywords):
        return True
    if any(k in text for k in honours_keywords):
        return True

    return False

def parse_languages_smart(text: str) -> list:
    """
    Parse une chaîne de langues même sans virgules.
    "francais natif anglais courant IELTS 8 allemand intermediaire B1"
    → ["Français natif", "Anglais courant (IELTS 8)", "Allemand intermédiaire (B1)"]
    """
    if not text:
        return []
    if "," in text:
        return [x.strip() for x in text.split(",") if x.strip()]

    lang_map = {
        "francais": "Français", "français": "Français",
        "anglais": "Anglais", "english": "Anglais",
        "allemand": "Allemand", "german": "Allemand",
        "espagnol": "Espagnol", "spanish": "Espagnol",
        "italien": "Italien", "italian": "Italien",
        "portugais": "Portugais", "portuguese": "Portugais",
        "chinois": "Chinois", "chinese": "Chinois",
        "japonais": "Japonais", "japanese": "Japonais",
        "arabe": "Arabe", "arabic": "Arabe",
        "russe": "Russe", "russian": "Russe",
        "neerlandais": "Néerlandais", "néerlandais": "Néerlandais", "dutch": "Néerlandais",
        "coreen": "Coréen", "coréen": "Coréen", "korean": "Coréen",
        "turc": "Turc", "turkish": "Turc",
    }

    text_low = text.lower()
    positions = []
    for key, canonical in lang_map.items():
        for m in re.finditer(r"\b" + re.escape(key) + r"\b", text_low):
            positions.append((m.start(), m.end(), canonical, key))

    if not positions:
        return [text.strip()] if text.strip() else []

    positions.sort(key=lambda x: x[0])
    seen_canonical = set()
    unique_positions = []
    for pos in positions:
        if pos[2] not in seen_canonical:
            seen_canonical.add(pos[2])
            unique_positions.append(pos)
    positions = unique_positions

    segments = []
    for i, (start, end, canonical, key) in enumerate(positions):
        next_start = positions[i + 1][0] if i + 1 < len(positions) else len(text)
        suffix = text[end:next_start].strip()
        suffix = re.sub(r"^[:\-–,\s]+", "", suffix)

        test_match = re.search(
            r"(TOEIC|TOEFL|IELTS|DELF|DALF|Cambridge|HSK)\s*[\:\s]?\s*(\d+[\.,]?\d*)",
            suffix, re.IGNORECASE
        )
        if test_match:
            score_text = test_match.group(0).strip()
            level_text = re.sub(
                r"(TOEIC|TOEFL|IELTS|DELF|DALF|Cambridge|HSK)\s*[\:\s]?\s*(\d+[\.,]?\d*)",
                "", suffix, flags=re.IGNORECASE
            ).strip().strip(" ,;-–")
            full = f"{canonical} {level_text} ({score_text})".strip() if level_text else f"{canonical} ({score_text})"
        else:
            full = f"{canonical} {suffix}".strip() if suffix else canonical
            # Normaliser niveaux CEFR isolés
            full = re.sub(r"\s+(b1|b2|c1|c2|a1|a2)\b",
                          lambda m: f" ({m.group(1).upper()})", full, flags=re.IGNORECASE)

        segments.append(full)

    return segments


def normalize_contract_type(t: str) -> str:
    if not t:
        return ""

    original = t.strip()
    t_clean = original.lower().strip()

    base_mapping = {
        "internship": "Stage",
        "intern": "Stage",
        "traineeship": "Stage",
        "apprenticeship": "Alternance",
        "full-time": "CDI",
        "full time": "CDI",
        "part-time": "Temps partiel",
        "part time": "Temps partiel",
        "part-time job": "Job étudiant",
        "student job": "Job étudiant",
        "summer job": "Job d'été",
        "temporary": "CDD",
        "contract": "CDD",
        "volunteering": "Volontariat",
        "volunteer": "Volontariat",
    }

    # Match exact
    if t_clean in base_mapping:
        return base_mapping[t_clean]

    # Match préfixe (ex: "part-time job - barista")
    for key, value in base_mapping.items():
        if t_clean.startswith(key + " "):
            suffix = original[len(key):].lstrip(" -–—")
            return value + (f" – {suffix}" if suffix else "")

    return original

def extract_start_year_month(dates_str: str) -> int:
    """
    Extrait l'année et le mois de début depuis une chaîne de dates.
    Retourne un entier YYYYMM pour permettre le tri chronologique décroissant.
    Ex: "Sept 2025 – Aujourd'hui" → 202509
        "Janv 2024 – Août 2024"  → 202401
    """
    if not dates_str:
        return 0

    month_map = {
        "janv": 1, "jan": 1, "janvier": 1, "january": 1,
        "fév": 2, "fev": 2, "février": 2, "fevrier": 2, "february": 2, "feb": 2,
        "mars": 3, "mar": 3, "march": 3,
        "avr": 4, "apr": 4, "avril": 4, "april": 4,
        "mai": 5, "may": 5,
        "juin": 6, "jun": 6, "june": 6,
        "juil": 7, "jul": 7, "juillet": 7, "july": 7,
        "août": 8, "aout": 8, "aug": 8, "august": 8,
        "sept": 9, "sep": 9, "septembre": 9, "september": 9,
        "oct": 10, "octobre": 10, "october": 10,
        "nov": 11, "novembre": 11, "november": 11,
        "déc": 12, "dec": 12, "décembre": 12, "decembre": 12, "december": 12,
    }

    # Prendre seulement la partie avant le premier tiret (date de début)
    start_part = re.split(r"\s*[–\-]\s*", dates_str)[0].strip()

    # Chercher mois + année (ex: "Sept 2025")
    m = re.search(r"(\b\w+\b)\s+(\d{4})", start_part, re.IGNORECASE)
    if m:
        month_str = m.group(1).lower()
        year = int(m.group(2))
        month = month_map.get(month_str, 0)
        return year * 100 + month

    # Fallback: juste une année (ex: "2024")
    m = re.search(r"(\d{4})", start_part)
    if m:
        return int(m.group(1)) * 100

    return 0


def is_student_job_exp(exp: dict) -> bool:
    """
    Retourne True si l'expérience est un job étudiant / alimentaire,
    pas directement lié à une trajectoire professionnelle.
    """
    type_ = (exp.get("type") or "").lower()
    role = (exp.get("role") or "").lower()
    company = (exp.get("company") or "").lower()
    all_text = f"{type_} {role} {company}"

    # Types de contrat explicitement étudiants
    # ✅ Jamais classer BDE / Junior Entreprise / associatif comme job étudiant
    ASSOCIATIF_KEYWORDS = [
        "bde", "bda", "bds", "junior entreprise", "junior-entreprise",
        "association", "asso", "club", "comité", "bureau", "cde", "vie étudiante",
    ]
    if any(kw in role or kw in company or kw in type_ for kw in ASSOCIATIF_KEYWORDS):
        return False

    student_types = [
        "job étudiant", "job etudiant", "job d'été", "job d'ete",
        "temps partiel", "part-time", "part time", "summer job",
        "saisonnier", "saisonnière", "saisonniere",
    ]
    if any(t in type_ for t in student_types):
        return True

    # Rôles typiquement alimentaires / étudiants
    student_roles = [
        "caissier", "caissière", "cassier", "cassière",
        "serveur", "serveuse",
        "barista", "barman", "barmaid",
        "vendeur", "vendeuse",
        "livreur", "livreuse",
        "hôte de caisse", "hôtesse de caisse",
        "hôte d'accueil", "hôtesse d'accueil",
        "animateur", "animatrice",
        "surveillant de baignade",
        "manutentionnaire",
        "baby-sitter", "babysitter",
        "employé polyvalent", "employe polyvalent",
        "employé de rayon", "employe de rayon",
        "agent d'accueil",
    ]
    if any(r in role for r in student_roles):
        return True

    return False


def _split_education_block_on_degree_titles(block: list[str]) -> list[list[str]]:
    """
    Si l'IA enchaîne plusieurs diplômes dans un même bloc (sans ligne vide),
    on découpe dès qu'une ligne commence par un mot typique de diplôme.
    Exemple :
      Master 2 Finance ...
      Licence Finance ...
      Baccalauréat ES ...

    devient 3 blocs distincts.
    """
    if not block:
        return []

    DEGREE_STARTERS = (
        "Master", "Master 1", "Master 2",
        "Programme Grande École", "Programme Grande Ecole", "Programme",
        "Licence", "License",
        "Baccalauréat", "Baccalaureat",
        "Classe préparatoire", "Classe préparatoire ECG",
        "Classe preparatoire", "Classe preparatoire ECG",
        "CPGE", "Prépa", "Prepa",
        "Échange académique", "Echange académique", "Exchange programme", "Exchange program",
        "BBA", "Bachelor"
    )

    new_blocks: list[list[str]] = []
    current: list[str] = []

    for raw in block:
        line = (raw or "").strip()
        if not line:
            # Lignes vides : on ferme le bloc en cours
            if current:
                new_blocks.append(current)
                current = []
            continue

        # Si on tombe sur une nouvelle ligne qui ressemble à un début de diplôme
        # on démarre un nouveau bloc
        if current and any(line.startswith(prefix) for prefix in DEGREE_STARTERS):
            new_blocks.append(current)
            current = [line]
        else:
            current.append(line)

    if current:
        new_blocks.append(current)

    return new_blocks

def collapse_blank_paragraphs(doc: Document, max_consecutive: int = 1):
    """
    Supprime les paragraphes vraiment vides,
    MAIS conserve ceux qui servent d'espacement (space_before/space_after > 0).
    """
    blanks = 0

    for p in list(doc.paragraphs):
        txt = (p.text or "")
        fmt = p.paragraph_format

        has_spacing = bool(
            (fmt.space_before and fmt.space_before.pt > 0) or
            (fmt.space_after and fmt.space_after.pt > 0)
        )

        is_blank_text = (txt.strip() == "")
        is_blank = is_blank_text and not has_spacing

        if is_blank:
            blanks += 1
            if blanks > max_consecutive:
                _remove_paragraph(p)
        else:
            blanks = 0

def normalize_section_titles_spacing(doc: Document, section_space: Pt, title_space_after: Pt):
    TITLES = {
        "FORMATION",
        "EXPÉRIENCES PROFESSIONNELLES",
        "COMPÉTENCES & OUTILS",
        "ACTIVITÉS & CENTRES D’INTÉRÊT",
        "ACTIVITÉS & CENTRES D'INTÉRÊT",
    }
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.upper() in TITLES:
            p.paragraph_format.space_before = section_space
            p.paragraph_format.space_after = title_space_after     

def _strip_blank_neighbors(doc: Document, p: Paragraph, before: int = 1, after: int = 1):
    """
    Supprime les paragraphes vides juste avant/après un paragraphe (souvent présents dans le template).
    Permet d'éviter le "double espace" (template + code).
    """
    paras = list(doc.paragraphs)

    idx = None
    for i, pp in enumerate(paras):
        if pp is p:
            idx = i
            break
    if idx is None:
        return

    # Remove blanks BEFORE
    for _ in range(before):
        j = idx - 1
        if j >= 0 and (paras[j].text or "").strip() == "":
            _remove_paragraph(paras[j])
            paras = list(doc.paragraphs)
            idx -= 1  # index shift

    # Remove blanks AFTER (remove up to `after` blank paras)
    removed = 0
    while removed < after:
        paras = list(doc.paragraphs)
        if idx + 1 < len(paras) and (paras[idx + 1].text or "").strip() == "":
            _remove_paragraph(paras[idx + 1])
            removed += 1
        else:
            break
            
def write_docx_from_template(template_path: str, cv_text: str, out_path: str, payload: dict = None, compact_mode: bool = False) -> None:
    doc = Document(template_path)
    # spacing appliqué plus bas selon le secteur

    # On mesure la longueur du texte pour savoir si on doit "tailler" ou pas.
    raw_text = cv_text or ""
    nb_lines = raw_text.count("\n") + 1  # nombre de lignes brutes

    # Longueur SANS espaces (celle que tu mesures dans Word)
    chars_no_space = len(re.sub(r"\s+", "", raw_text))

    # Au-delà d’environ 2225 caractères sans espaces → CV considéré comme "long"
    cv_is_long = (chars_no_space > 2225) or (nb_lines > 85)
    cv_is_short = (chars_no_space < 1150) or (nb_lines < 42)

    # Marges plus petites pour mieux utiliser la largeur
    for section in doc.sections:
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
        section.top_margin = Cm(1.0)      
        section.bottom_margin = Cm(1.0)   

    # ✅ Mode compact : on compresse légèrement la mise en page si ça dépasse 1 page
    if compact_mode:
        for p in doc.paragraphs:
            try:
                # réduire les espaces verticaux
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
    
                # réduire l'interligne (très léger)
                p.paragraph_format.line_spacing = 1.0
    
                # réduire la taille de police (un peu)
                for run in p.runs:
                    if run.font.size is None:
                        run.font.size = Pt(9.5)
                    else:
                        # ne pas toucher au nom géant (20pt), on limite juste
                        if run.font.size.pt > 11:
                            continue
                        run.font.size = Pt(min(run.font.size.pt, 9.5))
            except Exception:
                pass

    # ------- Données générales -------
    payload = payload or {}
    is_legal = is_legal_sector(payload.get("sector", ""))
    is_audit = is_audit_sector(payload.get("sector", ""))
    is_finance = is_finance_sector(payload.get("sector", ""))
    if is_finance or is_audit or is_management_sector(payload.get("sector", "")):
        normalize_section_titles_spacing(doc, SECTION_SPACING, ITEM_SPACING)
    else:
        normalize_section_titles_spacing(doc, Pt(0), Pt(0))
    full_name = payload.get("full_name", "").strip() or "NOM Prénom"
    role = payload.get("role", "").strip()
    finance_type = payload.get("finance_type", "").strip()
    cv_title = finance_type if finance_type else role

    contact_line = " | ".join([
        x.strip()
        for x in [
            payload.get("phone", ""),
            normalize_email(payload.get("email", "") or ""),
            payload.get("linkedin", ""),
        ]
        if x and x.strip()
    ])

    sections = _split_sections(cv_text)

    # On garde en priorité les SKILLS générés par le LLM
    llm_skills = sections.get("SKILLS") or []
    llm_languages = sections.get("LANGUAGES") or []

    if not llm_skills:
        llm_skills = []

    # fallback minimal si le LLM n'a rien mis
    if not llm_skills:
        raw_certifications = (payload.get("certifications") or "").strip()
        raw_skills = (payload.get("skills") or "").strip()
        raw_languages = (payload.get("languages") or "").strip()

        if raw_certifications:
            llm_skills.append(f"Certifications : {raw_certifications}")
        if raw_skills:
            llm_skills.append(f"Maîtrise des logiciels : {raw_skills}")
        if raw_languages:
            llm_skills.append(f"Langues : {raw_languages}")

    # si le LLM a mis les langues en section séparée, on les réintègre
    if llm_languages:
        lang_text = ", ".join(x.strip() for x in llm_languages if x.strip())
        has_languages_line = any(
            (line or "").strip().lower().startswith("langues")
            for line in llm_skills
        )
        if lang_text and not has_languages_line:
            llm_skills.append(f"Langues : {lang_text}")

    sections["SKILLS"] = normalize_skills_block(llm_skills, payload)
    # ✅ Vérifier qu'Excel/PowerPoint ne manquent pas malgré la normalisation
    sections["SKILLS"] = [
        validate_skills_completeness(line, payload) if "logiciels" in (line or "").lower() else line
        for line in sections["SKILLS"]
    ]
    sections["LANGUAGES"] = []

    if not sections.get("SKILLS"):
        fallback_skills = []

        raw_certifications = (payload.get("certifications") or "").strip()
        raw_skills = (payload.get("skills") or "").strip()
        raw_languages = (payload.get("languages") or "").strip()

        if raw_certifications:
            fallback_skills.append(f"Certifications : {raw_certifications}")

        if raw_skills:
            fallback_skills.append(f"Maîtrise des logiciels : {raw_skills}")

        if raw_languages:
            fallback_skills.append(f"Langues : {raw_languages}")

        sections["SKILLS"] = fallback_skills

    # SKILLS : on garde plusieurs lignes, mais on filtre ce qui n'est pas dans l'input user
    if isinstance(sections.get("SKILLS"), list):
        raw_skills_input = (
            (payload.get("skills") or "") + " " +
            (payload.get("languages") or "")
        ).lower()
    
        cleaned = []
        for x in sections["SKILLS"]:
            txt = x.strip().lstrip("-").strip()
            low = txt.lower()
    
            # garde toujours les libellés
            if low.startswith("maîtrise des logiciels") or low.startswith("capacités professionnelles") or low.startswith("certifications") or low.startswith("langues"):
                # filtre les ajouts trop "magiques"
                banned = [
                    "logiciels de gestion financière",
                    "data visualisation",
                    "expertise avancée",
                    "connaissance approfondie",
                    "présentation claire et convaincante",
                    "outils analytiques avancés",
                    "logiciels de reporting",
                    "maîtrise approfondie",
                    "expertise en",
                    "solide maîtrise des outils",
                    "compétences avancées en",
                    "visualisation de données",
                    "gestion financière avancée",
                    "capacités analytiques",
                    "facilitant la communication interculturelle",
                    "compétences numériques en gestion documentaire",
                    "pensée critique",
                    "communication interculturelle",
                ]
                if any(b in low for b in banned):
                    continue
                cleaned.append(txt)
    
        sections["SKILLS"] = cleaned
        sections["SKILLS"] = clean_skills_lines(sections["SKILLS"])
    
    # ⬇️ Langues intégrées dans Compétences & Outils
    languages_raw = sections.get("LANGUAGES") or []
    if isinstance(languages_raw, list):
        lang_text = ", ".join(x.strip() for x in languages_raw if x.strip())
    else:
        lang_text = str(languages_raw).strip()
    
    if lang_text:
        skills_list = sections.get("SKILLS") or []
        has_languages_line = any(
            (line or "").strip().lower().startswith("langues")
            for line in skills_list
        )
        if not has_languages_line:
            skills_list.append(f"Langues : {lang_text}")
        sections["SKILLS"] = skills_list
    
    sections["LANGUAGES"] = []
    
    generated_activities = sections.get("ACTIVITIES") or []
    payload_activities = [line.strip() for line in (payload.get("interests") or "").splitlines() if line.strip()]
    
    # priorité à la sortie générée si elle existe, sinon fallback input utilisateur
    if generated_activities and isinstance(generated_activities, list):
        interests_source = [x.strip() for x in generated_activities if x and x.strip()]
    else:
        interests_source = payload_activities
    
    # on stabilise les activités : on garde la version générée ou utilisateur sans réécriture LLM
    if interests_source:
        interests_rewritten = interests_source
    else:
        interests_rewritten = []
    
    if isinstance(interests_rewritten, list):
        if is_legal:
            interests_value = trim_activities_droit(
                interests_rewritten,
                cv_is_short=cv_is_short,
                cv_is_long=cv_is_long,
            )
        else:
            interests_value = trim_activities(
                interests_rewritten,
                cv_is_long=cv_is_long,
                cv_is_short=cv_is_short,
            )
    else:
        interests_value = []
    
    sections["SKILLS"] = normalize_skills_block(sections.get("SKILLS", []), payload)
        
    mapping = {
        "%%FULL_NAME%%": full_name,
        "%%CONTACT_LINE%%": contact_line,
        "%%CV_TITLE%%": cv_title,
        "%%EDUCATION%%": sections.get("EDUCATION", []),
        "%%EXPERIENCE%%": sections.get("EXPERIENCES", []),
        "%%SKILLS%%": sections.get("SKILLS", []),
        "%%LANGUAGES%%": sections.get("LANGUAGES", []),
        "%%INTERESTS%%": interests_value,
    }

    for ph, value in mapping.items():
        p = _find_paragraph_containing(doc, ph)
        if not p:
            continue

        _strip_blank_neighbors(doc, p, before=2, after=2)
        _clear_paragraph(p)

        # ------- COMPÉTENCES & OUTILS -------
        if ph == "%%SKILLS%%" and isinstance(value, list):
            _render_skills(p, value or [], payload=payload)
            _remove_paragraph(p)
            continue

        # ------- ACTIVITÉS / CENTRES D'INTÉRÊT -------
        if ph == "%%INTERESTS%%" and isinstance(value, list):
            if not (value or []):
                # on récupère d'abord les paragraphes et la position du placeholder
                paras = list(doc.paragraphs)
                idx = None
                for i, pp in enumerate(paras):
                    if pp is p:
                        idx = i
                        break
            
                # supprime le titre juste avant s'il existe
                if idx is not None and idx - 1 >= 0:
                    prev_p = paras[idx - 1]
                    prev_text = (prev_p.text or "").strip().upper()
                    if "ACTIVITÉS" in prev_text:
                        _remove_paragraph(prev_p)
            
                # supprime ensuite le placeholder
                _remove_paragraph(p)
                continue
        
        
            _render_interests(p, value or [])
            _remove_paragraph(p)
            continue

        # ------- FORMATION -------
        if ph == "%%EDUCATION%%" and isinstance(value, list):

            # 🔹 CAS 1 : format structuré avec DEGREE:/SCHOOL:/LOCATION:/DATES:/DETAILS:
            if any((line or "").strip().startswith("DEGREE:") for line in value):
                programs = parse_education_structured(value)
                anchor = p
                first_edu = True

                source_courses_blocks = extract_source_courses_by_education_block(payload.get("education", ""))
                extra_detail_blocks = extract_non_course_details_by_education_block(payload.get("education", ""))

                for idx, edu in enumerate(programs):
                    degree = (edu.get("degree") or "").strip()
                    school = (edu.get("school") or "").strip()
                    location = (edu.get("location") or "").strip()
                    raw_education = payload.get("education", "")
                    if location:
                        # Matching flexible : on compare juste la ville (avant la virgule) pour gérer "Lyon, France" vs "lyon"
                        city_part = re.sub(r"[,\s]+", " ", location.lower()).strip().split()[0] if location else ""
                        raw_edu_low = re.sub(r"[,\s]+", " ", raw_education.lower())
                        if city_part and city_part not in raw_edu_low:
                            location = ""
                    dates = (edu.get("dates") or "").strip()
                    details = edu.get("details") or []

                    source_courses = source_courses_blocks[idx] if idx < len(source_courses_blocks) else []
                    extra_details = extra_detail_blocks[idx] if idx < len(extra_detail_blocks) else []

                    details = filter_education_details(
                        details,
                        payload.get("education", ""),
                        is_legal=is_legal
                    )

                    merged_details = []

                    # 1) priorité absolue aux détails utilisateur
                    for d in extra_details:
                        d = clean_punctuation_text((d or "").strip())
                        if not d:
                            continue
                        if is_course_detail_line(d):
                            continue
                        merged_details.append(d)
                    
                    # 2) on ajoute UNE SEULE ligne matières depuis l'input utilisateur
                    if source_courses:
                        course_line = "Matières fondamentales : " + ", ".join(source_courses) + "."
                        merged_details.append(course_line)
                    
                    # 3) on ne prend les détails LLM QUE s'il n'y a rien côté utilisateur
                    if not merged_details:
                        for d in details:
                            d = clean_punctuation_text((d or "").strip())
                            if not d:
                                continue
                            if is_course_detail_line(d):
                                continue
                            merged_details.append(d)
                    
                    details = dedupe_preserve_order(merged_details)

                    # ✅ garde les classements fournis par l'utilisateur
                    user_edu_text = (payload.get("education") or "").lower()
                    has_user_ranking = any(word in user_edu_text for word in ["classement", "top", "rank", "mention", "major"])
                    details = [
                        d for d in details
                        if not re.search(r"(?i)classement|rank|top\s*\d+", d)
                        or has_user_ranking
                    ]
                    details = dedupe_preserve_order(details)

                    # ✅ fallback : si aucun détail n'existe, on ajoute une ligne courte pour éviter le trou visuel
                    if not details:
                        details = []    

                    # Création du tableau 2 colonnes
                    table = _add_table_after(anchor, rows=1, cols=2)
                    
                    left = table.cell(0, 0)
                    right = table.cell(0, 1)
                    left.text = ""
                    right.text = ""

                    # ---- Colonne gauche : diplôme + école + détails ----
                    lp = left.paragraphs[0]
                    _keep_lines(lp, keep_lines=True, keep_next=True)
                    try:
                        lp.style = doc.styles["Normal"]
                    except Exception:
                        pass
                    lp.paragraph_format.left_indent = Pt(0)
                    lp.paragraph_format.first_line_indent = Pt(0)

                    mention_value = ""
                    deg_low = (degree or "").lower()
                    if "mention" in deg_low:
                        parts = [p.strip() for p in degree.split("–")]
                        kept = []
                        for part in parts:
                            if part.lower().startswith("mention"):
                                mention_value = part.replace("Mention", "").strip()
                            else:
                                kept.append(part)
                        degree = " – ".join(kept).strip()

                    degree_clean = degree.strip()
                    school_clean = school.strip()

                    # ✅ Pour les échanges : l'école est le titre, "Semestre d'échange" en détail
                    EXCHANGE_LABELS = [
                        "exchange semester", "exchange program", "échange académique",
                        "semester abroad", "study abroad", "semestre d'échange",
                        "visiting student", "programme d'échange", "program d'échange",
                    ]
                    is_exchange = any(kw in degree_clean.lower() for kw in EXCHANGE_LABELS)

                    if is_exchange and school_clean:
                        title_line = school_clean
                        school_line = ""
                        # Injecter "Semestre d'échange" comme 1er détail si absent
                        if not any("échange" in d.lower() or "exchange" in d.lower() for d in details):
                            details = ["Semestre d'échange"] + [d for d in details if d.strip()]
                    elif degree_clean and school_clean and school_clean.lower() in degree_clean.lower():
                        title_line = degree_clean
                        school_line = ""
                    else:
                        title_line = degree_clean
                        school_line = school_clean

                    if title_line:
                        r_title = lp.add_run(title_line)
                        r_title.bold = True
                        r_title.font.size = Pt(11)

                    # École sur ligne séparée en italique
                    if school_line:
                        para_school = left.add_paragraph()
                        para_school.paragraph_format.space_before = Pt(0)
                        para_school.paragraph_format.space_after = Pt(0)
                        try:
                            para_school.style = doc.styles["Normal"]
                        except Exception:
                            pass
                        r_school = para_school.add_run(school_line)
                        r_school.italic = True
                        r_school.font.size = Pt(11)

                    if mention_value:
                        para_m = left.add_paragraph()
                        para_m.paragraph_format.space_before = Pt(0)
                        para_m.paragraph_format.space_after = Pt(0)
                        try:
                            para_m.style = doc.styles["Normal"]
                        except Exception:
                            pass
                        r1 = para_m.add_run("Mention :")
                        r1.underline = True
                        r1.font.size = Pt(11)
                        r2 = para_m.add_run(" " + mention_value)
                        r2.font.size = Pt(11)

                    # Détails sous le titre
                    is_last_program = (idx == len(programs) - 1)
                    detail_list = [d for d in details if (d or "").strip()]
                    for d_idx, d in enumerate(detail_list):
                        text = (d or "").strip()
                        if not text:
                            continue
                    
                        # ✅ On supprime BDE/Association dans EDUCATION (car ça va dans EXPERIENCES)
                        low = text.lower()
                        if "bde" in low or low.startswith("association"):
                            continue

                        para = left.add_paragraph()
                        para.paragraph_format.space_before = Pt(0)
                        # Sur le dernier détail du dernier bloc éducation, ajouter espace avant EXPÉRIENCES
                        is_last_detail = (d_idx == len(detail_list) - 1)
                        para.paragraph_format.space_after = Pt(4) if (is_last_program and is_last_detail) else Pt(0)
                        try:
                            para.style = doc.styles["Normal"]
                        except Exception:
                            pass
                        
                        para.paragraph_format.left_indent = Pt(0)
                        para.paragraph_format.first_line_indent = Pt(0)

                        label_text = None
                        after_text = None
                        lower = text.lower()

                        # Fix orthographe/accord
                        text = text.replace("Analyse financières", "Analyse financière")
                        lower = text.lower()
                        
                        # ✅ 1) Projets (avec ou sans ":") => label souligné
                        if lower.startswith("projets"):
                            label_text = "Projets"
                            after_text = re.sub(r"(?i)^projets(\s+de\s+groupe)?\s*", "", text).strip()
                            # enlève les ponctuations type ": :"
                            after_text = re.sub(r"^[\s:–-]+", "", after_text).strip()
                        
                        # ✅ 2) "Cours en ..." => Matières fondamentales
                        elif re.match(r"(?i)^cours\s+en\s+", text):
                            label_text = "Matières fondamentales"
                            after_text = re.sub(r"(?i)^cours\s+en\s+", "", text).strip().rstrip(".")
                        elif lower.startswith("cours") and ":" in text:
                            label_text = "Matières fondamentales"
                            _, _, after = text.partition(":")
                            after_text = after.strip().rstrip(".")
                        
                        # ✅ 3) Matières fondamentales / cours pertinents / key coursework
                        elif "matières fondamentales" in lower or "cours pertinents" in lower or "key coursework" in lower:
                            label_text = "Matières fondamentales"
                            if ":" in text:
                                _, _, after = text.partition(":")
                                after_text = after.strip()
                        
                        # ✅ 4) Autres labels courts "X: Y"
                        elif ":" in text:
                            before, sep, after = text.partition(":")
                            before_clean = before.strip()
                            if len(before_clean.split()) <= 4:
                                label_text = before_clean
                                after_text = after.strip()

                        if label_text:
                            r1 = para.add_run(label_text + " :")
                            r1.underline = True
                            r1.font.size = Pt(11)
                            if after_text and after_text.strip():
                                r2 = para.add_run(" " + after_text.strip())
                                r2.font.size = Pt(11)
                        else:
                            r = para.add_run(text)
                            r.font.size = Pt(11)

                    # ---- Colonne droite : dates + lieu ----
                    rp = right.paragraphs[0]
                    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    rp.paragraph_format.space_after = Pt(0)

                    if dates:
                        clean_date = dates.replace("\r", " ").replace("\n", " ")
                        clean_date = re.sub(r"\s+", " ", clean_date.strip())
                        clean_date = translate_months_fr(clean_date)
                        clean_date = clean_date.replace(" - ", " – ")
                        clean_date = clean_date.replace(" ", "\u00A0")
                        r_date = rp.add_run(clean_date)
                        r_date.italic = True
                        r_date.font.size = Pt(9)

                    if location:
                        rp.add_run("\n")
                        r_loc = rp.add_run(location.strip())
                        r_loc.italic = True
                        r_loc.font.size = Pt(9)

                    # ✅ spacer entre deux formations
                    if idx < len(programs) - 1:
                        spacer_elt = OxmlElement("w:p")
                        table._tbl.addnext(spacer_elt)
                        spacer = Paragraph(spacer_elt, p._parent)
                        spacer.paragraph_format.space_before = Pt(0)
                        spacer.paragraph_format.space_after = ITEM_SPACING
                        anchor = spacer
                    else:
                        # ✅ spacer léger entre dernière formation et titre EXPÉRIENCES
                        spacer_elt = OxmlElement("w:p")
                        table._tbl.addnext(spacer_elt)
                        spacer = Paragraph(spacer_elt, p._parent)
                        spacer.paragraph_format.space_before = Pt(0)
                        spacer.paragraph_format.space_after = Pt(0)
                        spacer.paragraph_format.line_spacing = 1.0
                        anchor = spacer
                
                _remove_paragraph(p)
                continue

            # 🔹 CAS 2 : ancien format libre (on garde ton ancien comportement)
            anchor = p

            # 1) Regrouper les lignes par formation (blocs séparés par ligne vide)
            blocks = []
            current_block = []
            for line in value:
                text = (line or "").rstrip()
                if not text:
                    if current_block:
                        blocks.append(current_block)
                        current_block = []
                else:
                    current_block.append(text)
            if current_block:
                blocks.append(current_block)

            # 2) Découper les blocs s'il y a plusieurs diplômes collés
            split_blocks = []
            for b in blocks:
                split_blocks.extend(_split_education_block_on_degree_titles(b))

            # 3) Tri du plus récent au plus ancien
            blocks_sorted = sorted(split_blocks, key=_education_end_year, reverse=True)

            # 4) Gestion du bac (on peut le masquer)
            non_bac_blocks = [b for b in blocks_sorted if not _is_bac_block(b)]

            # ✅ Si CV trop court : on garde le bac même si normal (mieux que d'inventer)
            # 4) Gestion du bac (on peut le masquer)
            non_bac_blocks = [b for b in blocks_sorted if not _is_bac_block(b)]
            
            # ✅ Si CV trop court : on garde le bac même si normal
            if len(non_bac_blocks) <= 1:
                filtered_blocks = blocks_sorted[:]
            elif cv_is_short and len(non_bac_blocks) == 1:
                filtered_blocks = blocks_sorted[:]
            else:
                # ✅ Sinon : on garde le bac uniquement s'il est "élite"
                filtered_blocks = []
                for b in blocks_sorted:
                    if _is_bac_block(b) and not _keep_bac_block(b):
                        continue
                    filtered_blocks.append(b)

            # 5) Pour chaque formation -> tableau 1 ligne / 2 colonnes
            for i, block in enumerate(filtered_blocks):
                if not block:
                    continue

                first_line = block[0]

                # Normalisation des termes d'échange
                lower_first = first_line.lower()
                if "exchange semester" in lower_first or "exchange program" in lower_first:
                    first_line = re.sub(r"(?i)exchange semester", "Échange académique", first_line)
                    first_line = re.sub(r"(?i)exchange program", "Échange académique", first_line)
                if "study abroad" in lower_first:
                    first_line = re.sub(r"(?i)study abroad", "Échange académique", first_line)

                # Séparation Titre / Dates en cherchant un VRAI intervalle de dates en fin de ligne
                title_part = first_line
                date_part = ""

                date_range_patterns = [
                    r"(Janv|Fév|Fev|Mars|Avr|Mai|Juin|Juil|Août|Aout|Sept|Oct|Nov|Déc|Dec)\s+\d{4}\s*[–-]\s*(Janv|Fév|Fev|Mars|Avr|Mai|Juin|Juil|Août|Aout|Sept|Oct|Nov|Déc|Dec)\s+\d{4}\s*$",
                    r"(0[1-9]|1[0-2])/\d{4}\s*[–-]\s*(0[1-9]|1[0-2])/\d{4}\s*$",
                    r"(19|20)\d{4}?\s*[–-]\s*(19|20)\d{4}?\s*$"
                ]

                for pat in date_range_patterns:
                    m = re.search(pat, first_line)
                    if m:
                        date_part = m.group(0).strip()
                        title_part = first_line[:m.start()].rstrip(" ,–-").strip()
                        break

                if not date_part:
                    for sep in ("–", "—", "-"):
                        idx = first_line.rfind(sep)
                        if idx != -1:
                            title_part = first_line[:idx].strip()
                            date_part = first_line[idx + 1:].strip()
                            break

                if date_part:
                    m = re.search(r"(19|20)\d{2}\s*$", title_part)
                    if m:
                        title_part = title_part[:m.start()].rstrip(" ,–-")

                table = _add_table_after(anchor, rows=1, cols=2)
                
                left = table.cell(0, 0)
                right = table.cell(0, 1)
                left.text = ""
                right.text = ""

                lp = left.paragraphs[0]
                try:
                    lp.style = doc.styles["Normal"]
                except Exception:
                    pass
                lp.paragraph_format.left_indent = Pt(0)
                lp.paragraph_format.first_line_indent = Pt(0)

                # ✅ Si "Mention ..." est dans le titre, on la sort pour la mettre en dessous
                mention_value = ""
                if "mention" in title_part.lower():
                    parts = [p.strip() for p in title_part.split("–")]
                    kept = []
                    for part in parts:
                        if part.lower().startswith("mention"):
                            mention_value = part.replace("Mention", "").strip()
                        else:
                            kept.append(part)
                    title_part = " – ".join(kept).strip()

                # ✅ Séparer école du programme : dernier segment après " – " si c'est un nom d'école
                school_part_cas2 = ""
                school_keywords = ["school", "university", "université", "institut", "college",
                                   "essca", "edhec", "hec", "em lyon", "emlyon", "audencia",
                                   "kedge", "skema", "neoma", "grenoble em", "iseg", "ieseg",
                                   "sciences po", "dauphine", "assas", "sorbonne", "panthéon"]
                if " – " in title_part:
                    parts_edu = title_part.split(" – ")
                    last_part = parts_edu[-1].strip()
                    if any(k in last_part.lower() for k in school_keywords):
                        school_part_cas2 = last_part
                        title_part = " – ".join(parts_edu[:-1]).strip()
                elif " - " in title_part:
                    parts_edu = title_part.split(" - ")
                    last_part = parts_edu[-1].strip()
                    if any(k in last_part.lower() for k in school_keywords):
                        school_part_cas2 = last_part
                        title_part = " - ".join(parts_edu[:-1]).strip()

                title_run = lp.add_run(title_part)
                title_run.bold = True
                title_run.font.size = Pt(11)

                # ✅ École sur ligne séparée en italique
                if school_part_cas2:
                    para_school2 = left.add_paragraph()
                    para_school2.paragraph_format.space_before = Pt(0)
                    para_school2.paragraph_format.space_after = Pt(0)
                    try:
                        para_school2.style = doc.styles["Normal"]
                    except Exception:
                        pass
                    r_school2 = para_school2.add_run(school_part_cas2)
                    r_school2.italic = True
                    r_school2.font.size = Pt(11)
                
                # ✅ Ligne en dessous : Mention : (souligné)
                if mention_value:
                    para = left.add_paragraph()
                    try:
                        para.style = doc.styles["Normal"]
                    except Exception:
                        pass
                    r1 = para.add_run("Mention :")
                    r1.underline = True
                    r1.font.size = Pt(11)
                    r2 = para.add_run(" " + mention_value)
                    r2.font.size = Pt(11)

                location = ""
                location_index = -1
                for idx_line, raw in enumerate(block):
                    t = (raw or "").strip()
                    lower_t = t.lower()
                    if not t:
                        continue

                    candidate = None
                    if "," in t:
                        parts = [pp.strip() for pp in t.split(",")]
                        if len(parts) == 2 and len(parts[0].split()) <= 3:
                            candidate = t
                    else:
                        if len(t.split()) <= 3:
                            bad_tokens = [
                                "cours", "course", "key", "ranked",
                                "mention", "option", "majeure",
                                "matières", "matieres", "gpa"
                            ]
                            if not any(bt in lower_t for bt in bad_tokens):
                                candidate = t

                    if candidate:
                        location = candidate
                        location_index = idx_line
                        break

                for idx_line, line in enumerate(block[1:], start=1):
                    if idx_line == location_index:
                        continue

                    text = (line or "").strip()
                    if not text:
                        continue

                    para = left.add_paragraph()
                    try:
                        para.style = doc.styles["Normal"]
                    except Exception:
                        pass
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.first_line_indent = Pt(0)

                    label_text = None
                    after_text = None

                    if ":" in text:
                        before, sep, after = text.partition(":")
                        before_clean = before.strip()
                        lower_before = before_clean.lower()

                        if "cours pertinents" in lower_before:
                            label_text = "Matières fondamentales"
                            after_text = after or ""
                        else:
                            word_count = len(before_clean.split())
                            keywords = [
                                "gpa", "hl", "matières", "matieres",
                                "option", "majeure",
                                "spécialité", "specialite",
                            ]
                            if word_count <= 4 or any(k in lower_before for k in keywords):
                                label_text = before_clean
                                after_text = after or ""

                    if label_text:
                        r1 = para.add_run(label_text + " :")
                        r1.underline = True
                        r1.font.size = Pt(11)
                        if after_text and after_text.strip():
                            r2 = para.add_run(" " + after_text.strip())
                            r2.font.size = Pt(11)
                    else:
                        run = para.add_run(text)
                        run.font.size = Pt(11)

                rp = right.paragraphs[0]
                rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                rp.paragraph_format.space_after = Pt(0)

                if date_part:
                    clean_date = date_part.replace("\r", " ").replace("\n", " ")
                    clean_date = re.sub(r"\s+", " ", clean_date.strip())
                    clean_date = translate_months_fr(clean_date)
                    clean_date = clean_date.replace(" - ", " – ")
                    clean_date = clean_date.replace(" ", "\u00A0")
                    r_date = rp.add_run(clean_date)
                    r_date.italic = True
                    r_date.font.size = Pt(9)

                if location:
                    rp.add_run("\n")
                    r_loc = rp.add_run(location.strip())
                    r_loc.italic = True
                    r_loc.font.size = Pt(9)
                    rp.paragraph_format.space_after = Pt(0)

                if i < len(filtered_blocks) - 1:
                    new_p_elt = OxmlElement("w:p")
                    table._tbl.addnext(new_p_elt)
                    anchor = Paragraph(new_p_elt, p._parent)
                    anchor.paragraph_format.space_after = ITEM_SPACING
                    anchor.paragraph_format.space_before = Pt(0)
                else:
                    # ✅ spacer léger entre dernière formation et titre EXPÉRIENCES
                    new_p_elt = OxmlElement("w:p")
                    table._tbl.addnext(new_p_elt)
                    anchor = Paragraph(new_p_elt, p._parent)
                    anchor.paragraph_format.space_after = Pt(0)
                    anchor.paragraph_format.space_before = Pt(0)
                    anchor.paragraph_format.line_spacing = 1.0

            # ⚠️ NE PAS supprimer anchor
            _remove_paragraph(p)

            # ✅ Forcer space_before=0 sur le titre EXPÉRIENCES PROFESSIONNELLES
            # (même s'il a été traité par normalize avant l'insertion des tables,
            #  certains styles héritent un space_before résiduel)
            for dp in doc.paragraphs:
                if (dp.text or "").strip().upper() in {
                    "EXPÉRIENCES PROFESSIONNELLES", "EXPERIENCES PROFESSIONNELLES"
                }:
                    dp.paragraph_format.space_before = Pt(2)
                    dp.paragraph_format.space_after = Pt(1)
                    break

            continue

        # ------- EXPÉRIENCES PROFESSIONNELLES -------
        if ph == "%%EXPERIENCE%%":
            exps_from_cv = parse_finance_experiences(value or [])

            if is_legal:
                exps = exps_from_cv if exps_from_cv else parse_raw_experiences_input(payload.get("experiences", ""))
                exps = trim_experiences_droit(exps, is_cv_long=cv_is_long, is_cv_short=cv_is_short)

            elif is_audit:
                exps = exps_from_cv if exps_from_cv else parse_raw_experiences_input(payload.get("experiences", ""))
                exps = trim_experiences_audit(exps, is_cv_long=cv_is_long, is_cv_short=cv_is_short)

            elif is_management_sector(payload.get("sector", "")):
                exps = exps_from_cv if exps_from_cv else parse_raw_experiences_input(payload.get("experiences", ""))
                exps = trim_experiences_management(exps, is_cv_long=cv_is_long, is_cv_short=cv_is_short)

            else:
                exps = exps_from_cv if exps_from_cv else parse_raw_experiences_input(payload.get("experiences", ""))
                exps = trim_finance_experiences(exps, is_cv_long=cv_is_long)

            # ✅ Séparer les expériences pro des jobs étudiants
            pro_exps = [e for e in exps if not is_student_job_exp(e)]
            student_exps = [e for e in exps if is_student_job_exp(e)]

            # ✅ Trier chaque groupe par date de début décroissante (plus récent en premier)
            pro_exps = sorted(pro_exps, key=lambda e: extract_start_year_month(e.get("dates", "")), reverse=True)
            student_exps = sorted(student_exps, key=lambda e: extract_start_year_month(e.get("dates", "")), reverse=True)

            # Combiner : pro d'abord, jobs étudiants ensuite (avec flag pour le séparateur)
            has_student_section = bool(student_exps)
            exps = pro_exps + student_exps

            anchor = p
            first_table = True
            student_separator_inserted = False

            # Si jamais le modèle n'a pas respecté le format structuré,
            # on retombe sur un simple rendu en liste pour ne pas tout casser.
            if not exps:
                _insert_lines_after(p, value or [], make_bullets=True)
                continue

            # Mots-clés qui correspondent plutôt à un type de contrat qu'à un vrai rôle
            CONTRACT_PREFIXES = [
                "stagiaire", "stage",
                "summer job", "part-time job", "student job",
                "volunteering", "volunteer",
                "internship", "intern", "traineeship",
                "apprenticeship",
                "full-time", "full time",
                "part-time", "part time",
            ]

            for idx, exp in enumerate(exps):
                raw_role = (exp.get("role") or "").strip()
                role = normalize_role_text(raw_role)

                # ✅ Insérer le séparateur "Autres expériences" avant le 1er job étudiant
                if (
                    has_student_section
                    and pro_exps  # on affiche le séparateur seulement s'il y a des exps pro avant
                    and not student_separator_inserted
                    and is_student_job_exp(exp)
                ):
                    student_separator_inserted = True
                    # Petit paragraphe label entre les deux groupes
                    sep_elt = OxmlElement("w:p")
                    if hasattr(anchor, "_tbl"):
                        anchor._tbl.addnext(sep_elt)
                    else:
                        anchor._p.addnext(sep_elt)
                    sep_para = Paragraph(sep_elt, p._parent)
                    sep_para.paragraph_format.space_before = Pt(4)
                    sep_para.paragraph_format.space_after = Pt(1)
                    r_sep = sep_para.add_run("Autres expériences")
                    r_sep.italic = True
                    r_sep.font.size = Pt(9)
                    anchor = sep_para

                if is_legal and raw_role:
                    raw_role_low = raw_role.lower()
                    if "stagiaire" in raw_role_low and "jurid" in raw_role_low:
                        role = raw_role
                
                raw_experiences_input = payload.get("experiences", "").lower()
                if role and role.lower() not in raw_experiences_input:
                    original_role = (exp.get("role") or "").strip()
                    if original_role.lower() not in raw_experiences_input:
                        # fallback fort : on essaie de récupérer le rôle depuis l'input brut parsé
                        parsed_original_exps = parse_raw_experiences_input(payload.get("experiences", ""))
                        for original_exp in parsed_original_exps:
                            original_company = (original_exp.get("company") or "").strip().lower()
                            current_company = (exp.get("company") or "").strip().lower()
                            if original_company and current_company and original_company == current_company:
                                role = normalize_role_text((original_exp.get("role") or "").strip())
                                break

                if len(role.strip()) <= 8 or role.strip().lower() in {
                    "rh", "audit", "finance", "juridique", "juridique contentieux", "service juridique"
                }:
                    parsed_original_exps = parse_raw_experiences_input(payload.get("experiences", ""))
                    for original_exp in parsed_original_exps:
                        original_company = (original_exp.get("company") or "").strip().lower()
                        current_company = (exp.get("company") or "").strip().lower()
                        original_role = (original_exp.get("role") or "").strip()
                        if original_company and current_company and original_company == current_company and original_role:
                            role = normalize_role_text(original_role)
                            break
                    parsed_original_exps = parse_raw_experiences_input(payload.get("experiences", ""))
                    for original_exp in parsed_original_exps:
                        original_company = (original_exp.get("company") or "").strip().lower()
                        current_company = (exp.get("company") or "").strip().lower()
                        original_role = (original_exp.get("role") or "").strip()
                        if original_company and current_company and original_company == current_company and original_role:
                            role = normalize_role_text(original_role)
                            break

                # 1) Cas du type "Stage en audit financier" -> on vire "Stage + en/dans/au/aux"
                if not is_legal:
                    role = re.sub(
                        r"^(stage|stagiaire|internship|intern|traineeship)\s+(en|dans|au|aux)\s+",
                        "",
                        role,
                        flags=re.IGNORECASE,
                    ).strip()

                lower_role = role.lower()

                # 2) Si le rôle commence encore par un type de contrat (hors "en ..."),
                #    on enlève juste ce préfixe, mais on garde la suite.
                for key in CONTRACT_PREFIXES:
                    if lower_role.startswith(key + " "):
                        role = role[len(key):].lstrip(" -–—")
                        lower_role = role.lower()
                        break

                # 3) Cas particulier "Student tutor"
                if "student tutor" in lower_role:
                    role = role.replace("Student tutor", "Tuteur bénévole").replace("student tutor", "Tuteur bénévole")

                # 4) On force une majuscule au début du rôle si besoin
                if role and role[0].islower():
                    role = role[0].upper() + role[1:]

                company = (exp.get("company") or "").strip()

                # ✅ petit espace entre le TITRE de section et la 1ère expérience (sans ligne vide)
                if first_table:
                    try:
                        anchor.paragraph_format.space_after = ITEM_SPACING
                        anchor.paragraph_format.space_before = Pt(0)
                    except Exception:
                        pass
                anchor_for_table = anchor
                
                # Tableau 2 colonnes (mêmes tailles qu'avant via _add_table_after)
                table = _add_table_after(anchor_for_table, rows=1, cols=2)
                
                # ✅ On supprime UNIQUEMENT le placeholder la première fois
                if first_table:
                    try:
                        _remove_paragraph(anchor)
                    except Exception:
                        pass
                    first_table = False
                    
                    
                left = table.cell(0, 0)
                right = table.cell(0, 1)
                left.text = ""
                right.text = ""

                # ----- Colonne gauche : rôle + entreprise (2 lignes) + bullets -----
                lp = left.paragraphs[0]
                _keep_lines(lp, keep_lines=True, keep_next=True)
                try:
                    lp.style = doc.styles["Normal"]
                except Exception:
                    pass
                lp.paragraph_format.left_indent = Pt(0)
                lp.paragraph_format.first_line_indent = Pt(0)
                lp.paragraph_format.space_after = Pt(0)

                # Ligne 1 : rôle en GRAS
                if role:
                    title_run = lp.add_run(role)
                    title_run.bold = True
                    title_run.font.size = Pt(11)

                # Ligne 2 : entreprise en ITALIQUE (comme l'école sous le diplôme)
                if company:
                    para_company = left.add_paragraph()
                    para_company.paragraph_format.space_before = Pt(0)
                    para_company.paragraph_format.space_after = Pt(1)
                    try:
                        para_company.style = doc.styles["Normal"]
                    except Exception:
                        pass
                    para_company.paragraph_format.left_indent = Pt(0)
                    para_company.paragraph_format.first_line_indent = Pt(0)
                    r_company = para_company.add_run(company)
                    r_company.italic = True
                    r_company.font.size = Pt(11)

                bullets = (exp.get("bullets") or [])[:3]
                for b in bullets:
                    if not b:
                        continue
                
                    if is_legal:
                        b = soften_legal_overclaiming(b.strip())
                    else:
                        b = soften_overclaiming(b.strip())

                    b = clean_punctuation_text(b)
                    b_clean = b.strip().lower()
                
                    if b_clean in {"n/a", "na", "not applicable", "non applicable", "non-applicable"}:
                        continue
                
                    bp = left.add_paragraph()
                    try:
                        bp.style = doc.styles["Normal"]
                    except Exception:
                        pass
                    # Tiret long (–) au lieu de puce ronde
                    r_dash = bp.add_run("–  ")
                    r_dash.font.size = Pt(11)
                    r_content = bp.add_run(b)
                    r_content.font.size = Pt(11)

                    bp.paragraph_format.space_after = Pt(0)
                    bp.paragraph_format.left_indent = Pt(8)
                    bp.paragraph_format.first_line_indent = Pt(-8)
                    _keep_lines(bp, keep_lines=True, keep_next=False)

                # ----- Colonne droite : dates, lieu, type -----
                rp = right.paragraphs[0]
                rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                rp.paragraph_format.space_after = Pt(0)

                dates_raw = (exp.get("dates") or "").strip()
                if dates_raw:
                    clean_date = dates_raw.replace("\r", " ").replace("\n", " ")
                    clean_date = re.sub(r"\s+", " ", clean_date.strip())
                    clean_date = translate_months_fr(clean_date)
                    clean_date = clean_date.replace(" - ", " – ")
                    clean_date = clean_date.replace(" ", "\u00A0")  # espaces insécables
                    r_date = rp.add_run(clean_date)
                    r_date.italic = True
                    r_date.font.size = Pt(9)

                location = (exp.get("location") or "").strip()
                raw_experiences = payload.get("experiences", "")
                if location:
                    city_part = location.split(",")[0].strip().lower()
                    raw_exp_low = raw_experiences.lower()
                    if city_part and city_part not in raw_exp_low:
                        location = ""

                if location:
                    rp.add_run("\n")
                    r_loc = rp.add_run(location)
                    r_loc.italic = True
                    r_loc.font.size = Pt(9)

                type_raw = (exp.get("type") or "").strip()
                type_ = normalize_contract_type(type_raw)
                if type_:
                    rp.add_run("\n")
                    r_type = rp.add_run(type_)
                    r_type.italic = True
                    r_type.font.size = Pt(9)

                # ✅ spacer UNIQUEMENT entre deux expériences
                if idx < len(exps) - 1:
                    spacer_elt = OxmlElement("w:p")
                    table._tbl.addnext(spacer_elt)
                    spacer = Paragraph(spacer_elt, p._parent)
                    spacer.paragraph_format.space_before = Pt(0)
                    spacer.paragraph_format.space_after = ITEM_SPACING
                    anchor = spacer
                else:
                    # ❌ pas d'anchor vide après la dernière expérience
                    anchor = p
    

            _remove_paragraph(p)
            continue

        # ------- Texte simple (nom, titre, contact) -------
        if isinstance(value, str):
            run = p.add_run(value)
            if ph == "%%FULL_NAME%%":
                run.bold = True
                run.font.size = Pt(20)
            elif ph == "%%CV_TITLE%%":
                run.bold = True
                run.font.size = Pt(12)
            elif ph == "%%CONTACT_LINE%%":
                run.font.size = Pt(10)
            continue

        # ------- Listes classiques (si jamais) -------
        _insert_lines_after(p, value or [], make_bullets=True)
        
    # Nettoyage des paragraphes vides en fin de document pour éviter la page blanche
    try:
        for p in reversed(doc.paragraphs):
            if (p.text or "").strip():
                break
            _remove_paragraph(p)
    except Exception:
        pass
    collapse_blank_paragraphs(doc, max_consecutive=1)
    doc.save(out_path)

def convert_docx_to_pdf(docx_path: str, pdf_path: str) -> None:
    """
    Convertit un DOCX en PDF via LibreOffice.
    Rend le PDF identique au template Word.
    """
    import os
    import subprocess
    import shutil

    out_dir = os.path.dirname(pdf_path) or "."
    os.makedirs(out_dir, exist_ok=True)

    # Sur Linux/Docker, la commande peut être "soffice" ou "libreoffice"
    cmd = None
    for candidate in ["soffice", "libreoffice"]:
        if shutil.which(candidate):
            cmd = candidate
            break
    if not cmd:
        raise RuntimeError("LibreOffice/soffice introuvable dans l'environnement.")

    subprocess.run(
        [
            cmd,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to", "pdf",
            "--outdir", out_dir,
            docx_path,
        ],
        check=True,
    )

    generated_pdf = os.path.join(
        out_dir,
        os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
    )

    if generated_pdf != pdf_path:
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
        os.rename(generated_pdf, pdf_path)

def make_download_urls(job_id: str) -> Dict[str, str]:
    token = (jobs.get(job_id) or {}).get("download_token", "")
    token_param = f"?token={token}" if token else ""
    return {
        "pdf":  f"{PUBLIC_BASE_DOWNLOAD}/download/{job_id}/cv.pdf{token_param}",
        "docx": f"{PUBLIC_BASE_DOWNLOAD}/download/{job_id}/cv.docx{token_param}",
    }

@app.get("/quota")
def quota_check(email: str):
    email = normalize_email(email)
    if not email:
        raise HTTPException(status_code=400, detail="Email manquant.")
    current_month = month_key()
    try:
        with db_conn() as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT month FROM quota WHERE email = %s", (email,))
                row = cur.fetchone()
    except Exception:
        raise HTTPException(status_code=503, detail="DB unavailable")
    if not row or row[0] != current_month:
        return {"ok": True, "free": True, "message": "✅ Tu as encore ton CV gratuit ce mois-ci."}
    return {"ok": True, "free": False, "message": "ℹ️ Ton CV gratuit du mois est déjà utilisé. Le prochain sera payant."}

@app.post("/start")
async def start(payload: Dict[str, Any], request: Request):
    # ✅ Rate limiting par IP : max 10 générations par heure
    client_ip = request.headers.get("X-Forwarded-For", request.client.host if request.client else "unknown").split(",")[0].strip()
    if not _check_ip_rate_limit(client_ip, "start", max_hits=10, window_seconds=3600):
        raise HTTPException(status_code=429, detail="Trop de générations depuis cette adresse IP. Réessaie dans une heure.")

    # Emails de test — bypass vérification email ET quota
    DEV_WHITELIST = {
        "louis.bonnamour@essca.eu",
        "viktoria.aureau--bobillon@essca.eu"
    }

    # Vérification que l'email a bien été vérifié côté backend
    email_check = normalize_email(payload.get("email") or "")
    if email_check not in DEV_WHITELIST:
        if email_check not in _verified_emails or dt.datetime.utcnow() > _verified_emails[email_check]:
            raise HTTPException(status_code=403, detail="Email non vérifié. Veuillez vérifier votre email avant de générer un CV.")
        _verified_emails.pop(email_check, None)  # usage unique

    required = ["email", "sector", "company", "role", "job_posting", "full_name", "city", "phone"]

    for k in required:
        if not payload.get(k):
            raise HTTPException(status_code=400, detail=f"Champ manquant: {k}")

    # Limite anti-abus et anti-prompt injection
    if len(payload.get("job_posting", "")) > 8000:
        raise HTTPException(status_code=400, detail="Offre d'emploi trop longue.")
    if len(payload.get("experiences", "")) > 5000:
        raise HTTPException(status_code=400, detail="Expériences trop longues.")
    if len(payload.get("education", "")) > 3000:
        raise HTTPException(status_code=400, detail="Formation trop longue.")

    email = normalize_email(payload["email"])

    # Validation email basique anti-bot
    if len(email) > 200 or "@" not in email or "." not in email.split("@")[-1]:
        raise HTTPException(status_code=400, detail="Email invalide.")

    # Emails exemptés du quota (testeurs internes)
    if email in DEV_WHITELIST:
        job_id = await generate_and_store(payload)
        return {"mode": "free", "downloads": make_download_urls(job_id)}

    current_month = month_key()

    # Vérifie et consomme le quota de façon atomique
    with db_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """INSERT INTO quota (email, month)
                   VALUES (%s, %s)
                   ON CONFLICT (email)
                   DO UPDATE SET month = EXCLUDED.month
                   WHERE quota.month != EXCLUDED.month
                   RETURNING email""",
                (email, current_month)
            )
            inserted = cur.fetchone()

    if inserted is None:
        raise HTTPException(
            status_code=402,
            detail="CV gratuit déjà utilisé. Paiement requis."
        )

    job_id = await generate_and_store(payload)
    return {"mode": "free", "downloads": make_download_urls(job_id)}

@app.post("/create-checkout")
async def create_checkout(payload: Dict[str, Any], request: Request):
    # Rate limiting par IP
    client_ip = request.headers.get("X-Forwarded-For", request.client.host).split(",")[0].strip()
    _check_ip_rate_limit(client_ip)
    """
    Crée une session Stripe Checkout sécurisée.
    Le payload CV est stocké en mémoire côté serveur.
    Le frontend reçoit uniquement l'URL de paiement.
    """
    if not STRIPE_SECRET:
        raise HTTPException(status_code=500, detail="Stripe non configuré.")

    plan = payload.pop("plan", "unite")  # "unite" ou "mensuel"
    price_id = STRIPE_PRICE_MENSUEL if plan == "mensuel" else STRIPE_PRICE_UNITE

    # Validation minimale du payload
    required = ["email", "sector", "company", "role", "job_posting", "full_name", "city", "phone"]
    for k in required:
        if not payload.get(k):
            raise HTTPException(status_code=400, detail=f"Champ manquant: {k}")

    email = normalize_email(payload["email"])
    app_url = os.getenv("APP_URL", "https://mycvcopilote.com")

    try:
        session = stripe.checkout.Session.create(
            payment_method_types=["card"],
            mode="payment" if plan == "unite" else "subscription",
            line_items=[{"price": price_id, "quantity": 1}],
            customer_email=email,
            success_url=f"{app_url}/success.html?session_id={{CHECKOUT_SESSION_ID}}",
            cancel_url=f"{app_url}/app.html",
            metadata={"email": email, "plan": plan},
        )
    except stripe.error.StripeError as e:
        raise HTTPException(status_code=500, detail=f"Erreur Stripe : {str(e)}")

    # Stocker le payload côté serveur, associé à la session Stripe
    pending_stripe_sessions[session.id] = payload

    return {"checkout_url": session.url}


@app.post("/stripe-webhook")
async def stripe_webhook(request: Request):
    """
    Webhook Stripe — seul endroit où le paiement est confirmé.
    On vérifie la signature cryptographique pour être sûr que
    c'est bien Stripe qui envoie l'événement.
    """
    payload_bytes = await request.body()
    sig_header = request.headers.get("stripe-signature", "")
    webhook_secret = STRIPE_WEBHOOK_SECRET

    if not webhook_secret:
        raise HTTPException(status_code=500, detail="Webhook secret non configuré.")

    try:
        event = stripe.Webhook.construct_event(
            payload_bytes, sig_header, webhook_secret
        )
    except stripe.error.SignatureVerificationError:
        # Signature invalide = tentative de fraude
        raise HTTPException(status_code=400, detail="Signature invalide.")
    except Exception:
        raise HTTPException(status_code=400, detail="Webhook invalide.")

    # Paiement unique confirmé
    if event["type"] == "checkout.session.completed":
        session = event["data"]["object"]
        session_id = session["id"]

        cv_payload = pending_stripe_sessions.pop(session_id, None)
        if cv_payload:
            try:
                job_id = await generate_and_store(cv_payload)
                jobs[session_id] = jobs.get(job_id, {})
                jobs[session_id]["job_id"] = job_id
                jobs[session_id]["ready"] = True
            except Exception as e:
                print(f"=== ERREUR GÉNÉRATION après paiement {session_id}: {e} ===")

    return {"ok": True}


@app.get("/payment-status/{session_id}")
async def payment_status(session_id: str):
    """
    Appelé par le frontend sur la page de succès.
    Renvoie les liens de téléchargement quand le CV est prêt.
    """
    if session_id not in jobs:
        return {"ready": False}

    entry = jobs[session_id]
    if not entry.get("ready"):
        return {"ready": False}

    job_id = entry.get("job_id")
    if not job_id:
        return {"ready": False}

    cv_payload = jobs[session_id].get("payload") or {}
    suggested_filename = build_cv_filename(cv_payload)
    return {
        "ready": True,
        "downloads": make_download_urls(job_id),
        "filename": suggested_filename
    }

@app.get("/download/{job_id}/{filename}")
def download(job_id: str, filename: str, token: str = ""):
    from fastapi.responses import FileResponse

    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Inconnu.")

    # ✅ Vérification du token de téléchargement
    stored_token = jobs[job_id].get("download_token", "")
    if stored_token and token != stored_token:
        raise HTTPException(status_code=403, detail="Accès non autorisé.")

    payload = jobs[job_id].get("payload") or {}
    download_base = build_cv_filename(payload)

    if filename == "cv.pdf":
        path = jobs[job_id].get("pdf_path")
        download_name = f"{download_base}.pdf"
    elif filename == "cv.docx":
        path = jobs[job_id].get("docx_path")
        download_name = f"{download_base}.docx"
    else:
        raise HTTPException(status_code=404, detail="Fichier inconnu.")

    if not path or not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Fichier non prêt.")

    return FileResponse(path, filename=download_name)

async def generate_and_store(payload: Dict[str, Any], job_id: Optional[str] = None) -> str:
    async with _cv_semaphore:
        return await _generate_and_store_inner(payload, job_id)

async def _generate_and_store_inner(payload: Dict[str, Any], job_id: Optional[str] = None) -> str:
    job_id = job_id or str(uuid.uuid4())
    os.makedirs("out", exist_ok=True)

    base_filename = build_cv_filename(payload)
    internal_filename = f"{base_filename}_{job_id}"
    
    docx_path = os.path.join("out", f"{internal_filename}.docx")
    pdf_path = os.path.join("out", f"{internal_filename}.pdf")

    tpl = sector_to_template(payload["sector"])

    # 1) baseline
    cv_text = generate_cv_text(payload)
    last_action = None
    compact_mode = False
    expand_count = 0
    best_1page_text = None
    best_1page_fill = 0.0

    # 2) boucle max 5 essais (baseline + 2 corrections)
    for attempt in range(5):
        import asyncio
        try:
            await asyncio.to_thread(write_docx_from_template, tpl, cv_text, docx_path, payload=payload, compact_mode=compact_mode)
            await asyncio.to_thread(convert_docx_to_pdf, docx_path, pdf_path)
        except Exception as e:
            print(f"=== ERREUR GÉNÉRATION attempt {attempt}: {e} ===")
            if attempt >= 3:
                raise HTTPException(status_code=500, detail="Erreur lors de la génération du CV. Réessaie dans quelques secondes.")
            continue
            
        pages = pdf_page_count(pdf_path)
        fill = pdf_fill_ratio_first_page(pdf_path) if pages == 1 else 0.0
        print("attempt", attempt, "pages", pages, "fill", round(fill, 2))
        if pages == 1 and fill > best_1page_fill:
            best_1page_fill = fill
            best_1page_text = cv_text
        
        # 1) Trop long => revenir au meilleur résultat 1 page si dispo, sinon shrink léger
        if pages > 1:
            if best_1page_text:
                # ✅ On revient TOUJOURS au meilleur 1-page connu — le shrink LLM est trop risqué
                cv_text = best_1page_text
                await asyncio.to_thread(write_docx_from_template, tpl, cv_text, docx_path, payload=payload, compact_mode=compact_mode)
                await asyncio.to_thread(convert_docx_to_pdf, docx_path, pdf_path)
                break
            # Pas de best_1page (premier attempt déjà 2 pages) → shrink prudent
            if last_action == "shrink" and attempt >= 2:
                compact_mode = True
            else:
                cv_text = safe_apply_llm_edit(cv_text, llm_shrink_cv(cv_text), payload=payload, allow_drop_exp=True)
                last_action = "shrink"
            if attempt >= 2:
                compact_mode = True
            continue
    
        # 2) 1 page mais trop vide => expand
        # _is_short calculé sur le PAYLOAD (stable) pas sur cv_text qui peut être court si edu manquante
        raw_exp_payload = payload.get("experiences", "") or ""
        raw_edu_payload = payload.get("education", "") or ""
        payload_content = raw_exp_payload + raw_edu_payload
        payload_chars = len(re.sub(r"\s+", "", payload_content))
        payload_lines = payload_content.count("\n") + 1
        _is_short = (payload_chars < 900) or (payload_lines < 20)
        fill_threshold = 0.75 if _is_short else 0.93
        if pages == 1 and fill < fill_threshold:
            sector = payload.get("sector", "")
            max_expand = 6 if _is_short else 10
        
            if expand_count >= max_expand:
                break
        
            if is_legal_sector(sector):
                cv_text = safe_apply_llm_edit(cv_text, llm_expand_cv_droit(cv_text), payload=payload)
                last_action = "expand"
                expand_count += 1
                continue
        
            if is_audit_sector(sector):
                cv_text = safe_apply_llm_edit(cv_text, llm_expand_cv_audit(cv_text), payload=payload)
                last_action = "expand"
                expand_count += 1
                continue
        
            if is_management_sector(sector):
                cv_text = safe_apply_llm_edit(cv_text, llm_expand_cv_management(cv_text), payload=payload)
                last_action = "expand"
                expand_count += 1
                continue
        
            if is_finance_sector(sector):
                finance_max_expand = 5
                if expand_count >= finance_max_expand:
                    break
                cv_text = safe_apply_llm_edit(cv_text, llm_expand_cv(cv_text), payload=payload)
                last_action = "expand"
                expand_count += 1
                continue
        
            cv_text = safe_apply_llm_edit(cv_text, llm_expand_cv(cv_text), payload=payload)
            last_action = "expand"
            expand_count += 1
            continue
            
        # 3) OK
        break

    download_token = str(uuid.uuid4())
    jobs[job_id] = {"docx_path": docx_path, "pdf_path": pdf_path, "payload": payload, "download_token": download_token}
    # Sécurité finale : si encore 2 pages, on force un shrink compact
    try:
        if pdf_page_count(pdf_path) > 1:
            cv_text = safe_apply_llm_edit(cv_text, llm_shrink_cv(cv_text), payload=payload)
            await asyncio.to_thread(write_docx_from_template, tpl, cv_text, docx_path, payload=payload, compact_mode=True)
            await asyncio.to_thread(convert_docx_to_pdf, docx_path, pdf_path)
    except Exception:
        pass

    # Nettoyage des fichiers vieux de plus de 2 heures
    try:
        import time
        cutoff = time.time() - 7200
        for f in os.listdir("out"):
            fp = os.path.join("out", f)
            if os.path.isfile(fp) and os.path.getmtime(fp) < cutoff:
                os.remove(fp)
    except Exception:
        pass

    return job_id
    
import psycopg2
from psycopg2.extras import RealDictCursor
import psycopg2
import os

import os
import psycopg2
from fastapi import HTTPException

DATABASE_URL = os.getenv("DATABASE_URL", "")

import psycopg2.pool as _pg_pool

_db_pool = None

def _get_pool():
    global _db_pool
    if _db_pool is None:
        _db_pool = _pg_pool.SimpleConnectionPool(1, 10, dsn=DATABASE_URL)
    return _db_pool

from contextlib import contextmanager

@contextmanager
def db_conn():
    pool = _get_pool()
    conn = pool.getconn()
    try:
        yield conn
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        pool.putconn(conn)

# ============================================================
# VÉRIFICATION EMAIL PAR CODE — ZOHO SMTP
# ============================================================

# Anti-abus : nb de tentatives d'envoi par email
_send_attempts: Dict[str, list] = {}   # email -> liste de datetimes
_verify_attempts: Dict[str, int] = {}  # email -> nb de mauvais codes
_verified_emails: Dict[str, dt.datetime] = {}  # email -> expiration (1h après vérification)

def _check_send_rate_limit(email: str):
    """Bloque si l'email a demandé plus de 3 codes en 1 heure."""
    now = dt.datetime.utcnow()
    history = _send_attempts.get(email, [])
    # On garde seulement les demandes des 60 dernières minutes
    history = [t for t in history if (now - t).seconds < 3600]
    if len(history) >= 3:
        raise HTTPException(
            status_code=429,
            detail="Trop de tentatives. Attends 1 heure avant de redemander un code."
        )
    history.append(now)
    _send_attempts[email] = history


def send_verification_email(to_email: str, code: str):
    """Envoie le code par email via SMTP Brevo."""
    if not BREVO_LOGIN or not BREVO_PASSWORD:
        raise HTTPException(status_code=500, detail="Serveur mail non configuré.")

    msg = MIMEMultipart("alternative")
    msg["Subject"] = "Ton code de vérification MyCVCopilote"
    msg["From"] = SENDER_EMAIL
    msg["To"] = to_email

    html = f"""
    <html><body style="font-family:Arial,sans-serif;max-width:480px;margin:auto">
      <h2 style="color:#2563eb;">MyCVCopilote 👋</h2>
      <p>Voici ton code de vérification :</p>
      <div style="font-size:36px;font-weight:bold;letter-spacing:10px;
                  color:#1e293b;background:#f1f5f9;padding:20px;
                  border-radius:8px;text-align:center;">{code}</div>
      <p style="color:#64748b;font-size:13px;">
        Ce code est valable <strong>10 minutes</strong>.<br>
        Si tu n'as pas demandé ce code, ignore cet email.
      </p>
    </body></html>
    """
    msg.attach(MIMEText(html, "html"))

    try:
        with smtplib.SMTP("smtp-relay.brevo.com", 587, timeout=10) as server:
            server.starttls()
            server.login(BREVO_LOGIN, BREVO_PASSWORD)
            server.sendmail(SENDER_EMAIL, to_email, msg.as_string())
    except smtplib.SMTPAuthenticationError:
        raise HTTPException(status_code=500, detail="Erreur d'authentification mail.")
    except Exception as e:
        print(f"[MAIL ERROR] {e}")
        raise HTTPException(status_code=500, detail="Impossible d'envoyer l'email.")


@app.post("/send-verification-code")
async def send_verification_code(body: EmailRequest, request: Request):
    email = normalize_email(body.email or "")

    # Validation basique
    if not email or "@" not in email or "." not in email.split("@")[-1]:
        raise HTTPException(status_code=400, detail="Email invalide.")
    if len(email) > 200:
        raise HTTPException(status_code=400, detail="Email invalide.")

    # ✅ Rate limiting par IP : max 5 codes par IP par 10 minutes
    client_ip = request.headers.get("X-Forwarded-For", request.client.host if request.client else "unknown").split(",")[0].strip()
    if not _check_ip_rate_limit(client_ip, "send-code", max_hits=5, window_seconds=600):
        raise HTTPException(status_code=429, detail="Trop de tentatives. Réessaie dans 10 minutes.")

    # ✅ Vérification Cloudflare Turnstile
    turnstile_token = body.turnstile_token if hasattr(body, "turnstile_token") else (request.headers.get("X-Turnstile-Token", ""))
    if not await verify_turnstile(turnstile_token, client_ip):
        raise HTTPException(status_code=403, detail="Vérification anti-bot échouée. Recharge la page et réessaie.")

    # Anti-abus
    _check_send_rate_limit(email)

    # Génère un code à 6 chiffres
    code = str(random.randint(100000, 999999))
    expires = dt.datetime.utcnow() + dt.timedelta(minutes=10)

    # Stocke le code (écrase l'ancien si existant)
    email_verification_codes[email] = {
        "code": code,
        "expires": expires,
    }
    # Remet le compteur de mauvais codes à 0
    _verify_attempts[email] = 0

    # Envoie l'email
    send_verification_email(email, code)

    return {"ok": True, "message": "Code envoyé !"}


@app.post("/verify-code")
async def verify_code(body: VerifyCodeRequest):
    email = normalize_email(body.email or "")
    code = (body.code or "").strip()

    if not email or not code:
        raise HTTPException(status_code=400, detail="Email ou code manquant.")

    # Trop de mauvais essais ?
    if _verify_attempts.get(email, 0) >= 5:
        raise HTTPException(
            status_code=429,
            detail="Trop de mauvaises tentatives. Redemande un nouveau code."
        )

    entry = email_verification_codes.get(email)
    if not entry:
        raise HTTPException(status_code=400, detail="Aucun code pour cet email. Redemande un code.")

    # Code expiré ?
    if dt.datetime.utcnow() > entry["expires"]:
        del email_verification_codes[email]
        raise HTTPException(status_code=400, detail="Code expiré. Redemande un nouveau code.")

    # Mauvais code ?
    if entry["code"] != code:
        _verify_attempts[email] = _verify_attempts.get(email, 0) + 1
        raise HTTPException(status_code=400, detail="Code incorrect.")

    # ✅ Code valide → on nettoie tout
    del email_verification_codes[email]
    _verify_attempts.pop(email, None)
    _send_attempts.pop(email, None)

    # Marque l'email comme vérifié (valable 1 heure)
    _verified_emails[email] = dt.datetime.utcnow() + dt.timedelta(hours=1)

    return {"ok": True, "verified": True}
